# app.py
# Final, unified, and production-ready application code for the AI Ministry & Content Suite.
import pandas as pd
import streamlit as st
import os
import io
import re
import json
import wave
import asyncio
import time
from datetime import datetime

# --- Dependency Imports ---
import docx
import markdown2
import pypdf
from google import genai
from google.generativeai import types
from crewai import Agent, Task, Crew, Process, LLM
from crewai_tools import SerperDevTool, ScrapeWebsiteTool
import google.generativeai as gen
from streamlit import markdown


# ==============================================================================
## 1. Helper Functions
# ==============================================================================

def create_downloadable_docx(content):
    """Converts text content to a downloadable DOCX file in memory."""
    doc = docx.Document()
    for line in content.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def parse_json_from_text(text):
    """Safely extracts a JSON object from a string."""
    match = re.search(r'```json\n({.*?})\n```', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            st.error("AI returned an invalid JSON structure. Please try again.")
            return None
    st.warning("Could not parse the AI's JSON response. The structure might be incorrect.")
    return None

def parse_chapters_from_outline(outline_text):
    """Uses regex to find chapter titles in the generated outline."""
    if not isinstance(outline_text, str): return ["Invalid Outline Format"]
    chapters = re.findall(r"Chapter \d+[:\.]\s*(.*)", outline_text, re.IGNORECASE)
    if not chapters:
        chapters = re.findall(r"Kapitel \d+[:\.]\s*(.*)", outline_text, re.IGNORECASE)
    return [ch.strip() for ch in chapters] if chapters else ["Could not parse chapters automatically."]

@st.cache_data
def get_available_models(_api_key, task="generateContent"):
    """Fetches and caches the list of available Gemini models for a specific task."""
    if not _api_key: return []
    try:
        gen.configure(api_key=_api_key)
        models = [m.name for m in gen.list_models() if task in m.supported_generation_methods]
        if task == "text-to-speech":
            models = [m.replace("models:", "gemini") for m in models if '-tts' in m]
        if task == "video-generation":
            models = [m.replace("models/","") for m in models if 'veo' in m]
        if task == "image-generation":
            models = [m.replace("models/","") for m in models if 'image' in m]
        return sorted([name.replace("models", "gemini") for name in models])
    except Exception:
        st.sidebar.error("Error fetching models: Invalid API Key.")
        return []

def render_download_buttons(content, filename_base):
    """Renders a consistent set of download buttons for text-based content."""
    st.markdown("---")
    st.subheader("Export Your Content")
    col1, col2, col3, col4 = st.columns(4)
    html_content = markdown2.markdown(content)
    with col1:
        st.download_button("⬇️ DOCX", create_downloadable_docx(content), f"{filename_base}.docx")
    with col2:
        st.download_button("⬇️ Markdown", content.encode('utf-8'), f"{filename_base}.md")
    with col3:
        st.download_button("⬇️ HTML", html_content.encode('utf-8'), f"{filename_base}.html")
    with col4:
        st.download_button("⬇️ Text", content.encode('utf-8'), f"{filename_base}.txt")

def pcm_to_wav(pcm_data, channels=1, sample_width=2, sample_rate=24000):
    """Converts raw PCM data to a WAV file in memory."""
    buffer = io.BytesIO()
    with wave.open(buffer, 'wb') as wf:
        wf.setnchannels(channels)
        wf.setsampwidth(sample_width)
        wf.setframerate(sample_rate)
        wf.writeframes(pcm_data)
    buffer.seek(0)
    return buffer.getvalue()

# ==============================================================================
## 2. AI Crew Classes
# ==============================================================================
class SermonCrew:
    def __init__(self, model_name, topic, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.topic = topic
        self.language = language

    def run(self):
        agents = [
            Agent(role='Pentecostal Theologian', goal=f'Create a biblically sound outline for a sermon on "{self.topic}".', backstory="Experienced theologian skilled in structuring compelling sermons.", llm=self.llm, verbose=True),
            Agent(role='Bible Scripture Researcher', goal=f'Find relevant Bible verses for a sermon on "{self.topic}".', backstory="Meticulous Bible scholar dedicated to scriptural accuracy.", llm=self.llm, verbose=True),
            Agent(role='Gifted Pentecostal Preacher', goal=f'Write a complete, engaging sermon on "{self.topic}" in English.', backstory="Seasoned pastor known for powerful storytelling.", llm=self.llm, verbose=True),
            Agent(role='Expert Theological Translator', goal=f'Translate the final sermon accurately into {self.language}.', backstory=f"Professional translator specializing in theological texts, native in {self.language}.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f'Create a comprehensive outline for a sermon on "{self.topic}".', agent=agents[0])
        task2 = Task(description='Find relevant Bible verses for each point in the sermon outline.', agent=agents[1], context=[task1])
        task3 = Task(description='Write a complete sermon in English using the outline and scriptures.', agent=agents[2], context=[task1, task2])
        task4 = Task(description=f'Translate the final sermon into {self.language}.', agent=agents[3], context=[task3])
        
        crew = Crew(agents=agents, tasks=[task1, task2, task3, task4], process=Process.sequential, verbose=True)
        return crew.kickoff()

class FlyerCrew:
    def __init__(self, text_model, image_model, topic, text_element, flyer_type, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        os.environ["SERPER_API_KEY"] = st.session_state.get('serper_key', '')
        self.llm = LLM(model=text_model, temperature=0.8, api_key=os.environ["GOOGLE_API_KEY"])
        self.image_model = image_model
        self.topic = topic; self.text_element = text_element; self.flyer_type = flyer_type; self.language = language
    
    def run_design_crew(self):
        agents = [
            Agent(role='Creative Brief Specialist', goal='Develop a clear creative brief for a flyer.', backstory="Marketing expert skilled in distilling ideas into actionable briefs.", llm=self.llm, tools=[SerperDevTool()], verbose=True),
            Agent(role='Visual Concept Developer', goal='Brainstorm strong visual concepts based on a creative brief.', backstory="Seasoned Art Director with a modern aesthetic sense.", llm=self.llm, tools=[SerperDevTool()], verbose=True),
            Agent(role='Google Imagen Prompt Engineer', goal='Craft a detailed, effective image generation prompt for Google\'s Imagen model.', backstory="Technical artist who knows how to translate concepts into effective AI prompts.", llm=self.llm, verbose=True),
            Agent(role='Multilingual Social Media Copywriter', goal=f'Write a short, engaging social media post in {self.language} to accompany the flyer image.', backstory="Viral marketing specialist who crafts words that stop the scroll in multiple languages.", llm=self.llm, tools=[SerperDevTool()], verbose=True)
        ]
        task1 = Task(description=f"Analyze the request (Topic: '{self.topic}', Text: '{self.text_element}', Type: '{self.flyer_type}') to create a Creative Brief.", agent=agents[0])
        task2 = Task(description="Based on the brief, develop a full visual concept.", agent=agents[1], context=[task1])
        task3 = Task(description="Synthesize the brief and concept into a single, masterful image generation prompt.", agent=agents[2], context=[task2])
        task4 = Task(description=f"Based on the brief and concept, write a compelling social media post in {self.language}.", agent=agents[3], context=[task2])
        
        crew = Crew(agents=agents, tasks=[task1, task2, task3, task4], process=Process.sequential, verbose=True)
        result = crew.kickoff()
        return {"image_prompt": result.tasks_output[2].raw, "social_copy": result.tasks_output[3].raw}

    def generate_image(self, prompt):
        try:
            client = genai.Client()
            response = client.models.generate_images(model=self.image_model, prompt=prompt, config=types.GenerateImagesConfig(number_of_images=1))
            img_byte_arr = io.BytesIO()
            response.generated_images[0].image.save(img_byte_arr, format='PNG')
            return img_byte_arr.getvalue()
        except Exception as e:
            st.error(f"Image generation failed: {e}"); return None

class LanguageAcademyCrew:
    def __init__(self, model_name, native_language, target_language, level):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.native_language = native_language; self.target_language = target_language; self.level = level

    def run_curriculum_crew(self):
        import json
        agent = Agent(role='CEFR Curriculum Director', goal=f"Design a course outline for a {self.native_language} speaker learning {self.target_language} at {self.level}.", backstory="A lead curriculum designer for international language certification bodies.", llm=self.llm, verbose=True)
        file_name="outline.md"
        #st.dataframe(pd.json_normalize(file_name))
        task = Task(description=f"Generate a JSON with 'lessons' and 'final_exam' keys. For A1,"
                                f" first lesson must be 'The Alphabet & Core Pronunciation'.",
                    agent=agent, expected_output="A list of lessons and titles  fields lesson"
                                                 " and their content content formatted  as markdown  without ``´   ```",  output_file=file_name, markdown=True)

        result = Crew(agents=[agent], tasks=[task]).kickoff()
        st.markdown(result.raw)
        #st.dataframe(result.raw)
        return parse_json_from_text(result.raw)

    def run_lesson_crew(self, lesson_title):
        if lesson_title == "The Alphabet & Core Pronunciation":
            agent = Agent(role='Phonetics Specialist', goal=f"Create a guide to the {self.target_language} alphabet.", backstory=f"A linguist specializing in {self.target_language} phonology.", llm=self.llm, verbose=True)
            task = Task(description=f"Generate a guide to the {self.target_language} alphabet for a {self.native_language} speaker.", agent=agent)
            agents, tasks = [agent], [task]
        else:
            agents = [
                Agent(role='Grammar Professor', goal=f"Explain grammar for '{lesson_title}'.", backstory=f"A university professor known for making complex topics simple.", llm=self.llm, verbose=True),
                Agent(role='Vocabulary Linguist', goal=f"Create a vocabulary list for '{lesson_title}'.", backstory="A lexicographer who creates practical, thematic vocabulary lists.", llm=self.llm, verbose=True),
                Agent(role='Exercise Designer', goal=f"Create exercises for '{lesson_title}'.", backstory="An expert in crafting effective, engaging exercises for language textbooks.", llm=self.llm, verbose=True)
            ]
            tasks = [
                Task(description="Write a clear grammar explanation with examples.", agent=agents[0]),
                Task(description=f"Create a vocabulary table with translations to {self.native_language} and English.", agent=agents[1]),
                Task(description="Create 5-7 practice exercises and a separate answer key.", agent=agents[2])
            ]
        return Crew(agents=agents, tasks=tasks, process=Process.sequential, verbose=True).kickoff()

class BookStudioCrew:
    def __init__(self, model_name, topic, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        os.environ["SERPER_API_KEY"] = st.session_state.get('serper_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.topic = topic; self.language = language

    def run_outline_crew(self, user_prompt):
        agent = Agent(role='Chief Outline Architect', goal=f'Create a chapter-by-chapter outline for a book on "{self.topic}".', backstory='A seasoned developmental editor, you excel at structuring complex ideas.', llm=self.llm, tools=[SerperDevTool()], verbose=True)
        task = Task(description=f"Analyze book idea (Topic: '{self.topic}', Prompt: '{user_prompt}') to develop an outline.", agent=agent)
        return Crew(agents=[agent], tasks=[task]).kickoff()

    def run_chapter_crew(self, full_outline, selected_chapter):
        agents = [
            Agent(role='Research Specialist', goal='Gather detailed information for the assigned chapter.', backstory='A meticulous researcher with a Ph.D.', llm=self.llm, tools=[SerperDevTool(), ScrapeWebsiteTool()], verbose=True),
            Agent(role='Narrative Crafter', goal=f'Write an engaging, well-structured chapter in {self.language}.', backstory='A master storyteller who brings ideas to life.', llm=self.llm, verbose=True),
            Agent(role='Senior Editor', goal=f'Review and polish the drafted chapter in {self.language}.', backstory='A top-tier editor with a sharp eye for detail.', llm=self.llm, verbose=True)
        ]
        research_task = Task(description=f"Research the chapter: '{selected_chapter}'.", agent=agents[0])
        writing_task = Task(description=f"Write the chapter '{selected_chapter}'.", agent=agents[1], context=[research_task])
        editing_task = Task(description="Edit the chapter for publication.", agent=agents[2], context=[writing_task], output_file='book_chapter.md')
        crew = Crew(agents=agents, tasks=[research_task, writing_task, editing_task], process=Process.sequential, memory=True, embedder={"provider": "google", "config": {"model": "models/embedding-001"}}, verbose=True)
        crew.kickoff()
        with open('book_chapter.md', 'r', encoding='utf-8') as f:
            return f.read()

class ChefStudioCrew:
    def __init__(self, model_name, country, food_type, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.country = country; self.food_type = food_type; self.language = language

    def run_crew(self):
        agents = [
            Agent(role='World Cuisine Specialist', goal=f"Generate 3 meal ideas for {self.food_type} cuisine from {self.country}.", backstory="An acclaimed food historian who understands the soul of a country's food.", llm=self.llm, verbose=True),
            Agent(role='Executive Chef', goal=f"Write clear recipes in {self.language}.", backstory="A Michelin-trained chef with a passion for teaching home cooks.", llm=self.llm, verbose=True),
            Agent(role='Food Blogger', goal="Format recipes into markdown and create a separate list of image prompts. Return both as a JSON object.", backstory="A top-tier food blogger who crafts irresistible food content and expert AI image prompts.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f"Brainstorm 3 meal plans ({self.food_type}) from {self.country}.", agent=agents[0])
        task2 = Task(description=f"Write a full recipe for each meal plan in {self.language}.", agent=agents[1], context=[task1])
        task3 = Task(description="Create a JSON object with 'recipes_markdown' and 'image_prompts' keys.", agent=agents[2], context=[task2])
        result = Crew(agents=agents, tasks=[task1, task2, task3], process=Process.sequential, verbose=True).kickoff()
        return parse_json_from_text(result)

class PodcastStudioCrew:
    def __init__(self, model_name, country, language, topic, speaker1, speaker2):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.country = country; self.language = language; self.topic = topic
        self.speaker1 = speaker1; self.speaker2 = speaker2

    def run_script_crew(self):
        agents = [
            Agent(role='Cultural Researcher & Theologian', goal=f"Research '{self.topic}' for an audience in {self.country}.", backstory="A passionate pastor and scholar who connects timeless biblical truths to specific cultural contexts.", llm=self.llm, verbose=True),
            Agent(role='Engaging Podcast Scriptwriter', goal=f"Write a 10-minute script in {self.language} between {self.speaker1} and {self.speaker2}.", backstory="A gifted storyteller for a popular Christian podcast with billions of followers.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f"Develop 3-4 key talking points for a podcast on '{self.topic}'.", agent=agents[0])
        task2 = Task(description=f"Write a full podcast script using the key points, with clear labels for '{self.speaker1}:' and '{self.speaker2}:'.", agent=agents[1], context=[task1])
        return Crew(agents=agents, tasks=[task1, task2], process=Process.sequential, verbose=True).kickoff()

    @staticmethod
    def generate_audio(tts_model_name, transcript, speaker_configs):
        try:
            client = genai.Client()
            response = client.models.generate_content(
                model=tts_model_name, contents=transcript,
                config=types.GenerateContentConfig(response_modalities=["AUDIO"], speech_config=types.SpeechConfig(multi_speaker_voice_config=types.MultiSpeakerVoiceConfig(speaker_voice_configs=speaker_configs)))
            )
            return response.candidates[0].content.parts[0].inline_data.data
        except Exception as e:
            st.error(f"Audio generation failed: {e}"); return None

class BibleStudyCrew:
    def __init__(self, model_name, language, translation):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        os.environ["SERPER_API_KEY"] = st.session_state.get('serper_key', '')
        self.llm = LLM(model=model_name, temperature=0.5, api_key=os.environ["GOOGLE_API_KEY"])
        self.language = language
        self.translation = translation

    def run_book_study(self, english_book_name):
        agents = [
            Agent(role='Biblical Historian & Archaeologist', goal=f'Provide a comprehensive historical, cultural, and literary background for {english_book_name}, in {self.language}.', backstory="With a PhD from Jerusalem University, you provide the crucial context that makes the biblical text come alive.", llm=self.llm, tools=[SerperDevTool()], verbose=True),
            Agent(role='Exegetical Theologian', goal=f'Analyze the text of {english_book_name} to uncover its main theological themes, key verses, and structure, presenting findings in {self.language} using the {self.translation} translation for any scripture quotes.', backstory="As a systematic theologian, you are an expert at exegesis—drawing out the intended meaning of the text.", llm=self.llm, tools=[SerperDevTool()], verbose=True),
            Agent(role='Pastoral Guide & Counselor', goal=f'Create practical, thought-provoking application questions and prayer points based on the themes of {english_book_name}, written in {self.language}.', backstory="A seasoned pastor skilled in crafting questions that bridge the gap between ancient text and modern life.", llm=self.llm, verbose=True),
            Agent(role='Senior Editor for Christian Publishing', goal=f'Compile all sections into a single, cohesive, and beautifully formatted Bible study guide in {self.language}.', backstory="You work for an international Christian publishing house, ensuring every manuscript is professional and theologically sound.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f"Create the 'Historical Background' section for a study guide on **{english_book_name}**. Your output MUST be in {self.language}.", agent=agents[0])
        task2 = Task(description=f"Create the 'Theological Themes & Key Verses' section for **{english_book_name}**. Your output MUST be in {self.language}. Use the {self.translation} Bible translation for quotes.", agent=agents[1])
        task3 = Task(description=f"Create the 'Practical Application & Reflection' section for **{english_book_name}**. Your output MUST be in {self.language}.", agent=agents[2])
        task4 = Task(description=f"Compile all sections into a single study guide. The main title should be the {self.language} translation for 'A Study Guide to the Book of {english_book_name}'.", agent=agents[3], context=[task1, task2, task3], output_file='final_study_guide.md')
        
        crew = Crew(agents=agents, tasks=[task1, task2, task3, task4], process=Process.sequential, verbose=True)
        crew.kickoff()
        with open('final_study_guide.md', 'r', encoding='utf-8') as f:
            return f.read()

    def run_topic_study(self, topic, testament):
        agents = [
            Agent(role='Bible Scholar and Researcher', goal=f"Conduct a comprehensive search of the Bible to find at least 10 relevant verses for the topic: '{topic}'. The search must be limited to the {testament} Testament(s) and use the {self.translation} translation.", backstory="You are a meticulous and knowledgeable Bible scholar with a deep understanding of biblical languages and contexts.", tools=[SerperDevTool()], llm=self.llm, verbose=True),
            Agent(role='Pentecostal Pastor and Theologian', goal=f"Write a short, inspiring devotional or sermon outline based on the Bible verses provided. The message should be written in {self.language} and reflect a Pentecostal passion for Jesus and love for people.", backstory="You are a seasoned pastor with a gift for teaching.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f"Search for scriptures related to '{topic}' within the {testament} Testament(s) using the {self.translation} translation. Compile a list of the top 10-15 most relevant verses.", expected_output=f"A markdown-formatted list of 10-15 bible verses, each with its reference.", agent=agents[0])
        task2 = Task(description="Using the list of scriptures from the scholar, write an inspiring devotional.", expected_output=f"A complete devotional message in {self.language}, approximately 300-500 words long.", agent=agents[1], context=[task1])
        
        crew = Crew(agents=agents, tasks=[task1, task2], process=Process.sequential, verbose=True)
        return crew.kickoff()

class SchoolTutorCrew:
    def __init__(self, model_name, country, grade, subject, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.country = country; self.grade = grade; self.subject = subject; self.language = language

    def run(self, question):
        agents = [
            Agent(role='Curriculum Analyst', backstory="Expert in global K-12 education systems.", goal=f"Analyze the educational context for a {self.grade} student in {self.country} studying {self.subject}.", llm=self.llm, verbose=True),
            Agent(role=f'{self.subject.title()} Subject Matter Expert', backstory=f"Renowned teacher in {self.subject} with a passion for clarity.", goal=f"Accurately solve the student's homework question about {self.subject}.", llm=self.llm, verbose=True),
            Agent(role='Pedagogy and Language Expert', backstory="Master educator skilled at adapting complex information for different age groups.", goal=f"Rewrite the expert's solution into an engaging answer in {self.language} for a {self.grade} student.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f"Analyze context: Country {self.country}, Grade {self.grade}, Subject {self.subject}. Plan how to best explain the answer to: '{question}'", agent=agents[0])
        task2 = Task(description=f"Solve this homework question: '{question}'", agent=agents[1], context=[task1])
        task3 = Task(description=f"Take the expert's solution and rewrite it in {self.language} as a friendly, clear markdown explanation.", agent=agents[2], context=[task2])
        
        crew = Crew(agents=agents, tasks=[task1, task2, task3], process=Process.sequential, verbose=True)
        return crew.kickoff()

class UniversityTutorCrew:
    def __init__(self, model_name, course, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.course = course; self.language = language

    def run(self, question):
        agents = [
            Agent(role='University Curriculum Specialist', backstory="Academic advisor with encyclopedic knowledge of university course structures.", goal=f"Analyze the curriculum for a university course titled '{self.course}'.", llm=self.llm, verbose=True),
            Agent(role=f'University Professor of {self.course}', backstory=f"Distinguished professor with a Ph.D. and extensive research experience relevant to {self.course}.", goal="Provide a rigorous, technically correct solution to the student's question.", llm=self.llm, verbose=True),
            Agent(role='Senior Academic Tutor', backstory="Award-winning teaching assistant skilled at making complex ideas click.", goal=f"Refine the professor's solution into a comprehensive explanation in {self.language}.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f"Analyze the academic framework for the course '{self.course}' to answer the question: '{question}'.", agent=agents[0])
        task2 = Task(description=f"Provide an expert, in-depth solution to the question: '{question}'.", agent=agents[1], context=[task1])
        task3 = Task(description=f"Synthesize the solution into a high-quality tutorial explanation in {self.language}.", agent=agents[2], context=[task2])
        
        crew = Crew(agents=agents, tasks=[task1, task2, task3], process=Process.sequential, verbose=True)
        return crew.kickoff()

class NewsroomHQCrew:
    def __init__(self, model_name, scope, location, topics, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        os.environ["SERPER_API_KEY"] = st.session_state.get('serper_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.scope = scope; self.location = location; self.topics = topics; self.language = language

    def run(self):
        editor = Agent(role='Managing Editor', goal=f'Oversee the creation of a high-quality newspaper in {self.language}.', backstory="With decades of experience at major news outlets, you are the final word on journalistic integrity.", llm=self.llm, allow_delegation=True, verbose=True)
        wire_service = Agent(role='News Wire Service', goal='Continuously scan the web for the latest, most significant news stories.', backstory="You are the digital equivalent of the Associated Press, the first to know about any breaking event.", llm=self.llm, tools=[SerperDevTool()], verbose=True)
        reporters = [Agent(role=f'{topic.title()} Reporter', goal=f'Develop in-depth, accurate, and engaging news articles on {topic} in {self.language}.', backstory=f"You are a seasoned journalist with a deep specialization in {topic}.", llm=self.llm, tools=[SerperDevTool()], verbose=True) for topic in self.topics]

        query_location = self.location if self.scope in ["Local", "National"] else "world"
        fetch_task = Task(description=f"Fetch the most recent and significant news stories for today, {datetime.now().strftime('%Y-%m-%d')}, for a {self.scope} newspaper focused on {query_location}.", agent=wire_service)
        reporting_tasks = [Task(description=f"Using the news wire data, write a concise and compelling news article in {self.language} on your beat: '{topic}'.", agent=reporter, context=[fetch_task]) for reporter, topic in zip(reporters, self.topics)]
        editing_task = Task(description=f"Review all drafted articles and assemble them into a single, cohesive newspaper format in {self.language}.", agent=editor, context=reporting_tasks, output_file='final_newspaper.md')
        
        crew = Crew(agents=[editor, wire_service] + reporters, tasks=[fetch_task] + reporting_tasks + [editing_task], process=Process.sequential, verbose=True)
        crew.kickoff()
        with open('final_newspaper.md', 'r', encoding='utf-8') as f:
            return f.read()

class ViralVideoCrew:
    def __init__(self, model_name, topic_or_verse, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.topic_or_verse = topic_or_verse
        self.language = language

    def run(self):
        agents = [
            Agent(role='Viral Content Strategist', goal=f'Develop a high-level concept and narrative arc for a 45-second, 5-part video series about "{self.topic_or_verse}".', backstory='You are a master of digital storytelling, able to create compelling narratives that unfold across multiple short clips.', llm=self.llm, verbose=True),
            Agent(role='Viral Video Storyboard Artist', goal='Break down a narrative arc into 5 distinct, visually compelling 8-second video concepts.', backstory='You are a visual thinker, able to translate a story into a sequence of powerful, attention-grabbing shots.', llm=self.llm, verbose=True),
            Agent(role='Google VEO Prompt Engineer', goal=f'Write 5 unique, detailed VEO prompts in {self.language}, one for each part of a video storyboard.', backstory='You are an expert in generative video AI, knowing the precise keywords to achieve cinematic quality for a series of related clips.', llm=self.llm, verbose=True),
            Agent(role='Social Media Hook Writer', goal=f'Write 5 unique, engaging social media hooks in {self.language}, one for each video clip in a series.', backstory=f'You craft irresistible, scroll-stopping text that makes people want to watch the next part.', llm=self.llm, verbose=True),
            Agent(role='Content Editor', goal=f'Compile all generated content into a single, well-formatted Markdown document in {self.language}.', backstory='You are a meticulous editor who ensures the final output is clean, organized, and ready for the user.', llm=self.llm, verbose=True)
        ]
        
        task1 = Task(description=f'Based on the topic "{self.topic_or_verse}", create a cohesive narrative arc for a 45-second video series.', agent=agents[0])
        task2 = Task(description='Based on the narrative arc, create a detailed storyboard for the 5 video clips.', agent=agents[1], context=[task1])
        task3 = Task(description=f'For each of the 5 storyboard parts, write a unique and highly detailed VEO prompt in {self.language}.', agent=agents[2], context=[task2])
        task4 = Task(description=f'For each of the 5 storyboard parts, write a unique and compelling social media hook in {self.language}.', agent=agents[3], context=[task2])
        task5 = Task(description=f'Compile all sections into a single, clean Markdown document.', agent=agents[4], context=[task1, task2, task3, task4])

        crew = Crew(agents=agents, tasks=[task1, task2, task3, task4, task5], process=Process.sequential, verbose=True)
        return crew.kickoff()

class SingleVideoCrew:
    @staticmethod
    def generate_video(model_name, prompt):
        try:
            client = genai.Client()
            operation = client.models.generate_videos(model=model_name, prompt=prompt)
            
            status_placeholder = st.empty()
            with st.spinner("Generating video... This may take several minutes."):
                while not operation.done:
                    status_placeholder.info(f"Waiting for video generation... Status: {operation.state.name}")
                    time.sleep(10)
                    operation = client.operations.get(operation)
            
            status_placeholder.success("Video generation complete!")
            generated_video = operation.response.generated_videos[0]
            video_bytes = client.files.download(file=generated_video.video)
            return video_bytes
        except Exception as e:
            st.error(f"An error occurred during video generation: {e}")
            st.info("Please ensure your Google Cloud project has access to the VEO model.")
            return None

class AudioSuiteCrew:
    def __init__(self, model_name):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])

    def translate_text(self, text, language):
        agent = Agent(role='Expert Translator', goal=f"Translate the given text accurately into {language}.", backstory=f"You are a professional translator with expertise in many languages, ensuring natural and fluent translations.", llm=self.llm, verbose=True)
        task = Task(description=f"Translate the following text to {language}: '{text}'", agent=agent)
        crew = Crew(agents=[agent], tasks=[task], verbose=True)
        return crew.kickoff()

    @staticmethod
    def generate_audio(tts_model_name, text, voice_name):
        try:
            client = genai.Client()
            response = client.models.generate_content(
                model=tts_model_name, contents=[f"Say this with a clear and engaging tone: {text}"],
                config=types.GenerateContentConfig(response_modalities=["AUDIO"], speech_config=types.SpeechConfig(voice_config=types.VoiceConfig(prebuilt_voice_config=types.PrebuiltVoiceConfig(voice_name=voice_name))))
            )
            return response.candidates[0].content.parts[0].inline_data.data
        except Exception as e:
            st.error(f"Audio generation failed: {e}"); return None

    @staticmethod
    def transcribe_audio(audio_file):
        try:
            client = genai.Client()
            response = client.models.transcribe(model='models/chirp-2', media=[audio_file])
            return response.text
        except Exception as e:
            st.error(f"Audio transcription failed: {e}"); return None

# ==============================================================================
## 3. Page Rendering Functions
# ==============================================================================

def render_language_academy_page():
    st.title("🗣️ AI Language Academy")
    st.markdown("Generate a full curriculum, then study lesson by lesson.")
    if 'course_outline' not in st.session_state: st.session_state.course_outline = None
    if 'lesson_content' not in st.session_state: st.session_state.lesson_content = None
    available_models = get_available_models(st.session_state.get('gemini_key'))

    with st.form("curriculum_form"):
        st.header("Step 1: Design Your Course")
        col1, col2 = st.columns(2); native_language = col1.text_input("Your Language", "English"); target_language = col2.text_input("Language to Learn", "French")
        col3, col4 = st.columns(2); level = col3.selectbox("Select Level (CEFR)", ["A1", "A2", "B1", "B2", "C1", "C2"]); selected_model = col4.selectbox("Choose AI Model", available_models) if available_models else None
        if st.form_submit_button("Generate Course Outline", use_container_width=True):
            if not all([native_language, target_language, selected_model]): st.error("Please fill all fields.")
            else:
                with st.spinner("AI Director is designing your learning path..."):
                    st.session_state.update(lang_model=selected_model, native_lang=native_language, target_lang=target_language, lesson_content=None)
                    crew = LanguageAcademyCrew(selected_model, native_language, target_language, level)
                    st.session_state.course_outline = crew.run_curriculum_crew()
                    #st.markdown()

    if st.session_state.get('course_outline'):
        st.markdown("---"); st.header("Step 2: Choose a Lesson to Study")
        st.markdown(st.session_state.course_outline)

        outline = st.session_state.course_outline; lesson_options = outline.get('lessons', []) + ["Final Exam Preparation"]
        selected_lesson = st.text_area("Select a lesson", placeholder="Enter any of the above lessons you want to learn")
        if st.button(f"Generate Content for: {selected_lesson}", use_container_width=True):
            with st.spinner(f"AI teaching crew is preparing '{selected_lesson}'..."):
                crew = LanguageAcademyCrew(st.session_state.lang_model, st.session_state.native_lang, st.session_state.target_lang, level)
                if selected_lesson == "Final Exam Preparation":
                    exam_plan = outline.get('final_exam', {}); st.session_state.lesson_content = f"## Final Exam Guide\n\n**Description:** {exam_plan.get('description', 'N/A')}\n\n**Structure:** {exam_plan.get('structure', 'N/A')}\n\n**Tips:** {exam_plan.get('preparation_tips', 'N/A')}"
                else: st.session_state.lesson_content = crew.run_lesson_crew(selected_lesson)
    
    if st.session_state.get('lesson_content'):
        st.markdown("---"); st.header(f"Study Material: {selected_lesson}"); st.markdown(st.session_state.lesson_content)
        render_download_buttons(st.session_state.lesson_content, "language_lesson")

def render_book_page():
    st.title("📚 AI Book Writing Studio")
    st.markdown("Outline your book and then write it, one chapter at a time.")
    if 'book_outline' not in st.session_state: st.session_state.book_outline = None
    if 'chapter_content' not in st.session_state: st.session_state.chapter_content = None
    available_models = get_available_models(st.session_state.get('gemini_key'))

    with st.form("outline_form"):
        st.header("Step 1: Generate Your Book's Outline")
        language = st.selectbox("Language:", ["English", "German", "French", "Spanish"])
        topic = st.text_input("Book Topic:", placeholder="e.g., The History of Ancient Rome")
        user_prompt = st.text_area("Detailed Description:", height=150)
        selected_model = st.selectbox("Choose AI Model", available_models) if available_models else None
        if st.form_submit_button("Generate Outline"):
            if not all([topic, user_prompt, selected_model]): st.error("Please fill all fields.")
            else:
                st.session_state.update(book_topic=topic, book_language=language, book_model=selected_model)
                crew = BookStudioCrew(selected_model, topic, language)
                st.session_state.book_outline = crew.run_outline_crew(user_prompt)

    if st.session_state.get('book_outline'):
        st.markdown("---"); st.header("Step 2: Write Your Book, Chapter by Chapter")
        st.markdown(st.session_state.book_outline)
        chapter_list = parse_chapters_from_outline(st.session_state.book_outline)
        with st.form("chapter_form"):
            selected_chapter = st.selectbox("Choose a chapter to write:", chapter_list)
            if st.form_submit_button(f"Write '{selected_chapter}'"):
                with st.spinner("Initializing the writing crew..."):
                    crew = BookStudioCrew(st.session_state.book_model, st.session_state.book_topic, st.session_state.book_language)
                    st.session_state.chapter_content = crew.run_chapter_crew(st.session_state.book_outline, selected_chapter)
    
    if st.session_state.get('chapter_content'):
        st.markdown("---"); st.header("Your Newly Written Chapter"); st.markdown(st.session_state.chapter_content)
        render_download_buttons(st.session_state.chapter_content, "book_chapter")

def render_chef_page():
    st.title("🍳 AI Chef Studio")
    st.markdown("Get multilingual meal plans with recipes and AI image prompts.")
    if 'chef_recipes' not in st.session_state: st.session_state.chef_recipes = None
    if 'chef_prompts' not in st.session_state: st.session_state.chef_prompts = None
    available_models = get_available_models(st.session_state.get('gemini_key'))

    with st.form("chef_form"):
        st.subheader("What are you in the mood for?")
        col1, col2 = st.columns(2); country = col1.text_input("Country or Region", "Italy"); food_type = col2.selectbox("Food Type", ["Any", "Meat", "Vegetarian", "Vegan"])
        col3, col4 = st.columns(2); language = col3.text_input("Language for Recipe", "English"); selected_model = col4.selectbox("Choose AI Model", available_models) if available_models else None
        if st.form_submit_button("Generate Meal Ideas", use_container_width=True):
            if not all([country, language, selected_model]): st.error("Please fill all fields.")
            else:
                with st.spinner("👩‍🍳 The AI Chef Crew is crafting your menu..."):
                    crew = ChefStudioCrew(selected_model, country, food_type, language)
                    structured_result = crew.run_crew()
                    if structured_result:
                        st.session_state.chef_recipes = structured_result.get("recipes_markdown")
                        st.session_state.chef_prompts = structured_result.get("image_prompts")

    if st.session_state.get('chef_recipes'):
        st.markdown("---"); st.subheader("Your Custom Meal & Recipe Plan"); st.markdown(st.session_state.chef_recipes)
        render_download_buttons(st.session_state.chef_recipes, "recipe_plan")
    if st.session_state.get('chef_prompts'):
        st.markdown("---"); st.subheader("🎨 AI Image Prompts")
        for i, prompt in enumerate(st.session_state.chef_prompts):
            st.text_area(f"Prompt for Meal {i+1}", prompt, height=100, key=f"prompt_{i}")

def render_podcast_studio_page():
    st.title("🎙️ AI Podcast Studio")
    st.markdown("Generate a multi-speaker podcast from scratch with AI.")
    if 'podcast_transcript' not in st.session_state: st.session_state.podcast_transcript = ""
    text_models = get_available_models(st.session_state.get('gemini_key'), task="generateContent")
    tts_models = get_available_models(st.session_state.get('gemini_key'), task="text-to-speech")
    voices = ["Kore", "Puck", "Chipp", "Sadachbia", "Lyra", "Arpy", "Fable", "Onyx"]

    st.header("Step 1: Generate Script")
    with st.form("podcast_script_form"):
        col1, col2 = st.columns(2); topic = col1.text_input("Topic", "The Power of Forgiveness"); country = col2.text_input("Country", "USA")
        language = col1.text_input("Language", "English"); transcript_model = col2.selectbox("Transcript Model", text_models)
        st.markdown("**Speakers**"); col_s1, col_s2 = st.columns(2); speaker1_name = col_s1.text_input("Speaker 1", "Dr. Anya"); speaker2_name = col_s2.text_input("Speaker 2", "Liam")
        if st.form_submit_button("Generate Podcast Script", use_container_width=True):
            with st.spinner("AI Crew is writing your transcript..."):
                crew = PodcastStudioCrew(transcript_model, country, language, topic, speaker1_name, speaker2_name)
                st.session_state.podcast_transcript = crew.run_script_crew()
                st.session_state.update(speaker1_name=speaker1_name, speaker2_name=speaker2_name, podcast_topic=topic)
    
    if st.session_state.podcast_transcript:
        st.text_area("Review/Edit Transcript:", value=st.session_state.podcast_transcript, height=250, key="transcript_editor")
        render_download_buttons(st.session_state.transcript_editor, "podcast_script")
        st.markdown("---"); st.header("Step 2: Generate Audio")
        with st.form("podcast_audio_form"):
            s1_name = st.session_state.get('speaker1_name', 'Speaker 1'); s2_name = st.session_state.get('speaker2_name', 'Speaker 2')
            col1, col2, col3 = st.columns(3); speaker1_voice = col1.selectbox(f"Voice for {s1_name}", voices); speaker2_voice = col2.selectbox(f"Voice for {s2_name}", voices); tts_model = col3.selectbox("Audio Model", tts_models)
            if st.form_submit_button("Create Podcast Audio", use_container_width=True):
                with st.spinner("🎙️ AI is recording your podcast..."):
                    speaker_configs = [types.SpeakerVoiceConfig(speaker=s1_name, voice_config=types.VoiceConfig(prebuilt_voice_config=types.PrebuiltVoiceConfig(voice_name=speaker1_voice))), types.SpeakerVoiceConfig(speaker=s2_name, voice_config=types.VoiceConfig(prebuilt_voice_config=types.PrebuiltVoiceConfig(voice_name=speaker2_voice)))]
                    audio_data = PodcastStudioCrew.generate_audio(tts_model, st.session_state.transcript_editor, speaker_configs)
                    if audio_data:
                        st.success("Audio generated!"); st.audio(audio_data, format='audio/wav')
                        st.download_button("⬇️ Download Podcast (.wav)", audio_data, f"{st.session_state.podcast_topic.replace(' ', '_')}.wav")

def render_sermon_page():
    st.title("📖 AI Sermon Generator Crew")
    st.markdown("This application uses a team of AI agents to help you create a detailed, biblically-sound sermon.")
    available_models = get_available_models(st.session_state.get('gemini_key'))
    LANGUAGES = ("English", "German", "French", "Swahili", "Italian", "Spanish", "Portuguese")
    
    with st.form("sermon_form"):
        sermon_topic = st.text_input("Enter the Sermon Topic:", "The Power of Forgiveness")
        target_language = st.selectbox("Choose Sermon Language:", LANGUAGES)
        selected_model = st.selectbox("Choose a Gemini Model:", available_models) if available_models else None
        
        if st.form_submit_button("✨ Generate Sermon"):
            if not all([sermon_topic, target_language, selected_model]):
                st.error("Please fill all fields.")
            else:
                with st.spinner(f"The AI Sermon Crew is at work..."):
                    crew = SermonCrew(selected_model, sermon_topic, target_language)
                    st.session_state.sermon_content = crew.run()

    if 'sermon_content' in st.session_state and st.session_state.sermon_content:
        st.markdown("---"); st.subheader("Your Generated Sermon")
        st.markdown(st.session_state.sermon_content)
        render_download_buttons(st.session_state.sermon_content, "sermon")

def render_flyer_page():
    st.title("🚀 AI Flyer Production Studio")
    st.markdown("From idea to share-ready asset in minutes.")
    text_models = get_available_models(st.session_state.get('gemini_key'))
    image_models = get_available_models(st.session_state.get('gemini_key'), task="image-generation")
    LANGUAGES = ("English", "German", "French", "Swahili", "Italian", "Spanish", "Portuguese")

    with st.form("flyer_form"):
        st.header("Step 1: Describe Your Flyer")
        topic = st.text_input("Topic or Theme:", "Youth Camp")
        text_element = st.text_input("Key Text (Verse, Slogan):", "Summer Blast 2025")
        flyer_type = st.selectbox("Flyer Type:", ("Social Media Post (Square)", "Poster (Portrait)"))
        language = st.selectbox("Social Media Copy Language:", LANGUAGES)
        text_model = st.selectbox("Choose Text Model:", text_models) if text_models else None
        image_model = st.selectbox("Choose Image Model:", image_models) if image_models else None
        
        if st.form_submit_button("🚀 Generate Concept"):
            if not all([topic, text_element, text_model, image_model]):
                st.error("Please fill all fields.")
            else:
                with st.spinner("AI Design Studio is developing the concept..."):
                    crew = FlyerCrew(text_model, image_model, topic, text_element, flyer_type, language)
                    st.session_state.flyer_concept = crew.run_design_crew()

    if 'flyer_concept' in st.session_state and st.session_state.flyer_concept:
        st.markdown("---"); st.header("Step 2: Generate Image")
        concept = st.session_state.flyer_concept
        st.subheader("🎨 Generated Image Prompt")
        st.code(concept['image_prompt'], language='text')
        
        if st.button("Generate Flyer Image"):
            with st.spinner("Sending prompt to Imagen..."):
                crew = FlyerCrew(text_model, image_model, topic, text_element, flyer_type, language)
                st.session_state.flyer_image = crew.generate_image(concept['image_prompt'])

    if 'flyer_image' in st.session_state and st.session_state.flyer_image:
        st.markdown("---"); st.header("✅ Your Final Assets")
        st.image(st.session_state.flyer_image, caption="Generated Flyer")
        st.download_button("⬇️ Download Flyer", st.session_state.flyer_image, "flyer.png", "image/png")
        st.subheader(f"✍️ Your Social Media Caption in {language}")
        st.text_area("", concept['social_copy'], height=150)
        render_download_buttons(concept['social_copy'], "social_media_copy")

# def render_music_page():
#     st.title("🎶 AI Worship Song Studio")
#     st.markdown("Partner with our AI Worship Team to compose a song concept and generate a production-ready prompt for **Google's Lyria AI**.")
#     if 'lyria_prompt' not in st.session_state: st.session_state.lyria_prompt = None
#     available_models = get_available_models(st.session_state.get('gemini_key'))
#     LANGUAGES = ("English", "German", "French", "Swahili", "Italian", "Spanish", "Portuguese")
#
#     with st.form("music_form"):
#         st.header("Step 1: Share Your Inspiration")
#         col1, col2 = st.columns(2)
#         genre = col1.selectbox("Select the Musical Genre:", ("Worship (Hillsong/Bethel style)", "Praise (Elevation/Upbeat style)", "African Gospel Praise"))
#         topic = col1.text_input("Core Topic or Theme:", placeholder="e.g., God's Grace, Faithfulness")
#         verses = col2.text_area("Enter Bible Verses or Inspirational Text:", placeholder="e.g., John 3:16, Psalm 23", height=150)
#         target_language = col2.selectbox("Choose Lyrics Language:", LANGUAGES)
#         selected_model = st.selectbox("Choose a Gemini Model for Songwriting:", available_models) if available_models else None
#
#         if st.form_submit_button("Compose & Generate Lyria Prompt"):
#             if not selected_model: st.error("Please select a model.")
#             elif not topic and not verses: st.error("Please provide a Topic or Verses.")
#             else:
#                 with st.spinner("Your AI Worship Team is gathering..."):
#                     crew = MusicStudioCrew(selected_model, genre, verses, topic, target_language)
#                     st.session_state.lyria_prompt = crew.run_crew()
#
#     if st.session_state.lyria_prompt:
#         st.markdown("---"); st.header("Step 2: Generate Your Audio")
#         st.subheader("✅ Your Final Lyria Prompt")
#         st.code(st.session_state.lyria_prompt, language="text")
#         render_download_buttons(st.session_state.lyria_prompt, "lyria_prompt")
#
#         if st.button("🎵 Generate 30-Second Audio Track"):
#             st.info("The Lyria model for real-time music generation is experimental and may require special access through Google Cloud.")
#             st.warning("Audio generation with Lyria is a complex process not fully implemented in this public-facing script.")

def render_bible_book_study_page():
    st.title("📖 Bible Book Study Generator")
    st.markdown("Powered by a team of AI Theologians, Pastors, and Historians to create an in-depth guide to any book of the Bible.")
    
    ENGLISH_BOOKS = ["Genesis", "Exodus", "Leviticus", "Numbers", "Deuteronomy", "Joshua", "Judges", "Ruth", "1 Samuel", "2 Samuel", "1 Kings", "2 Kings", "1 Chronicles", "2 Chronicles", "Ezra", "Nehemiah", "Esther", "Job", "Psalms", "Proverbs", "Ecclesiastes", "Song of Solomon", "Isaiah", "Jeremiah", "Lamentations", "Ezekiel", "Daniel", "Hosea", "Joel", "Amos", "Obadiah", "Jonah", "Micah", "Nahum", "Habakkuk", "Zephaniah", "Haggai", "Zechariah", "Malachi", "Matthew", "Mark", "Luke", "John", "Acts", "Romans", "1 Corinthians", "2 Corinthians", "Galatians", "Ephesians", "Philippians", "Colossians", "1 Thessalonians", "2 Thessalonians", "1 Timothy", "2 Timothy", "Titus", "Philemon", "Hebrews", "James", "1 Peter", "2 Peter", "1 John", "2 John", "3 John", "Jude", "Revelation"]
    BIBLE_BOOKS_TRANSLATIONS = {"English": ENGLISH_BOOKS, "German": ["Genesis", "Exodus", "Levitikus", "Numeri", "Deuteronomium", "Josua", "Richter", "Ruth", "1. Samuel", "2. Samuel", "1. Könige", "2. Könige", "1. Chronik", "2. Chronik", "Esra", "Nehemia", "Esther", "Hiob", "Psalmen", "Sprüche", "Prediger", "Hohelied", "Jesaja", "Jeremia", "Klagelieder", "Hesekiel", "Daniel", "Hosea", "Joel", "Amos", "Obadja", "Jona", "Micha", "Nahum", "Habakuk", "Zefanja", "Haggai", "Sacharja", "Maleachi", "Matthäus", "Markus", "Lukas", "Johannes", "Apostelgeschichte", "Römer", "1. Korinther", "2. Korinther", "Galater", "Epheser", "Philipper", "Kolosser", "1. Thessalonicher", "2. Thessalonicher", "1. Timotheo", "2. Timotheo", "Titus", "Philemon", "Hebräer", "Jakobus", "1. Petrus", "2. Petrus", "1. Johannes", "2. Johannes", "3. Johannes", "Judas", "Offenbarung"]}
    available_models = get_available_models(st.session_state.get('gemini_key'))

    with st.form("bible_book_form"):
        st.header("1. Select Your Study Parameters")
        col1, col2 = st.columns(2)
        selected_language = col1.selectbox("Choose your language:", list(BIBLE_BOOKS_TRANSLATIONS.keys()))
        selected_book_translated = col2.selectbox("Choose a book to study:", BIBLE_BOOKS_TRANSLATIONS[selected_language])
        
        col3, col4 = st.columns(2)
        bible_translation = col3.text_input("Enter Bible Translation", placeholder="e.g., NIV, KJV, Hoffnung für alle")
        selected_model = col4.selectbox("Select Gemini Model", available_models) if available_models else None
        
        if st.form_submit_button(f"Create Study Guide for {selected_book_translated}"):
            if not all([selected_model, bible_translation]):
                st.error("❌ Please select a model and enter a Bible translation.")
            else:
                book_index = BIBLE_BOOKS_TRANSLATIONS[selected_language].index(selected_book_translated)
                english_book_name = ENGLISH_BOOKS[book_index]
                with st.spinner(f"Your AI Bible Study Team is preparing your guide..."):
                    crew = BibleStudyCrew(selected_model, selected_language, bible_translation)
                    st.session_state.study_guide_content = crew.run_book_study(english_book_name)

    if "study_guide_content" in st.session_state and st.session_state.study_guide_content:
        st.markdown("---"); st.header("3. Your Custom Study Guide"); st.markdown(st.session_state.study_guide_content)
        render_download_buttons(st.session_state.study_guide_content, f"{selected_book_translated.replace(' ', '_')}_study_guide")

def render_bible_topic_study_page():
    st.title("🙏 AI Bible Topic Study")
    st.markdown("Your personal theology research partner. Enter a topic to discover relevant scriptures and receive a custom devotional.")
    available_models = get_available_models(st.session_state.get('gemini_key'))

    with st.form("bible_topic_form"):
        st.header("Start Your Study")
        topic = st.text_input("Enter a Topic, Theme, or Question", placeholder="e.g., Faith, Forgiveness, Who is the Holy Spirit?")
        col1, col2 = st.columns(2)
        language = col1.text_input("Language for Results", "English")
        testament = col2.selectbox("Select Testament(s)", ["All", "Old Testament", "New Testament"])
        
        col3, col4 = st.columns(2)
        bible_translation = col3.text_input("Enter Bible Translation", placeholder="e.g., NIV, KJV, Hoffnung für alle")
        selected_model = col4.selectbox("Choose AI Model", available_models) if available_models else None
        
        if st.form_submit_button("Begin Bible Study", use_container_width=True):
            if not all([topic, language, selected_model, bible_translation]):
                st.error("Please fill all fields.")
            else:
                with st.spinner("The AI ministry team is studying God's Word for you..."):
                    crew = BibleStudyCrew(selected_model, language, bible_translation)
                    st.session_state.study_result = crew.run_topic_study(topic, testament)

    if "study_result" in st.session_state and st.session_state.study_result:
        st.markdown("---"); st.subheader(f"Study Results on '{topic}'"); st.markdown(st.session_state.study_result)
        render_download_buttons(st.session_state.study_result, f"bible_study_{topic.replace(' ', '_')}")

def render_news_page():
    st.title("📰 AI Newsroom Headquarters")
    st.markdown("Welcome, Editor-in-Chief! Commission a complete, up-to-the-minute digital newspaper from your AI journalist crew.")
    available_models = get_available_models(st.session_state.get('gemini_key'))
    LANGUAGES = ("English", "German", "French", "Swahili", "Italian", "Spanish", "Portuguese")

    with st.form("news_form"):
        st.header("Step 1: Define Your Newspaper's Focus")
        scope = st.selectbox("Select Newspaper Scope:", ["Global", "National", "Local"])
        location = ""
        if scope == "Local":
            location = st.text_input("Enter City:", "Nairobi")
        elif scope == "National":
            location = st.text_input("Enter Country:", "Kenya")

        st.markdown("**Select the sections to include:**")
        topic_options = ["Top Story", "Business & Stock Market", "Sports", "Technology", "Fashion & Trends"]
        selected_topics = [topic for topic in topic_options if st.checkbox(topic, True, key=f"topic_{topic}")]
        target_language = st.selectbox("Choose Newspaper Language:", LANGUAGES)
        selected_model = st.selectbox("Choose a Gemini Model for Reporting:", available_models) if available_models else None
        
        if st.form_submit_button("Assemble Today's Newspaper"):
            if not selected_model:
                st.error("❌ Please select a model.")
            elif not selected_topics:
                st.error("Please select at least one topic.")
            else:
                with st.spinner("Your AI Newsroom is on the story... This will take a few minutes."):
                    crew = NewsroomHQCrew(selected_model, scope, location, selected_topics, target_language)
                    st.session_state.newspaper_content = crew.run()
                    st.session_state.newspaper_location = location if location else scope

    if "newspaper_content" in st.session_state and st.session_state.newspaper_content:
        st.markdown("---")
        st.subheader(f"The {st.session_state.get('newspaper_location', 'Global')} Times")
        st.markdown(st.session_state.newspaper_content, unsafe_allow_html=True)
        render_download_buttons(st.session_state.newspaper_content, "newspaper")

def render_viral_video_page():
    st.title("🎬 AI Viral Video Series Studio")
    st.markdown("Create a powerful, 45-second video series concept for social media, broken into 5 compelling clips.")
    available_models = get_available_models(st.session_state.get('gemini_key'))
    LANGUAGES = ("English", "German", "French", "Swahili", "Italian", "Spanish", "Portuguese")

    with st.form("viral_video_form"):
        st.header("Step 1: What's Your Message?")
        topic_or_verse = st.text_input("Enter a Christian Topic or Bible Verse:", placeholder="e.g., John 3:16, Saved by Grace")
        target_language = st.selectbox("Choose Prompt & Hook Language:", LANGUAGES)
        selected_model = st.selectbox("Choose a Gemini Model for Video Concepting:", available_models) if available_models else None
        
        if st.form_submit_button("Generate Viral Video Series Concept"):
            if not all([topic_or_verse, selected_model]):
                st.error("🚨 Please provide a topic/verse and select a model.")
            else:
                with st.spinner("Your AI Social Media team is brainstorming a viral series..."):
                    crew = ViralVideoCrew(selected_model, topic_or_verse, target_language)
                    st.session_state.video_series_content = crew.run()

    if "video_series_content" in st.session_state and st.session_state.video_series_content:
        st.markdown("---")
        st.subheader("✅ Your Viral Video Series Concept")
        st.markdown(st.session_state.video_series_content)
        render_download_buttons(st.session_state.video_series_content, "viral_video_series_concept")

def render_single_video_page():
    st.title("📹 Single Video Studio")
    st.markdown("Unleash your creativity! Use AI to generate a compelling short video from your text prompt.")
    
    video_models = get_available_models(st.session_state.get('gemini_key'), task="video-generation")
    
    st.warning("**Cost Warning:** Generating videos with VEO models can be expensive. A single second of video can cost approximately **€0.75**. Please be mindful of prompt length and usage.", icon="⚠️")

    with st.form("single_video_form"):
        prompt = st.text_area("Enter your video prompt here:", height=150, placeholder="A photorealistic video of a hummingbird flying in slow motion in a tropical garden...")
        selected_model = st.selectbox("Choose VEO Model:", video_models) if video_models else None
        
        submitted = st.form_submit_button("Generate Video", use_container_width=True)

    if submitted:
        if not prompt:
            st.error("Please enter a prompt to generate a video.")
        elif not selected_model:
            st.error("No VEO models available. Check your API Key and access permissions.")
        else:
            video_bytes_content = SingleVideoCrew.generate_video(selected_model, prompt)
            if video_bytes_content:
                st.session_state.single_video = video_bytes_content

    if 'single_video' in st.session_state and st.session_state.single_video:
        st.markdown("---")
        st.subheader("Your Generated Video")
        st.video(st.session_state.single_video)
        st.download_button(
            label="⬇️ Download Video (.mp4)",
            data=st.session_state.single_video,
            file_name="generated_veo_video.mp4",
            mime="video/mp4"
        )

def render_school_tutor_page():
    st.title("🎓 AI Tutor (Grades 1-12)")
    st.markdown("Your personal AI learning assistant. Get step-by-step help with your homework.")
    available_models = get_available_models(st.session_state.get('gemini_key'))
    LANGUAGES = ("English", "German", "French", "Swahili", "Italian", "Spanish", "Portuguese")

    with st.form("school_tutor_form"):
        st.subheader("Tell us about your homework")
        col1, col2 = st.columns(2)
        country = col1.text_input("Country", "Germany")
        subject = col1.selectbox("Subject", ["Mathematics", "History", "Science", "Literature", "Physics"])
        grade = col2.selectbox("Grade / Class", [f"Grade {i}" for i in range(1, 14)])
        language = col2.selectbox("Language for Explanation", LANGUAGES)
        selected_model = st.selectbox("Choose a Gemini Model:", available_models) if available_models else None
        question = st.text_area("📝 Enter your question here", height=150)
        
        if st.form_submit_button("Get Help from AI Tutors!", use_container_width=True):
            if not all([question, selected_model]):
                st.error("Please enter a question and select a model.")
            else:
                with st.spinner("🚀 Your AI Tutors are working on it..."):
                    crew = SchoolTutorCrew(selected_model, country, grade, subject, language)
                    st.session_state.school_result = crew.run(question)

    if "school_result" in st.session_state and st.session_state.school_result:
        st.markdown("---"); st.subheader("✨ Here's your explanation:")
        st.markdown(st.session_state.school_result)
        render_download_buttons(st.session_state.school_result, "school_homework_help")

def render_university_tutor_page():
    st.title("🧑‍🏫 University AI Professor")
    st.markdown("Get expert-level academic help for your university courses.")
    available_models = get_available_models(st.session_state.get('gemini_key'))
    LANGUAGES = ("English", "German", "French", "Swahili", "Italian", "Spanish", "Portuguese")

    with st.form("uni_tutor_form"):
        st.subheader("Provide your course and question details")
        course = st.text_input("University Course Name", placeholder="e.g., Experimental Physics, Linear Algebra II")
        language = st.selectbox("Language for Explanation", LANGUAGES)
        selected_model = st.selectbox("Choose a Gemini Model:", available_models) if available_models else None
        question = st.text_area("📝 Enter your question or problem here", height=150)
        
        if st.form_submit_button("Consult the Professor", use_container_width=True):
            if not all([course, question, selected_model]):
                st.error("Please fill all fields and select a model.")
            else:
                with st.spinner("🚀 Consulting with the AI academic team..."):
                    crew = UniversityTutorCrew(selected_model, course, language)
                    st.session_state.uni_result = crew.run(question)

    if "uni_result" in st.session_state and st.session_state.uni_result:
        st.markdown("---"); st.subheader("✨ Professor's Explanation:")
        st.markdown(st.session_state.uni_result)
        render_download_buttons(st.session_state.uni_result, "university_homework_help")

def render_audio_suite_page():
    st.title("🎧 AI Audio Suite")
    st.markdown("Your one-stop shop for audio generation, transcription, and translation.")
    LANGUAGES = ("English", "German", "French", "Swahili", "Italian", "Spanish", "Portuguese")
    
    tab1, tab2, tab3 = st.tabs(["**Text-to-Audio**", "**Text Translation**", "**Audio Transcription & Translation**"])

    with tab1:
        st.subheader("Convert Text into High-Quality Audio")
        tts_models = get_available_models(st.session_state.get('gemini_key'), task="text-to-speech")
        voices = ["Kore", "Puck", "Chipp", "Sadachbia", "Lyra", "Arpy", "Fable", "Onyx"]

        with st.form("tts_form"):
            text_to_convert = st.text_area("Enter text to convert to audio:", height=150, key="text_for_audio")
            col1, col2 = st.columns(2)
            selected_voice = col1.selectbox("Choose a Voice:", voices)
            selected_tts_model = col2.selectbox("Choose Audio Model:", tts_models) if tts_models else None
            
            if st.form_submit_button("🎤 Generate Audio", use_container_width=True):
                if not text_to_convert: st.error("Please enter some text.")
                elif not selected_tts_model: st.error("Please select an audio model.")
                else:
                    with st.spinner("Generating audio..."):
                        audio_data = AudioSuiteCrew.generate_audio(selected_tts_model, text_to_convert, selected_voice)
                        if audio_data:
                            st.success("Audio Generated!")
                            wav_bytes = pcm_to_wav(audio_data)
                            st.audio(wav_bytes, format='audio/wav')
                            st.download_button("⬇️ Download Audio (.wav)", wav_bytes, "generated_audio.wav", "audio/wav")

    with tab2:
        st.subheader("Translate Text to a Different Language")
        text_models = get_available_models(st.session_state.get('gemini_key'))
        with st.form("translation_form"):
            text_to_translate = st.text_area("Enter text to translate:", height=150)
            col1, col2 = st.columns(2)
            translation_language = col1.selectbox("Translate to:", LANGUAGES, key="translate_lang_text")
            selected_text_model = col2.selectbox("Choose AI Model:", text_models) if text_models else None
            
            if st.form_submit_button("✍️ Translate Text", use_container_width=True):
                if not text_to_translate: st.error("Please enter text to translate.")
                elif not selected_text_model: st.error("Please select a model.")
                else:
                    with st.spinner(f"Translating to {translation_language}..."):
                        crew = AudioSuiteCrew(selected_text_model)
                        st.session_state.translated_text = crew.translate_text(text_to_translate, translation_language)
        
        if 'translated_text' in st.session_state and st.session_state.translated_text:
            st.markdown("---"); st.subheader(f"Translation ({translation_language}):")
            st.markdown(st.session_state.translated_text)
            render_download_buttons(st.session_state.translated_text, "translated_text")

    with tab3:
        st.subheader("Upload an Audio File to Transcribe and Translate")
        text_models = get_available_models(st.session_state.get('gemini_key'))
        with st.form("transcription_form"):
            uploaded_audio = st.file_uploader("Upload an audio file:", type=["wav", "mp3", "m4a"])
            translation_language_audio = st.selectbox("Translate transcript to:", LANGUAGES, key="translate_lang_audio")
            selected_transcribe_model = st.selectbox("Choose AI Model for Translation:", text_models) if text_models else None

            if st.form_submit_button("🔬 Transcribe & Translate Audio", use_container_width=True):
                if not uploaded_audio: st.error("Please upload an audio file.")
                elif not selected_transcribe_model: st.error("Please select a model.")
                else:
                    with st.spinner("Uploading and transcribing audio..."):
                        transcribed_text = AudioSuiteCrew.transcribe_audio(uploaded_audio)
                        st.session_state.transcribed_text = transcribed_text
                    
                    if st.session_state.transcribed_text:
                        with st.spinner(f"Translating transcript to {translation_language_audio}..."):
                            crew = AudioSuiteCrew(selected_transcribe_model)
                            st.session_state.translated_transcript = crew.translate_text(st.session_state.transcribed_text, translation_language_audio)

        if 'transcribed_text' in st.session_state and st.session_state.transcribed_text:
            st.markdown("---"); st.subheader("Original Transcript:")
            st.markdown(st.session_state.transcribed_text)
        
        if 'translated_transcript' in st.session_state and st.session_state.translated_transcript:
            st.markdown("---"); st.subheader(f"Translated Transcript ({translation_language_audio}):")
            st.markdown(st.session_state.translated_transcript)
            render_download_buttons(st.session_state.translated_transcript, "translated_transcript")

# ==============================================================================
## 4. Main Application Router
# ==============================================================================

def main():
    st.set_page_config(page_title="AI Ministry & Content Suite", layout="wide")

    st.markdown("""<style>.main .block-container { padding-top: 2rem; padding-left: 5rem; padding-right: 5rem; } .stButton>button { border-radius: 20px; border: 1px solid #4CAF50; background-color: #4CAF50; color: white; padding: 10px 24px; font-size: 16px; margin: 4px 2px; cursor: pointer; transition-duration: 0.4s; } .stButton>button:hover { background-color: white; color: black; } h1, h2, h3 { color: #2E4053; }</style>""", unsafe_allow_html=True)

    st.sidebar.title("🔐 Central Configuration")
    st.session_state['gemini_key'] = st.sidebar.text_input("Google Gemini API Key", type="password", value=st.session_state.get('gemini_key', ''))
    st.session_state['serper_key'] = st.sidebar.text_input("Serper.dev API Key", type="password", value=st.session_state.get('serper_key', ''))
    st.sidebar.markdown("---")
    
    st.sidebar.title("Navigation")
    page_options = {
        "Home": "🏠", "Sermon Generator": "📖", "Flyer Production Studio": "🚀", "Worship Song Studio": "🎶",
        "Book Writing Studio": "📚", "Bible Book Study": "🌍", "Bible Topic Study": "🙏", "Newsroom HQ": "📰", 
        "Viral Video Series Studio": "🎬", "Single Video Studio": "📹", "AI Podcast Studio": "🎙️", "AI Chef Studio": "🍳", 
        "AI Language Academy": "🗣️", "AI Tutor (Grades 1-12)": "🎓", "University AI Professor": "🧑‍🏫", "AI Audio Suite": "🎧"
    }
    selection = st.sidebar.radio("Go to", list(page_options.keys()))
    
    keys_needed = {
        "Sermon Generator": ['gemini_key'], "Flyer Production Studio": ['gemini_key', 'serper_key'],
        "Worship Song Studio": ['gemini_key', 'serper_key'], "Book Writing Studio": ['gemini_key', 'serper_key'],
        "Bible Book Study": ['gemini_key', 'serper_key'], "Bible Topic Study": ['gemini_key', 'serper_key'],
        "Newsroom HQ": ['gemini_key', 'serper_key'], "Viral Video Series Studio": ['gemini_key'], "Single Video Studio": ['gemini_key'],
        "AI Podcast Studio": ['gemini_key'], "AI Chef Studio": ['gemini_key'], "AI Language Academy": ['gemini_key'],
        "AI Tutor (Grades 1-12)": ['gemini_key'], "University AI Professor": ['gemini_key'], "AI Audio Suite": ['gemini_key'],
    }

    if selection != "Home" and not all(st.session_state.get(key) for key in keys_needed.get(selection, [])):
        st.warning(f"Please enter the required API Key(s) to use the {selection}."); st.stop()
    
    if selection == "Home":
        st.title("✨ Welcome to the AI Ministry & Content Suite!")
        st.markdown("This suite combines powerful tools to help you create, communicate, and learn effectively. Select a tool from the sidebar to begin your creative, learning, or ministry journey.")
        
        st.markdown("---")
        st.header("🔗 Connect With Me")
        st.markdown("""
        - **Facebook:** Dive into our community discussions and live events. [Join the Conversation](https://www.facebook.com/)
        - **TikTok:** Catch daily inspiration and creative shorts. [Watch Now](https://www.tiktok.com/)
        - **Instagram:** Explore a visual journey of faith and creativity. [Follow Us](https://www.instagram.com/)
        - **YouTube:** Watch in-depth teachings, sermons, and tutorials. [Subscribe Here](https://www.youtube.com/)
        - **Smart Generative App:** Explore another one of my powerful AI applications. [Try it Out](https://smart-app-generative-ai-crew-ai-de.streamlit.app/)
        """)

        st.markdown("---")
        st.header("🔑 How to Get API Keys")
        with st.expander("▶️ How to get your Google Gemini API Key"):
            st.markdown("""
            1.  **Go to Google AI Studio:** Navigate to the [Google AI Studio website](https://aistudio.google.com/).
            2.  **Sign In:** Log in with your Google account.
            3.  **Create API Key:** Click on the "**Get API Key**" button, then choose "**Create API key in new project**".
            4.  **Copy & Paste:** Copy your new key and paste it into the sidebar configuration.
            """)
        with st.expander("▶️ How to get your Serper.dev API Key"):
            st.markdown("""
            1.  **Go to Serper.dev:** Navigate to the [Serper website](https://serper.dev/).
            2.  **Sign Up:** Create a free account (2,500 free queries).
            3.  **Find Your API Key:** Your key is available on your account dashboard.
            4.  **Copy & Paste:** Copy the key and paste it into the sidebar configuration.
            """)

    elif selection == "Sermon Generator": render_sermon_page()
    elif selection == "Flyer Production Studio": render_flyer_page()
    #elif selection == "Worship Song Studio": render_music_page()
    elif selection == "Book Writing Studio": render_book_page()
    elif selection == "Bible Book Study": render_bible_book_study_page()
    elif selection == "Bible Topic Study": render_bible_topic_study_page()
    elif selection == "Newsroom HQ": render_news_page()
    elif selection == "Viral Video Series Studio": render_viral_video_page()
    elif selection == "Single Video Studio": render_single_video_page()
    elif selection == "AI Podcast Studio": render_podcast_studio_page()
    elif selection == "AI Chef Studio": render_chef_page()
    elif selection == "AI Language Academy": render_language_academy_page()
    elif selection == "AI Tutor (Grades 1-12)": render_school_tutor_page()
    elif selection == "University AI Professor": render_university_tutor_page()
    elif selection == "AI Audio Suite": render_audio_suite_page()

if __name__ == "__main__":
    main()
__import__('pysqlite3')

import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')

import streamlit as st
import os
import io
import re
import json
import wave
import asyncio
import time
import base64
from datetime import datetime
from PIL import Image
import google.generativeai as gen
# --- Dependency Imports ---
import docx
import markdown2
import pypdf
from click import option
from google import genai
from google.generativeai import types
from crewai import Agent, Task, Crew, Process, LLM
from crewai_tools import SerperDevTool, ScrapeWebsiteTool
import streamlit.components.v1 as components

import  language as  l
import crew_utis
from general import render_swimming_page, render_fitness_page, render_driving_license_page, render_language_academy_page
from user_guid import render_user_guide_page
from stock_health import  render_health_support_page, render_stock_analyzer_page
# This file holds global configurations and variables for your app.
voice_names = [
    "Zephyr",
    "Puck",
    "Charon",
    "Kore",
    "Fenrir",
    "Leda",
    "Orus",
    "Aoede",
    "Callirrhoe",
    "Autonoe",
    "Enceladus",
    "Iapetus",
    "Umbriel",
    "Algieba",
    "Despina",
    "Erinome",
    "Algenib",
    "Rasalgethi",
    "Laomedeia",
    "Achernar",
    "Alnilam",
    "Schedar",
    "Gacrux",
    "Pulcherrima",
    "Achird",
    "Zubenelgenubi",
    "Vindemiatrix",
    "Sadachbia",
    "Sadaltager",
    "Sulafat"
]
gemini_supported_languages = (
    "Arabic",
    "Bengali",
    "Bulgarian",
    "Chinese (Simplified)",
    "Chinese (Traditional)",
    "Croatian",
    "Czech",
    "Danish",
    "Dutch",
    "English",
    "Estonian",
    "Farsi",
    "Finnish",
    "French",
    "German",
    "Greek",
    "Gujarati",
    "Hebrew",
    "Hindi",
    "Hungarian",
    "Indonesian",
    "Italian",
    "Japanese",
    "Kannada",
    "Korean",
    "Latvian",
    "Lithuanian",
    "Malayalam",
    "Marathi",
    "Norwegian",
    "Polish",
    "Portuguese",
    "Romanian",
    "Russian",
    "Serbian",
    "Slovak",
    "Slovenian",
    "Spanish",
    "Swahili",
    "Swedish",
    "Tamil",
    "Telugu",
    "Thai",
    "Turkish",
    "Ukrainian",
    "Vietnamese"
)
LANGUAGES=gemini_supported_languages

# ==============================================================================
## 1. Helper Functions
# ==============================================================================

def create_downloadable_docx(content):
    """Converts text content to a downloadable DOCX file in memory."""
    doc = docx.Document()
    for line in content.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def parse_json_from_text(text):
    """Safely extracts a JSON object from a string."""
    match = re.search(r'```json\n({.*?})\n```', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            st.error("AI returned an invalid JSON structure. Please try again.")
            return None
    st.warning("Could not parse the AI's JSON response. The structure might be incorrect.")
    return None

def parse_chapters_from_outline(outline_text):
    """Uses regex to find chapter titles in the generated outline."""
    if not isinstance(outline_text, str): return ["Invalid Outline Format"]
    chapters = re.findall(r"Chapter \d+[:\.]\s*(.*)", outline_text, re.IGNORECASE)
    if not chapters:
        chapters = re.findall(r"Kapitel \d+[:\.]\s*(.*)", outline_text, re.IGNORECASE)
    return [ch.strip() for ch in chapters] if chapters else ["Could not parse chapters automatically."]

@st.cache_data
def get_available_models(_api_key, task="generateContent"):
    """Fetches and caches the list of available Gemini models for a specific task."""
    if not _api_key: return []
    try:
        gen.configure(api_key=_api_key)
        if task=='generateContent':
            models = [m.name.replace("models","gemini") for m in gen.list_models() if task in m.supported_generation_methods]
        elif task == "text-to-speech":
            models = [m.name.replace("models/","") for m in gen.list_models() if '-tts' in m.name]
        elif task == "video-generation":
            models = [m.name.replace("models/","") for m in gen.list_models()  if 'veo' in m.name]
        elif task == "image-generation":
             models = [m.name.replace("models/","") for m in gen.list_models() if 'image-generation' in m.name or 'imagen' in m.name]
        else:
            models = [m.name.replace("models/", "") for m in gen.list_models() if 'flash' in m.name or '2.5' in m.name ]
        return sorted(models)
    except Exception:
        st.sidebar.error("Error fetching models: Invalid API Key.")
        return []

def render_download_buttons(content, filename_base):
    """Renders a consistent set of download buttons for text-based content."""
    st.markdown("---")
    st.subheader("Export Your Content")
    col1, col2, col3, col4 = st.columns(4)
    html_content = markdown2.markdown(content)
    with col1:
        st.download_button("⬇️ DOCX", create_downloadable_docx(content), f"{filename_base}.docx")
    with col2:
        st.download_button("⬇️ Markdown", content.encode('utf-8'), f"{filename_base}.md")
    with col3:
        st.download_button("⬇️ HTML", html_content.encode('utf-8'), f"{filename_base}.html")
    with col4:
        st.download_button("⬇️ Text", content.encode('utf-8'), f"{filename_base}.txt")

def pcm_to_wav(pcm_data, channels=1, sample_width=2, sample_rate=24000):
    """Converts raw PCM data to a WAV file in memory."""
    buffer = io.BytesIO()
    with wave.open(buffer, 'wb') as wf:
        wf.setnchannels(channels)
        wf.setsampwidth(sample_width)
        wf.setframerate(sample_rate)
        wf.writeframes(pcm_data)
    buffer.seek(0)
    return buffer.getvalue()

# ==============================================================================
## 2. AI Crew Classes
# ==============================================================================

class SermonCrew:
    def __init__(self, model_name, topic, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.topic = topic
        self.language = language

    def run(self):
        agents = [
            Agent(role='Pentecostal Theologian', goal=f'Create a biblically sound outline for a sermon on "{self.topic}".', backstory="Experienced theologian skilled in structuring compelling sermons.", llm=self.llm, verbose=True),
            Agent(role='Bible Scripture Researcher', goal=f'Find relevant Bible verses for a sermon on "{self.topic}".', backstory="Meticulous Bible scholar dedicated to scriptural accuracy.", llm=self.llm, tools=[SerperDevTool()], verbose=True),
            Agent(role='Gifted Pentecostal Preacher', goal=f'Write a complete, engaging sermon on "{self.topic}" in English.', backstory="Seasoned pastor known for powerful storytelling.", llm=self.llm, verbose=True),
            Agent(role='Expert Theological Translator', goal=f'Translate the final sermon accurately into {self.language}.', backstory=f"Professional translator specializing in theological texts, native in {self.language}.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f'Create a comprehensive outline for a sermon on "{self.topic}".', agent=agents[0], expected_output="A detailed sermon outline in markdown format.")
        task2 = Task(description='Find relevant Bible verses for each point in the sermon outline.', agent=agents[1], context=[task1], expected_output="A list of scriptures organized by outline point.")
        task3 = Task(description='Write a complete sermon in English using the outline and scriptures.', agent=agents[2], context=[task1, task2], expected_output="A complete sermon text in English.")
        task4 = Task(description=f'Translate the final sermon into {self.language}.', agent=agents[3], context=[task3], expected_output=f"The full sermon text translated into {self.language}.", output_file="sermon.md")
        
        crew = Crew(agents=agents, tasks=[task1, task2, task3, task4], process=Process.sequential, verbose=True)
        crew.kickoff()
        with open("sermon.md", "r", encoding="utf-8") as f:
            return f.read()

class FlyerCrew:
    def __init__(self, text_model, image_model, topic, text_element, flyer_type, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        os.environ["SERPER_API_KEY"] = st.session_state.get('serper_key', '')
        self.llm = LLM(model=text_model, temperature=0.8, api_key=os.environ["GOOGLE_API_KEY"])
        self.image_model = image_model
        self.topic = topic; self.text_element = text_element; self.flyer_type = flyer_type; self.language = language
    
    def run_design_crew(self):
        agents = [
            Agent(role='Creative Brief Specialist', goal='Develop a clear creative brief for a flyer.', backstory="Marketing expert skilled in distilling ideas into actionable briefs.", llm=self.llm, tools=[SerperDevTool()], verbose=True),
            Agent(role='Visual Concept Developer', goal='Brainstorm strong visual concepts based on a creative brief.', backstory="Seasoned Art Director with a modern aesthetic sense.", llm=self.llm, tools=[SerperDevTool()], verbose=True),
            Agent(role='Google Imagen Prompt Engineer', goal='Craft a detailed, effective image generation prompt for Google\'s Imagen model.', backstory="Technical artist who knows how to translate concepts into effective AI prompts.", llm=self.llm, verbose=True),
            Agent(role='Multilingual Social Media Copywriter', goal=f'Write a short, engaging social media post in {self.language} to accompany the flyer image.', backstory="Viral marketing specialist who crafts words that stop the scroll in multiple languages.", llm=self.llm, tools=[SerperDevTool()], verbose=True)
        ]
        task1 = Task(description=f"Analyze the request (Topic: '{self.topic}', Text: '{self.text_element}', Type: '{self.flyer_type}') to create a Creative Brief.", agent=agents[0], expected_output="A concise creative brief.")
        task2 = Task(description="Based on the brief, develop a full visual concept.", agent=agents[1], context=[task1], expected_output="A detailed visual concept document.")
        task3 = Task(description="Synthesize the brief and concept into a single, masterful image generation prompt with max 100 words.", agent=agents[2], context=[task2], expected_output="A single paragraph: the final image prompt with max 100 words.")
        task4 = Task(description=f"Based on the brief and concept, write a compelling social media post in {self.language}.", agent=agents[3], context=[task2], expected_output=f"A complete social media post with text and hashtags, written in {self.language} with max 100 words.", output_file="flyer_copy.txt")
        
        crew = Crew(agents=agents, tasks=[task1, task2, task3, task4], process=Process.sequential, verbose=True)
        result = crew.kickoff()
        
        with open("flyer_copy.txt", "r", encoding="utf-8") as f:
            social_copy = f.read()
            
        return {"image_prompt": result.tasks_output[2].raw, "social_copy": social_copy}

    def generate_image(self, prompt):
        try:
            from google import genai as gen
            from google.genai import types

            client = gen.Client(api_key=st.session_state.get('gemini_key', ''))
            response = client.models.generate_images(model=self.image_model, prompt=prompt, config=gen.types.GenerateImagesConfig(number_of_images=1))
            image_bytes_list = []
            for i, generated_image in enumerate(response.generated_images):
                generated_image.image.save(f"flyer{i}.png")
                image_bytes_list.append(f"image{i}.png")
            # The new SDK returns a result object with a list of GeneratedImage objects
            return image_bytes_list
        except Exception as e:
            st.error(f"Image generation failed: {e}"); return None



class BookStudioCrew:
    def __init__(self, model_name, topic, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        os.environ["SERPER_API_KEY"] = st.session_state.get('serper_key', '')
        self.llm = LLM(model=model_name, temperature=0.3, api_key=os.environ["GOOGLE_API_KEY"])
        self.topic = topic; self.language = language

    def run_outline_crew(self, user_prompt):
        agent = Agent(role='Chief Outline Architect', goal=f'Create a chapter-by-chapter outline for a book on "{self.topic}".', backstory='A seasoned developmental editor, you excel at structuring complex ideas.', llm=self.llm, tools=[SerperDevTool()], verbose=True)
        task = Task(description=f"Analyze book idea (Topic: '{self.topic}', Prompt: '{user_prompt}') to develop an outline.", agent=agent, expected_output="A well-structured book outline and Table of Contents.", output_file="book_outline.md")
        crew = Crew(agents=[agent], tasks=[task]).kickoff()
        with open("book_outline.md", "r", encoding="utf-8") as f:
            return f.read()

    def create_and_run_book_crew(self, user_prompt):
        """Initializes and runs the book writing crew."""
        try:




            architect = Agent(role='Chief Outline Architect',
                              goal=f'Create a comprehensive, chapter-by-chapter outline for a ~10-page book on "{self.topic}" in {self.language}.',
                              backstory='A seasoned developmental editor and bestselling author, you excel at structuring complex ideas into engaging book formats.',
                              llm=self.llm, tools=[SerperDevTool(api_key=os.environ["SERPER_API_KEY"])], verbose=True)
            researcher = Agent(role='Research Specialist',
                               goal='Gather, verify, and compile detailed information for each point in the book outline.',
                               backstory='A meticulous multilingual researcher with a Ph.D., you can find a needle in a digital haystack.',
                               llm=self.llm, tools=[SerperDevTool(api_key=os.environ["SERPER_API_KEY"])], verbose=True)
            writer = Agent(role='Narrative Crafter',
                           goal=f'Write engaging, well-structured chapters in {self.language}, based on the provided outline and research.',
                           backstory='A master storyteller and ghostwriter, you bring ideas to life with native-level fluency in several languages.',
                           llm=self.llm, tools=[SerperDevTool(api_key=os.environ["SERPER_API_KEY"])], verbose=True)
            editor = Agent(role='Senior Editor',
                           goal=f'Review, edit, and polish the drafted chapters to ensure stylistic consistency, grammatical correctness, and overall narrative coherence in {self.language}.',
                           backstory='With a red pen sharpened by years at top publishing houses, you are the final gatekeeper of quality.',
                           llm=self.llm, verbose=True)
            output_filename_outline = f'book_outline_output_{self.language.lower()}.md'

            outline_task = Task(
                description=f"Analyze the user's book idea (Topic: '{self.topic}', Prompt: '{user_prompt}') and develop a comprehensive chapter-by-chapter outline. The entire output MUST be in {self.language}.",
                agent=architect, expected_output=f"A detailed book outline, written entirely in {self.language}.",
                output_file=output_filename_outline)
            research_task = Task(
                description=f"For each chapter in the outline, conduct thorough research. Compile all findings into a structured document, tailored for a writer working in {self.language}.",
                agent=researcher, context=[outline_task], expected_output="A well-organized research document.")
            writing_task = Task(
                description=f"Using the outline and research, write the full content for the first three chapters of the book. The entire text must be in {self.language}.",
                agent=writer, context=[research_task],
                expected_output=f"The complete text for the first three chapters, written in {self.language}.")

            output_filename = f'book_final_output_{self.language.lower()}.md'
            editing_task = Task(
                description=f"Perform a comprehensive edit of the drafted chapters. Your review and all final edits must be in {self.language}, ensuring the text sounds like it was written by a native speaker.",
                agent=editor, context=[writing_task],
                expected_output=f"The final, polished text for the written chapters, ready for publication in {self.language}.",
                output_file=output_filename)

            book_crew = Crew(agents=[architect, researcher, writer, editor],
                             tasks=[outline_task, research_task, writing_task, editing_task],
                             process=Process.sequential,
                             Verbose=True)
            result = book_crew.kickoff()
            with open(output_filename, 'r', encoding='utf-8') as file:
                final_output = file.read()
            st.session_state['book_content'] = final_output
        except Exception as e:
            st.error(f"An error occurred while running the AI crew: {e}")



class ChefStudioCrew:
    def __init__(self, model_name, country, food_type, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.country = country; self.food_type = food_type; self.language = language

    def run_crew(self):
        cuisine_specialist = Agent(
            role='World Cuisine & Dietary Specialist',
            goal=f"Generate 3 diverse and exciting meal ideas (Appetizer, Main Course, Dessert) that fit the {self.food_type} category and are inspired by the cuisine of {self.country}.",
            backstory="An acclaimed food historian and globetrotter who understands the soul of a country's food and the principles of dietary choices like veganism. Your suggestions are authentic and inspiring.",
            llm=self.llm,
            verbose=True
        )

        master_chef = Agent(
            role='Executive Chef & Recipe Developer',
            goal="Write clear, concise, and easy-to-follow recipes for the meal ideas provided. Each recipe must include an ingredient list (with metric and imperial measurements), step-by-step instructions, and estimated prep/cook times.",
            backstory="A Michelin-trained chef with a passion for teaching home cooks. You can deconstruct any dish into simple, foolproof steps. Your recipes are reliable and always delicious.",
            llm=self.llm,
            verbose=True
        )

        food_stylist = Agent(
            role='Food Blogger & Creative Director',
            goal="Format the recipes into a beautiful markdown file. For each of the 3 complete meals, write a tantalizing description and a detailed, effective image generation prompt for Gemini to visualize the final dishes.",
            backstory="A top-tier food blogger and photographer who knows how to make food look irresistible. You are an expert in crafting prompts for AI image generators to create stunning, photorealistic food photography.",
            llm=self.llm,
            verbose=True
        )

        # Define the Culinary Tasks
        task_brainstorm = Task(
            description=f"Brainstorm 3 complete meal ideas (appetizer, main, dessert) based on {self.country}'s cuisine for a {self.food_type} diet.",
            expected_output=f"A list of 3 distinct meal plans, each containing a name for an appetizer, a main course, and a dessert.",
            agent=cuisine_specialist
        )

        task_develop_recipes = Task(
            description="For each of the 3 meal plans, write a full, detailed recipe for the appetizer, main course, and dessert.",
            expected_output="A complete set of recipes. Each recipe must have a title, ingredient list, and step-by-step instructions.",
            agent=master_chef,
            context=[task_brainstorm]
        )

        task_format_and_present = Task(
            description="Combine all the recipes into a single, beautifully formatted markdown document. For each of the 3 meals, add a mouth-watering description and a specific, copy-paste ready prompt for an AI image generator (like Gemini) to create a picture of the main course.",
            expected_output=f"A final, user-ready markdown document containing descriptions, recipes for 3 full meals, and 3 distinct in {self.language}, detailed image generation prompts in English.",
            agent=food_stylist,
            output_file="chef_output.md",
            context=[task_develop_recipes]
        )

        crew = Crew(
            agents=[cuisine_specialist, master_chef, food_stylist],
            tasks=[task_brainstorm, task_develop_recipes, task_format_and_present],
            process=Process.sequential
        )


        crew.kickoff()
        with open("chef_output.md", "r", encoding="utf-8") as f:
            return f.read()

class PodcastStudioCrew:
    def __init__(self, model_name, country, language, topic, speaker1, speaker2):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.country = country; self.language = language; self.topic = topic
        self.speaker1 = speaker1; self.speaker2 = speaker2

    def run_script_crew(self):
        agents = [
            Agent(role='Cultural Researcher & Theologian', goal=f"Research '{self.topic}' for an audience in {self.country}.", backstory="A passionate pastor and scholar who connects timeless biblical truths to specific cultural contexts.", llm=self.llm, verbose=True),
            Agent(role='Engaging Podcast Scriptwriter', goal=f"Write a 10-minute script in {self.language} between {self.speaker1} and {self.speaker2}.", backstory="A gifted storyteller for a popular Christian podcast with billions of followers.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f"Develop 3-4 key talking points for a podcast on '{self.topic}'.", agent=agents[0], expected_output="A bulleted list of key themes and talking points.")
        task2 = Task(description=f"Write a full podcast script using the key points, with clear labels for '{self.speaker1}:' and '{self.speaker2}:'.", agent=agents[1], context=[task1], expected_output="A complete podcast script.", output_file="podcast_script.txt")
        crew = Crew(agents=agents, tasks=[task1, task2], process=Process.sequential, verbose=True)
        crew.kickoff()
        with open("podcast_script.txt", "r", encoding="utf-8") as f:
            return f.read()

    @staticmethod
    def generate_audio(tts_model_name, transcript, speaker_configs):
        try:
            from google import genai as gen
            from google.genai import types
            client = gen.Client(api_key=st.session_state.get('gemini_key', ''))
            response = client.models.generate_content(
                model=tts_model_name, contents=transcript,
                config=gen.types.GenerateContentConfig(response_modalities=["AUDIO"], speech_config=gen.types.SpeechConfig(multi_speaker_voice_config=gen.types.MultiSpeakerVoiceConfig(speaker_voice_configs=speaker_configs)))
            )
            return response.candidates[0].content.parts[0].inline_data.data
        except Exception as e:
            st.error(f"Audio generation failed: {e}"); return None

class BibleStudyCrew:
    def __init__(self, model_name, language, translation):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        os.environ["SERPER_API_KEY"] = st.session_state.get('serper_key', '')
        self.llm = LLM(model=model_name, temperature=0.3, api_key=os.environ["GOOGLE_API_KEY"])
        self.language = language
        self.translation = translation

    def run_book_study(self, english_book_name):
        agents = [
            Agent(role='Biblical Historian & Archaeologist', goal=f'Provide a comprehensive historical, cultural, and literary background for {english_book_name}, in {self.language}.', backstory="With a PhD from Jerusalem University, you provide the crucial context that makes the biblical text come alive.", llm=self.llm, tools=[SerperDevTool()], verbose=True),
            Agent(role='Exegetical Theologian', goal=f'Analyze the text of {english_book_name} to uncover its main theological themes, key verses, and structure, presenting findings in {self.language} using the {self.translation} translation for any scripture quotes.', backstory="As a systematic theologian, you are an expert at exegesis—drawing out the intended meaning of the text.", llm=self.llm, tools=[SerperDevTool()], verbose=True),
            Agent(role='Pastoral Guide & Counselor', goal=f'Create practical, thought-provoking application questions and prayer points based on the themes of {english_book_name}, written in {self.language}.', backstory="A seasoned pastor skilled in crafting questions that bridge the gap between ancient text and modern life.", llm=self.llm, verbose=True),
            Agent(role='Senior Editor for Christian Publishing', goal=f'Compile all sections into a single, cohesive, and beautifully formatted Bible study guide in {self.language}.', backstory="You work for an international Christian publishing house, ensuring every manuscript is professional and theologically sound.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f"Create the 'Historical Background' section for a study guide on **{english_book_name}**. Your output MUST be in {self.language}.", agent=agents[0], expected_output="A detailed markdown section on historical background.")
        task2 = Task(description=f"Create the 'Theological Themes & Key Verses' section for **{english_book_name}**. Your output MUST be in {self.language}. Use the {self.translation} Bible translation for quotes.", agent=agents[1], expected_output="A detailed markdown section on theological themes.")
        task3 = Task(description=f"Create the 'Practical Application & Reflection' section for **{english_book_name}**. Your output MUST be in {self.language}.", agent=agents[2], expected_output="An encouraging markdown section with discussion questions.")
        task4 = Task(description=f"Compile all sections into a single study guide. The main title should be the {self.language} translation for 'A Study Guide to the Book of {english_book_name}'.", agent=agents[3], context=[task1, task2, task3], output_file='final_study_guide.md', expected_output="A complete, well-formatted markdown document.")
        
        crew = Crew(agents=agents, tasks=[task1, task2, task3, task4], process=Process.sequential, verbose=True)
        crew.kickoff()
        with open('final_study_guide.md', 'r', encoding='utf-8') as f:
            return f.read()

    def run_topic_study(self, topic, testament):
        agents = [
            Agent(role='Bible Scholar and Researcher', goal=f"Conduct a comprehensive search of the Bible to find at least 10 relevant verses for the topic: '{topic}'. The search must be limited to the {testament} Testament(s) and use the {self.translation} translation.", backstory="You are a meticulous and knowledgeable Bible scholar with a deep understanding of biblical languages and contexts.", tools=[SerperDevTool()], llm=self.llm, verbose=True),
            Agent(role='Pentecostal Pastor and Theologian', goal=f"Write a short, inspiring devotional or sermon outline based on the Bible verses provided. The message should be written in {self.language} and reflect a Pentecostal passion for Jesus and love for people.", backstory="You are a seasoned pastor with a gift for teaching.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f"Search for scriptures related to '{topic}' within the {testament} Testament(s) using the {self.translation} translation. Compile a list of the top 10-15 most relevant verses.", expected_output=f"A markdown-formatted list of 10-15 bible verses, each with its reference.", agent=agents[0])
        task2 = Task(description="Using the list of scriptures from the scholar, write an inspiring devotional.", expected_output=f"A complete devotional message in {self.language}, approximately 300-500 words long.", agent=agents[1], context=[task1], output_file="topic_study.md")
        
        crew = Crew(agents=agents, tasks=[task1, task2], process=Process.sequential, verbose=True)
        crew.kickoff()
        with open("topic_study.md", "r", encoding="utf-8") as f:
            return f.read()

class SchoolTutorCrew:
    def __init__(self, model_name, country, grade, subject, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.3, api_key=os.environ["GOOGLE_API_KEY"])
        self.country = country; self.grade = grade; self.subject = subject; self.language = language

    def run(self, question):
        agents = [
            Agent(role='Curriculum Analyst', backstory="Expert in global K-12 education systems.", goal=f"Analyze the educational context for a {self.grade} student in {self.country} studying {self.subject}.", llm=self.llm, verbose=True),
            Agent(role=f'{self.subject.title()} Subject Matter Expert', backstory=f"Renowned teacher in {self.subject} with a passion for clarity.", goal=f"Accurately solve the student's homework question about {self.subject}.", llm=self.llm, verbose=True),
            Agent(role='Pedagogy and Language Expert', backstory="Master educator skilled at adapting complex information for different age groups.", goal=f"Rewrite the expert's solution into an engaging answer in {self.language} for a {self.grade} student.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f"Analyze context: Country {self.country}, Grade {self.grade}, Subject {self.subject}. Plan how to best explain the answer to: '{question}'", agent=agents[0], expected_output="A brief plan on key concepts, depth, and tone.")
        task2 = Task(description=f"Solve this homework question: '{question}'", agent=agents[1], context=[task1], expected_output="A correct, step-by-step solution.")
        task3 = Task(description=f"Take the expert's solution and rewrite it in {self.language} as a friendly, clear markdown explanation.", agent=agents[2], context=[task2], expected_output="A complete, well-formatted markdown document.", output_file="school_tutor.md")
        
        crew = Crew(agents=agents, tasks=[task1, task2, task3], process=Process.sequential, verbose=True)
        crew.kickoff()
        with open("school_tutor.md", "r", encoding="utf-8") as f:
            return f.read()

class UniversityTutorCrew:
    def __init__(self, model_name, course, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.course = course; self.language = language

    def run(self, question):
        agents = [
            Agent(role='University Curriculum Specialist', backstory="Academic advisor with encyclopedic knowledge of university course structures.", goal=f"Analyze the curriculum for a university course titled '{self.course}'.", llm=self.llm, verbose=True),
            Agent(role=f'University Professor of {self.course}', backstory=f"Distinguished professor with a Ph.D. and extensive research experience relevant to {self.course}.", goal="Provide a rigorous, technically correct solution to the student's question.", llm=self.llm, verbose=True),
            Agent(role='Senior Academic Tutor', backstory="Award-winning teaching assistant skilled at making complex ideas click.", goal=f"Refine the professor's solution into a comprehensive explanation in {self.language}.", llm=self.llm, verbose=True)
        ]
        task1 = Task(description=f"Analyze the academic framework for the course '{self.course}' to answer the question: '{question}'.", agent=agents[0], expected_output="An academic plan outlining the theoretical foundations and expected depth.")
        task2 = Task(description=f"Provide an expert, in-depth solution to the question: '{question}'.", agent=agents[1], context=[task1], expected_output="A detailed, technically accurate, step-by-step solution.")
        task3 = Task(description=f"Synthesize the solution into a high-quality tutorial explanation in {self.language}.", agent=agents[2], context=[task2], expected_output=f"A complete tutorial in {self.language} linking the solution to core concepts of '{self.course}'.", output_file="uni_tutor.md")
        
        crew = Crew(agents=agents, tasks=[task1, task2, task3], process=Process.sequential, verbose=True)
        crew.kickoff()
        with open("uni_tutor.md", "r", encoding="utf-8") as f:
            return f.read()

class NewsroomHQCrew:
    def __init__(self, model_name, scope, location, topics, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        os.environ["SERPER_API_KEY"] = st.session_state.get('serper_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.scope = scope; self.location = location; self.topics = topics; self.language = language

    def run(self):
        editor = Agent(role='Managing Editor', goal=f'Oversee the creation of a high-quality newspaper in {self.language}.', backstory="With decades of experience at major news outlets, you are the final word on journalistic integrity.", llm=self.llm, allow_delegation=True, verbose=True)
        wire_service = Agent(role='News Wire Service', goal='Continuously scan the web for the latest, most significant news stories.', backstory="You are the digital equivalent of the Associated Press, the first to know about any breaking event.", llm=self.llm, tools=[SerperDevTool()], verbose=True)
        reporters = [Agent(role=f'{topic.title()} Reporter', goal=f'Develop in-depth, accurate, and engaging news articles on {topic} in {self.language}.', backstory=f"You are a seasoned journalist with a deep specialization in {topic}.", llm=self.llm, tools=[SerperDevTool()], verbose=True) for topic in self.topics]

        query_location = self.location if self.scope in ["Local", "National"] else "world"
        fetch_task = Task(description=f"Fetch the most recent and significant news stories for today, {datetime.now().strftime('%Y-%m-%d')}, for a {self.scope} newspaper focused on {query_location}.", agent=wire_service, expected_output="A structured list of current news stories, each with a headline, a URL source, and a one-sentence summary.")
        reporting_tasks = [Task(description=f"Using the news wire data, write a concise and compelling news article in {self.language} on your beat: '{topic}'.", agent=reporter, context=[fetch_task], expected_output="A well-formatted news article including a headline, byline, body, and source URL.") for reporter, topic in zip(reporters, self.topics)]
        editing_task = Task(description=f"Review all drafted articles and assemble them into a single, cohesive newspaper format in {self.language}.", agent=editor, context=reporting_tasks, output_file='final_newspaper.md', expected_output="A single, well-formatted Markdown document containing the complete newspaper.")
        
        crew = Crew(agents=[editor, wire_service] + reporters, tasks=[fetch_task] + reporting_tasks + [editing_task], process=Process.sequential, verbose=True)
        crew.kickoff()
        with open('final_newspaper.md', 'r', encoding='utf-8') as f:
            return f.read()

class ViralVideoCrew:
    def __init__(self, model_name, topic_or_verse, language):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])
        self.topic_or_verse = topic_or_verse
        self.language = language

    def run(self):
        agents = [
            Agent(role='Viral Content Strategist', goal=f'Develop a high-level concept and narrative arc for a 45-second, 5-part video series about "{self.topic_or_verse}".', backstory='You are a master of digital storytelling, able to create compelling narratives that unfold across multiple short clips.', llm=self.llm, verbose=True),
            Agent(role='Viral Video Storyboard Artist', goal='Break down a narrative arc into 5 distinct, visually compelling 8-second video concepts.', backstory='You are a visual thinker, able to translate a story into a sequence of powerful, attention-grabbing shots.', llm=self.llm, verbose=True),
            Agent(role='Google VEO Prompt Engineer', goal=f'Write 5 unique, detailed VEO prompts in {self.language}, one for each part of a video storyboard.', backstory='You are an expert in generative video AI, knowing the precise keywords to achieve cinematic quality for a series of related clips.', llm=self.llm, verbose=True),
            Agent(role='Social Media Hook Writer', goal=f'Write 5 unique, engaging social media hooks in {self.language}, one for each video clip in a series.', backstory=f'You craft irresistible, scroll-stopping text that makes people want to watch the next part.', llm=self.llm, verbose=True),
            Agent(role='Content Editor', goal=f'Compile all generated content into a single, well-formatted Markdown document in {self.language}.', backstory='You are a meticulous editor who ensures the final output is clean, organized, and ready for the user.', llm=self.llm, verbose=True)
        ]
        
        task1 = Task(description=f'Based on the topic "{self.topic_or_verse}", create a cohesive narrative arc for a 45-second video series.', agent=agents[0], expected_output="A high-level summary of the 5-part video series concept.")
        task2 = Task(description='Based on the narrative arc, create a detailed storyboard for the 5 video clips.', agent=agents[1], context=[task1], expected_output="A 5-part storyboard, with a detailed visual description for each 8-second clip.")
        task3 = Task(description=f'For each of the 5 storyboard parts, write a unique and highly detailed VEO prompt in {self.language}.', agent=agents[2], context=[task2], expected_output=f"Five distinct, detailed VEO prompts, written in {self.language}.")
        task4 = Task(description=f'For each of the 5 storyboard parts, write a unique and compelling social media hook in {self.language}.', agent=agents[3], context=[task2], expected_output=f"Five numbered social media hooks, written in {self.language}.")
        task5 = Task(description=f'Compile all sections into a single, clean Markdown document.', agent=agents[4], context=[task1, task2, task3, task4], expected_output=f"A final, well-structured Markdown document containing the complete video series plan.", output_file="viral_video_concept.md")

        crew = Crew(agents=agents, tasks=[task1, task2, task3, task4, task5], process=Process.sequential, verbose=True)
        crew.kickoff()
        with open("viral_video_concept.md", "r", encoding="utf-8") as f:
            return f.read()

class SingleVideoCrew:
    @staticmethod
    def generate_video(model_name, prompt):
        try:
            from google import genai as gen
            from google.genai import types
            client = gen.Client(api_key=st.session_state.get('gemini_key', ''))
            operation = client.models.generate_videos(model=model_name, prompt=prompt)
            
            status_placeholder = st.empty()
            with st.spinner("Generating video... This may take several minutes."):
                while not operation.done:
                    status_placeholder.info(f"Waiting for video generation... Status:------")
                    time.sleep(10)
                    operation = client.operations.get(operation)
            
            status_placeholder.success("Video generation complete!")
            generated_video = operation.response.generated_videos[0]
            video_bytes = client.files.download(file=generated_video.video)
            return video_bytes
        except Exception as e:
            st.error(f"An error occurred during video generation: {e}")
            st.info("Please ensure your Google Cloud project has access to the VEO model.")
            return None

class AudioSuiteCrew:
    def __init__(self, model_name):
        os.environ["GOOGLE_API_KEY"] = st.session_state.get('gemini_key', '')
        self.llm = LLM(model=model_name, temperature=0.7, api_key=os.environ["GOOGLE_API_KEY"])

    def translate_text(self, text, language):
        agent = Agent(role='Expert Translator', goal=f"Translate the given text accurately into {language}.", backstory=f"You are a professional translator with expertise in many languages, ensuring natural and fluent translations.", llm=self.llm, verbose=True)
        task = Task(description=f"Translate the following text to {language}: '{text}'", agent=agent, expected_output=f"The translated text in {language}.", output_file="translation.txt")
        crew = Crew(agents=[agent], tasks=[task], verbose=True)
        crew.kickoff()
        with open("translation.txt", "r", encoding="utf-8") as f:
            return f.read()

    @staticmethod
    def generate_audio(tts_model_name, text, voice_name):
        try:
            from google import genai as gen
            from google.genai import types
            client = gen.Client(api_key=st.session_state.get('gemini_key', ''))
            response = client.models.generate_content(
                model=tts_model_name, contents=[f"Say this with a clear and engaging tone: {text}"],
                config=gen.types.GenerateContentConfig(response_modalities=["AUDIO"], speech_config=gen.types.SpeechConfig(voice_config=gen.types.VoiceConfig(prebuilt_voice_config=gen.types.PrebuiltVoiceConfig(voice_name=voice_name))))
            )
            return response.candidates[0].content.parts[0].inline_data.data
        except Exception as e:
            st.error(f"Audio generation failed: {e}"); return None

    @staticmethod
    def transcribe_audio(txt_model,audio_file):
        try:
            #input,
            from google import genai as gen
            from google.genai import types
            client = genai.Client(api_key=st.session_state.get('gemini_key', ''))
            response = client.models.generate_content( model=txt_model,
            contents=["Transcript this audio", types.Part.from_bytes(data=audio_file.read(), mime_type="audio/mpeg")])
            return response.text
        except Exception as e:
            st.error(f"Audio transcription failed: {e}"); return None

class ImageEditingCrew:
    @staticmethod
    def generate_edited_image(model_name, images, prompt):

        try:
            from google import genai as gen
            from google.genai import types
            client = gen.Client(api_key=st.session_state.get('gemini_key', ''))

            image_parts = []
            for img_file in images:
                image_parts.append(gen.types.Part.from_bytes(data=img_file.read(), mime_type="image/jpeg"))
            contents = [gen.types.Content(role="user", parts=image_parts + [gen.types.Part.from_text(text=prompt)])]

            generate_content_config = gen.types.GenerateContentConfig(response_modalities=["TEXT", "IMAGE"])

            response = client.models.generate_content_stream(model="gemini-2.0-flash-preview-image-generation", contents=contents, config=generate_content_config)

            for part in response.candidates[0].content.parts:
                if part.inline_data:
                    img_byte_arr = io.BytesIO(part.inline_data.data)
                    return img_byte_arr.getvalue()
            return None
        except Exception as e:
            st.error(f"An error occurred during image editing: {e}")
            return None

    @staticmethod
    def generate_images_from_prompt(api_key: str, prompt: str, num_images: int, aspect_ratio: str, person_gen: str,
                                    model_name: str):
        """
        Calls the Google GenAI API to generate images based on user inputs.
        Returns a list of PIL Image objects.
        """
        from google.api_core import exceptions
        try:
            os.environ["GOOGLE_API_KEY"] = api_key
            from google import genai as gen
            from google.genai import types
            from PIL import Image
            client = gen.Client(api_key=os.environ["GOOGLE_API_KEY"])

            # This client call is based on the new standalone Python SDK.
            # Ensure the model name is current.
            response = client.models.generate_images(
                model=model_name,  # Using a current model name as of late 2024/early 2025
                prompt=prompt,
                # The config maps directly to the user's selections
                config=types.GenerateImagesConfig(
                    number_of_images=num_images,
                    aspect_ratio=aspect_ratio,
                    personGeneration=person_gen

                )
            )
            image_bytes_list = []
            for i, generated_image in enumerate(response.generated_images):
                # img_byte_arr = BytesIO()
                generated_image.image.save(f"image{i}.png")
                # initial_image = img_byte_arr
                # st.image(f"image{i}.png", caption=f"Generated Initial Image {i}")
                image_bytes_list.append(f"image{i}.png")
            # The new SDK returns a result object with a list of GeneratedImage objects
            return image_bytes_list

        except exceptions.InvalidArgument as e:
            # Catches errors like a prompt being rejected by the safety filter.
            st.error(
                f"⚠️ Your request could not be processed. The prompt may have been rejected by the safety filter. Please try a different prompt. (Error: {e})")
            return []
        except Exception as e:
            st.error(f"An unexpected error occurred: {e}")
            return []

    # --- NEW: Image Studio Page ---
    @staticmethod
    def render_image_studio_page():
        """Renders the AI Image Studio page."""
        st.title("🎨 AI Image Studio")
        st.markdown("Bring your ideas to life. Generate stunning visuals with a simple text prompt.")
        available_image_models = get_available_models(st.session_state.get('gemini_key'), task="image-generation")
        if available_image_models:
            default_model = "gemini-1.5-pro-latest"
            selected_model = st.selectbox("Choose a Gemini Model:", available_image_models,
                                          index=available_image_models.index(
                                              default_model) if default_model in available_image_models else 0)
        else:
            st.warning("Please enter a valid Gemini API Key in the sidebar to load available models.")
            selected_model = None

        with st.form("image_form"):
            st.subheader("Craft Your Vision")
            prompt = st.text_area("Enter your image prompt", height=100,
                                  placeholder="e.g., A majestic lion wearing a crown, sitting on a throne in a futuristic jungle.")

            st.subheader("Configure Your Image")
            col1, col2, col3 = st.columns(3)

            num_images = col1.slider("Number of Images", 1, 4, 1)
            aspect_ratio = col2.selectbox("Aspect Ratio", ["1:1", "3:4", "4:3", "9:16", "16:9"], index=0)
            person_gen_options = {
                "Allow Adults (Default)": "allow_adult",
                "Allow Adults & Children": "allow_all",
                "Don't Allow People": "dont_allow"
            }
            person_gen_display = col3.selectbox("People in Image", options=list(person_gen_options.keys()))
            person_gen_value = person_gen_options[person_gen_display]

            submitted = st.form_submit_button("Generate Images", use_container_width=True)

        if submitted:
            if not prompt:
                st.error("Please enter a prompt to generate an image.")
            else:
                with st.spinner("🎨 The AI is painting your masterpiece..."):
                    # Call the backend function to get the images
                    generated_images = ImageEditingCrew.generate_images_from_prompt(st.session_state.get('gemini_key'), prompt,
                                                                   num_images, aspect_ratio,
                                                                   person_gen_value, selected_model)

                    if generated_images:
                        st.session_state.generated_images = generated_images
                    else:
                        st.session_state.generated_images = []

        if "generated_images" in st.session_state and st.session_state.generated_images:
            st.markdown("---")
            st.subheader("Your Generated Images")

            # Create a dynamic number of columns for the image gallery
            cols = st.columns(len(st.session_state.generated_images))
            for i, img in enumerate(st.session_state.generated_images):
                with cols[i]:
                    # The 'img' object from the SDK is a PIL.Image object
                    st.image(img, caption=f"Image {i + 1}", use_column_width=True)

                    # Provide a download button for each image
                    st.download_button(
                        label="Download",
                        data=img,
                        file_name=f"generated_image_{i + 1}.png",
                        mime="image/png"
                    )

# ==============================================================================
## 3. Page Rendering Functions
# ==============================================================================


def render_book_page():
    st.title("📚 AI Book Writing Studio")
    st.markdown("Outline your book and then write it, one chapter at a time.")
    if 'book_outline' not in st.session_state: st.session_state.book_outline = None
    if 'chapter_content' not in st.session_state: st.session_state.chapter_content = None
    available_models = get_available_models(st.session_state.get('gemini_key'))

    with st.form("outline_form"):
        st.header("Assemble Your Crew and Start Writing")
        language = st.selectbox("Language:", options=LANGUAGES)
        topic = st.text_input("Book Topic:", placeholder="e.g., The History of Ancient Rome")
        user_prompt = st.text_area("Detailed Description:", height=150)
        selected_model = st.selectbox("Choose AI Model", available_models) if available_models else None
        if st.form_submit_button(f"Start Writing My Book in {language}"):
            if not all([topic, user_prompt, selected_model]): st.error("Please fill all fields.")
            else:
                st.session_state.update(book_topic=topic, book_language=language, book_model=selected_model)
                crew = BookStudioCrew(selected_model, topic, language)
                results = crew.create_and_run_book_crew(user_prompt)
    if "book_content" in st.session_state:
        st.markdown("---")
        st.subheader("Your Generated Book Chapters")
        st.markdown(st.session_state["book_content"])
        st.markdown("---")
        render_download_buttons(st.session_state["book_content"], "book_chapters")


    
    if st.session_state.get('chapter_content'):
        st.markdown("---"); st.header("Your Newly Written Chapter"); st.markdown(st.session_state.chapter_content)
        render_download_buttons(st.session_state.chapter_content, "book_chapter")

def render_chef_page():
    st.title("🍳 AI Chef Studio")
    st.markdown("Get multilingual meal plans with recipes and AI image prompts.")
    if 'chef_recipes' not in st.session_state: st.session_state.chef_recipes = None
    if 'chef_prompts' not in st.session_state: st.session_state.chef_prompts = None
    available_models = get_available_models(st.session_state.get('gemini_key'))

    with (st.form("chef_form")):
        st.subheader("What are you in the mood for?")
        col1, col2 = st.columns(2); country = col1.text_input("Country or Region", "Italy"); food_type = col2.selectbox("Food Type", ["Any", "Meat", "Vegetarian", "Vegan"])
        col3, col4 = st.columns(2); language = col3.text_input("Language for Recipe", "English"); selected_model = col4.selectbox("Choose AI Model", available_models) if available_models else None
        if st.form_submit_button("Generate Meal Ideas", use_container_width=True):
            if not all([country, language, selected_model]): st.error("Please fill all fields.")
            else:
                with st.spinner("👩‍🍳 The AI Chef Crew is crafting your menu..."):
                    crew = ChefStudioCrew(selected_model, country, food_type, language)
                    structured_result = crew.run_crew()
                    if structured_result:
                        #st.markdown(structured_result)
                        st.session_state.chef_recipes = structured_result
                        #structured_result.get("recipes_markdown")
    if st.session_state.get('chef_recipes'):
        st.markdown("---"); st.subheader("Your Custom Meal & Recipe Plan"); st.markdown(st.session_state.chef_recipes)
        render_download_buttons(st.session_state.chef_recipes, "recipe_plan")


def render_podcast_studio_page():
    st.title("🎙️ AI Podcast Studio")
    st.markdown("Generate a multi-speaker podcast from scratch with AI.")
    if 'podcast_transcript' not in st.session_state: st.session_state.podcast_transcript = ""
    text_models = get_available_models(st.session_state.get('gemini_key'), task="generateContent")
    tts_models = get_available_models(st.session_state.get('gemini_key'), task="text-to-speech")
    voices =voice_names

    st.header("Step 1: Generate Script")
    with st.form("podcast_script_form"):
        col1, col2 = st.columns(2); topic = col1.text_input("Topic", "The Power of Forgiveness"); country = col2.text_input("Country", "USA")
        language = col1.text_input("Language", "English"); transcript_model = col2.selectbox("Transcript Model", text_models)
        st.markdown("**Speakers**"); col_s1, col_s2 = st.columns(2); speaker1_name = col_s1.text_input("Speaker 1", "Dr. Anya"); speaker2_name = col_s2.text_input("Speaker 2", "Liam")
        if st.form_submit_button("Generate Podcast Script", use_container_width=True):
            with st.spinner("AI Crew is writing your transcript..."):
                crew = PodcastStudioCrew(transcript_model, country, language, topic, speaker1_name, speaker2_name)
                st.session_state.podcast_transcript = crew.run_script_crew()
                st.session_state.update(speaker1_name=speaker1_name, speaker2_name=speaker2_name, podcast_topic=topic)
    
    if st.session_state.podcast_transcript:
        st.text_area("Review/Edit Transcript:", value=st.session_state.podcast_transcript, height=250, key="transcript_editor")
        render_download_buttons(st.session_state.transcript_editor, "podcast_script")
        st.markdown("---"); st.header("Step 2: Generate Audio")
        from google import genai as gen
        from google.genai import types
        with st.form("podcast_audio_form"):
            s1_name = st.session_state.get('speaker1_name', 'Speaker 1'); s2_name = st.session_state.get('speaker2_name', 'Speaker 2')
            col1, col2, col3 = st.columns(3); speaker1_voice = col1.selectbox(f"Voice for {s1_name}", voices); speaker2_voice = col2.selectbox(f"Voice for {s2_name}", voices); tts_model = col3.selectbox("Audio Model", tts_models)
            if st.form_submit_button("Create Podcast Audio", use_container_width=True):
                with st.spinner("🎙️ AI is recording your podcast..."):
                    speaker_configs = [gen.types.SpeakerVoiceConfig(speaker=s1_name, voice_config=gen.types.VoiceConfig(prebuilt_voice_config=gen.types.PrebuiltVoiceConfig(voice_name=speaker1_voice))), gen.types.SpeakerVoiceConfig(speaker=s2_name, voice_config=gen.types.VoiceConfig(prebuilt_voice_config=gen.types.PrebuiltVoiceConfig(voice_name=speaker2_voice)))]
                    st.session_state['audiodata_prod'] = PodcastStudioCrew.generate_audio(tts_model, st.session_state.transcript_editor, speaker_configs)

    if st.session_state.get('audiodata_prod') and  st.session_state.get('podcast_transcript') :
        st.success("Audio generated!")
        wav_bytes = pcm_to_wav(st.session_state.get("audiodata_prod"), channels=1, sample_width=2, sample_rate=24000)
        st.audio(wav_bytes, format='audio/wav')
        st.download_button("⬇️ Download Podcast", wav_bytes, f"{st.session_state.podcast_topic.replace(' ', '_')}.wav")

def render_sermon_page():
    st.title("📖 AI Sermon Generator Crew")
    st.markdown("This application uses a team of AI agents to help you create a detailed, biblically-sound sermon.")
    available_models = get_available_models(st.session_state.get('gemini_key'))
    LANGUAGES = gemini_supported_languages
    
    with st.form("sermon_form"):
        sermon_topic = st.text_input("Enter the Sermon Topic:", "The Power of Forgiveness")
        target_language = st.selectbox("Choose Sermon Language:", LANGUAGES)
        selected_model = st.selectbox("Choose a Gemini Model:", available_models) if available_models else None
        
        if st.form_submit_button("✨ Generate Sermon"):
            if not all([sermon_topic, target_language, selected_model]):
                st.error("Please fill all fields.")
            else:
                with st.spinner(f"The AI Sermon Crew is at work..."):
                    crew = SermonCrew(selected_model, sermon_topic, target_language)
                    st.session_state.sermon_content = crew.run()

    if 'sermon_content' in st.session_state and st.session_state.sermon_content:
        st.markdown("---"); st.subheader("Your Generated Sermon")
        st.markdown(st.session_state.sermon_content)
        render_download_buttons(st.session_state.sermon_content, "sermon")
        if st.button("🎧 Listen to this  sermon "):
           st.session_state['text_for_audio'] =st.session_state.sermon_content
           st.info("Go to the 'Audio Suite' page to generate the audio.")

def render_flyer_page():
    st.title("🚀 AI Flyer Production Studio")
    st.markdown("From idea to share-ready asset in minutes.")
    text_models = get_available_models(st.session_state.get('gemini_key'))
    image_models = get_available_models(st.session_state.get('gemini_key'), task="image-generation")
    LANGUAGES = gemini_supported_languages

    with st.form("flyer_form"):
        st.header("Step 1: Describe Your Flyer")
        topic = st.text_input("Topic or Theme:", "Youth Camp")
        text_element = st.text_input("Key Text (Verse, Slogan):", "Summer Blast 2025")
        flyer_type = st.selectbox("Flyer Type:", ("Social Media Post (Square)", "Poster (Portrait)"))
        language = st.selectbox("Social Media Copy Language:", LANGUAGES)
        text_model = st.selectbox("Choose Text Model:", text_models) if text_models else None
        image_model = st.selectbox("Choose Image Model:", image_models) if image_models else None
        
        if st.form_submit_button("🚀 Generate Concept"):
            if not all([topic, text_element, text_model, image_model]):
                st.error("Please fill all fields.")
            else:
                with st.spinner("AI Design Studio is developing the concept..."):
                    crew = FlyerCrew(text_model, image_model, topic, text_element, flyer_type, language)
                    st.session_state.flyer_concept = crew.run_design_crew()

    if 'flyer_concept' in st.session_state and st.session_state.flyer_concept:
        st.markdown("---"); st.header("Step 2: Generate Image")
        concept = st.session_state.flyer_concept
        st.subheader("🎨 Generated Image Prompt")
        st.code(concept['image_prompt'], language='text')
        
        if st.button("Generate Flyer Image"):
            with st.spinner("Sending prompt to Imagen..."):
                crew = FlyerCrew(text_model, image_model, topic, text_element, flyer_type, language)
                st.session_state.flyer_image = crew.generate_image(concept['image_prompt'])

    if 'flyer_image' in st.session_state and st.session_state.flyer_image:
        st.markdown("---"); st.header("✅ Your Final Assets")
        cols = st.columns(len(st.session_state.flyer_image))
        for i, img in enumerate(st.session_state.flyer_image):
            with cols[i]:
                # The 'img' object from the SDK is a PIL.Image object
                st.image(img, caption=f"Flyer Image", use_column_width=True)
                # Provide a download button for each image
                st.download_button(
                    label="Download",
                    data=img,
                    file_name=f"flyer_image_{i + 1}.png",
                    mime="image/png"
                )
        #st.download_button("⬇️ Download Flyer", st.session_state.flyer_image, "flyer.png", "image/png")
        st.subheader(f"✍️ Your Social Media Caption in {language}")
        st.text_area("", concept['social_copy'], height=150)
        render_download_buttons(concept['social_copy'], "social_media_copy")

# def render_music_page():
def render_music_page():
    st.title("🎶 AI Worship Song Studio")
    st.markdown(
        "Partner with our AI Worship Team to compose a song concept, generate a production-ready prompt for **Google's Lyria AI**, and then generate the audio.")

    available_models = get_available_models(st.session_state.get('gemini_key'))

    st.header("Step 1: Share Your Inspiration")
    col1, col2 = st.columns(2)
    with col1:
        genre = st.selectbox("**Select the Musical Genre:**", (
        "Worship (Hillsong/Bethel style)", "Praise (Elevation/Upbeat style)", "African Gospel Praise"))
        topic = st.text_input("**Core Topic or Theme:**", placeholder="e.g., God's Grace, Faithfulness, Salvation")
    with col2:
        verses = st.text_area("**Enter Bible Verses or Inspirational Text:**", placeholder="e.g., John 3:16, Psalm 23",
                              height=150)
        target_language = st.selectbox("Choose Lyrics Language:", LANGUAGES)

    if available_models:
        selected_model = st.selectbox("Choose a Gemini Model for Songwriting:", available_models,
                                      index=available_models.index(
                                          "gemini-1.5-pro-latest") if "gemini-1.5-pro-latest" in available_models else 0)
    else:
        st.warning("Please enter a valid Gemini API Key in the sidebar to load available models.")
        selected_model = None

    st.header("Step 2: Compose the Song")
    if st.button("Compose & Generate Lyria Prompt"):
        if not all([st.session_state.get(k) for k in ['gemini_key', 'serper_key']]):
            st.error("❌ Please enter your Gemini and Serper API keys in the sidebar.")
        elif not selected_model:
            st.error("❌ Cannot generate. Please provide a valid API key to load models.")
        elif not topic and not verses:
            st.error("🚨 Please provide a Topic or some Bible Verses to inspire the song.")
        else:
            create_and_run_music_crew(api_key=st.session_state['gemini_key'], genre=genre, verses=verses, topic=topic,
                                      language=target_language, model_name=selected_model)

    if "lyria_prompt" in st.session_state:
        st.subheader("✅ Your Final Lyria Prompt")
        st.code(st.session_state["lyria_prompt"], language="text")
        render_download_buttons(st.session_state["lyria_prompt"].raw, "lyria_prompt")

        st.header("Step 3: Generate Your Audio")
        if st.button("🎵 Generate 30-Second Audio Track"):
            generate_and_display_music(st.session_state['gemini_key'],st.session_state["lyria_prompt"])


async def generate_music_async(api_key, prompt, audio_chunks):
    """Async function to handle the real-time music generation."""
    try:
        from google import genai as gen
        from google.genai import types
        os.environ["GOOGLE_API_KEY"]=api_key
        #client = gen.Client(api_key=os.environ["GOOGLE_API_KEY"])
        client = gen.Client(api_key=os.environ["GOOGLE_API_KEY"],http_options={'api_version': 'v1alpha'})

        async def receive_audio(session):
            while True:
                async for message in session.receive():
                    if message.server_content and message.server_content.audio_chunks:
                        audio_chunks.append(message.server_content.audio_chunks[0].data)

        async with client.aio.live.music.connect(model='models/lyria-realtime-exp') as session:
            async with asyncio.TaskGroup() as tg:
                tg.create_task(receive_audio(session))
                await session.set_weighted_prompts(prompts=[gen.types.WeightedPrompt(text=prompt, weight=1.0)])
                await session.set_music_generation_config(
                    config=gen.types.LiveMusicGenerationConfig(bpm=120, temperature=1.0))
                await session.play()
                await asyncio.sleep(30)  # Generate for 30 seconds
                await session.stop()
    except Exception as e:
        st.error(f"Error during async music generation: {e}")


def pcm_to_wav(pcm_data, channels, sample_width, sample_rate):
    """Converts raw PCM data to a WAV file in memory."""
    buffer = io.BytesIO()
    with wave.open(buffer, 'wb') as wf:
        wf.setnchannels(channels)
        wf.setsampwidth(sample_width)
        wf.setframerate(sample_rate)
        wf.writeframes(pcm_data)
    buffer.seek(0)
    return buffer.getvalue()


def generate_and_display_music(api_key, prompt):
    """Main function to orchestrate async generation and display in Streamlit."""
    with st.spinner("Connecting to Lyria and generating your audio track... This will take about 45 seconds."):
        audio_chunks = []
        try:
            asyncio.run(generate_music_async(api_key,prompt, audio_chunks))

            if audio_chunks:
                st.success("Audio track generated successfully!")
                raw_audio_data = b''.join(audio_chunks)
                wav_bytes = pcm_to_wav(raw_audio_data, channels=2, sample_width=2, sample_rate=24000)

                st.audio(wav_bytes, format='audio/wav')
                st.download_button(
                    label="⬇️ Download WAV file",
                    data=wav_bytes,
                    file_name="generated_worship_song.wav",
                    mime="audio/wav"
                )
            else:
                st.error("No audio data was generated. Please try again.")
        except Exception as e:
            st.error(f"An error occurred: {e}")
            st.info(
                "This feature is experimental. Please ensure your Google Cloud account has access to the Lyria API.")


def create_and_run_music_crew(api_key, genre, verses, topic, language, model_name):
    """Initializes and runs the music creation crew."""
    try:
        os.environ["GOOGLE_API_KEY"] = api_key
        llm = LLM(model=model_name, temperature=0.3, api_key=os.environ["GOOGLE_API_KEY"])
        search_tool = SerperDevTool(api_key=st.session_state['serper_key'])
    except Exception as e:
        st.error(f"Error initializing language model: {e}")
        return

    with st.spinner("Your AI Worship Team is gathering... This may take a few minutes."):
        lyricist = Agent(role='Theological Lyricist',
                         goal='Extract core theological truths, emotions, and imagery from user input for a worship song.',
                         backstory="Bridges deep biblical study and heartfelt lyrical expression.", llm=llm,
                         tools=[search_tool], verbose=True)
        songwriter = Agent(role='Worship Songwriter',
                           goal=f'Craft compelling, structured song lyrics in {language} based on the theological concepts.',
                           backstory="Seasoned songwriter skilled in creating poetic and accessible lyrics for worship.",
                           llm=llm, verbose=True)
        arranger = Agent(role='Music Arranger',
                         goal='Define the musical arrangement and atmosphere for a song based on the chosen genre.',
                         backstory="Producer who creates sonic landscapes for any genre, from Bethel to African Gospel.",
                         llm=llm, verbose=True)
        prompt_technician = Agent(role='Lyria Prompt Technician',
                                  goal=f'Synthesize the {language} lyrics and arrangement into a comprehensive prompt for Google\'s Lyria music AI.',
                                  backstory="Specialist in generative AI for music, translating creative vision into precise AI instructions.",
                                  llm=llm, verbose=True)

        lyrical_concept_task = Task(
            description=f"Analyze provided verses ('{verses}') and topic ('{topic}') to create a Lyrical Concept Brief.",
            agent=lyricist, expected_output="A concise Lyrical Concept Brief.")
        arrangement_task = Task(
            description=f"Create a detailed Musical Arrangement Guide for a '{genre}' song about '{topic}'.",
            agent=arranger, expected_output="A detailed Musical Arrangement Guide.")
        song_writing_task = Task(
            description=f"Using the Lyrical Concept Brief, write a complete worship song in {language} with a clear structure (Verse, Chorus, Bridge).",
            agent=songwriter, context=[lyrical_concept_task],
            expected_output="A complete song with clearly labeled sections.")
        prompt_generation_task = Task(
            description=f"Combine the final lyrics (in {language}) and the Musical Arrangement Guide into one single,"
                        f" detailed prompt for Google's Lyria model.",
            agent=prompt_technician, context=[song_writing_task, arrangement_task],
            expected_output="A single, comprehensive text prompt for Lyria.")

        music_crew = Crew(agents=[lyricist, songwriter, arranger, prompt_technician],
                          tasks=[lyrical_concept_task, arrangement_task, song_writing_task, prompt_generation_task],
                          process=Process.sequential, Verbose=True)

        try:
            final_prompt = music_crew.kickoff()
            st.session_state["lyria_prompt"] = final_prompt
        except Exception as e:
            st.error(f"An error occurred during composition: {e}")

def render_bible_book_study_page():
    st.title("📖 Bible Book Study Generator")
    st.markdown("Powered by a team of AI Theologians, Pastors, and Historians to create an in-depth guide to any book of the Bible.")
    # This data is extensive, so it's defined within the page function for clarity.
    ENGLISH_BOOKS = ["Genesis", "Exodus", "Leviticus", "Numbers", "Deuteronomy", "Joshua", "Judges", "Ruth", "1 Samuel",
                     "2 Samuel", "1 Kings", "2 Kings", "1 Chronicles", "2 Chronicles", "Ezra", "Nehemiah", "Esther",
                     "Job", "Psalms", "Proverbs", "Ecclesiastes", "Song of Solomon", "Isaiah", "Jeremiah",
                     "Lamentations", "Ezekiel", "Daniel", "Hosea", "Joel", "Amos", "Obadiah", "Jonah", "Micah", "Nahum",
                     "Habakkuk", "Zephaniah", "Haggai", "Zechariah", "Malachi", "Matthew", "Mark", "Luke", "John",
                     "Acts", "Romans", "1 Corinthians", "2 Corinthians", "Galatians", "Ephesians", "Philippians",
                     "Colossians", "1 Thessalonians", "2 Thessalonians", "1 Timothy", "2 Timothy", "Titus", "Philemon",
                     "Hebrews", "James", "1 Peter", "2 Peter", "1 John", "2 John", "3 John", "Jude", "Revelation"]
    BIBLE_BOOKS_TRANSLATIONS = {
        "English": ENGLISH_BOOKS,
        "German": ["Genesis", "Exodus", "Levitikus", "Numeri", "Deuteronomium", "Josua", "Richter", "Ruth", "1. Samuel",
                   "2. Samuel", "1. Könige", "2. Könige", "1. Chronik", "2. Chronik", "Esra", "Nehemia", "Esther",
                   "Hiob", "Psalmen", "Sprüche", "Prediger", "Hohelied", "Jesaja", "Jeremia", "Klagelieder", "Hesekiel",
                   "Daniel", "Hosea", "Joel", "Amos", "Obadja", "Jona", "Micha", "Nahum", "Habakuk", "Zefanja",
                   "Haggai", "Sacharja", "Maleachi", "Matthäus", "Markus", "Lukas", "Johannes", "Apostelgeschichte",
                   "Römer", "1. Korinther", "2. Korinther", "Galater", "Epheser", "Philipper", "Kolosser",
                   "1. Thessalonicher", "2. Thessalonicher", "1. Timotheus", "2. Timotheus", "Titus", "Philemon",
                   "Hebräer", "Jakobus", "1. Petrus", "2. Petrus", "1. Johannes", "2. Johannes", "3. Johannes", "Judas",
                   "Offenbarung"],
        "French": ["Genèse", "Exode", "Lévitique", "Nombres", "Deutéronome", "Josué", "Juges", "Ruth", "1 Samuel",
                   "2 Samuel", "1 Rois", "2 Rois", "1 Chroniques", "2 Chroniques", "Esdras", "Néhémie", "Esther", "Job",
                   "Psaumes", "Proverbes", "Ecclésiaste", "Cantique des Cantiques", "Ésaïe", "Jérémie", "Lamentations",
                   "Ézéchiel", "Daniel", "Osée", "Joël", "Amos", "Abdias", "Jonas", "Michée", "Nahum", "Habacuc",
                   "Sophonie", "Aggée", "Zacharie", "Malachie", "Matthieu", "Marc", "Luc", "Jean", "Actes", "Romains",
                   "1 Corinthiens", "2 Corinthiens", "Galates", "Éphésiens", "Philippiens", "Colossiens",
                   "1 Thessaloniciens", "2 Thessaloniciens", "1 Timothée", "2 Timothée", "Tite", "Philémon", "Hébreux",
                   "Jacques", "1 Pierre", "2 Pierre", "1 Jean", "2 Jean", "3 Jean", "Jude", "Apocalypse"],
        "Swahili": ["Mwanzo", "Kutoka", "Walawi", "Hesabu", "Kumbukumbu la Torati", "Yoshua", "Waamuzi", "Ruthu",
                    "1 Samweli", "2 Samweli", "1 Wafalme", "2 Wafalme", "1 Mambo ya Nyakati", "2 Mambo ya Nyakati",
                    "Ezra", "Nehemia", "Esta", "Ayubu", "Zaburi", "Methali", "Mhubiri", "Wimbo Ulio Bora", "Isaya",
                    "Yeremia", "Maombolezo", "Ezekieli", "Danieli", "Hosea", "Yoeli", "Amosi", "Obadia", "Yona", "Mika",
                    "Nahumu", "Habakuki", "Sefania", "Hagai", "Zekaria", "Malaki", "Mathayo", "Marko", "Luka", "Yohana",
                    "Matendo", "Warumi", "1 Wakorintho", "2 Wakorintho", "Wagalatia", "Waefeso", "Wafilipi",
                    "Wakolosai", "1 Wathesalonike", "2 Wathesalonike", "1 Timotheo", "2 Timotheo", "Tito", "Filemoni",
                    "Waebrania", "Yakobo", "1 Petro", "2 Petro", "1 Yohana", "2 Yohana", "3 Yohana", "Yuda", "Ufunuo"],
        "Italian": ["Genesi", "Esodo", "Levitico", "Numeri", "Deuteronomio", "Giosuè", "Giudici", "Rut", "1 Samuele",
                    "2 Samuele", "1 Re", "2 Re", "1 Cronache", "2 Cronache", "Esdra", "Neemia", "Ester", "Giobbe",
                    "Salmi", "Proverbi", "Ecclesiaste", "Cantico dei Cantici", "Isaia", "Geremia", "Lamentazioni",
                    "Ezechiele", "Daniele", "Osea", "Gioele", "Amos", "Abdia", "Giona", "Michea", "Naum", "Abacuc",
                    "Sofonia", "Aggeo", "Zaccaria", "Malachia", "Matteo", "Marco", "Luca", "Giovanni", "Atti", "Romani",
                    "1 Corinzi", "2 Corinzi", "Galati", "Efesini", "Filippesi", "Colossesi", "1 Tessalonicesi",
                    "2 Tessalonicesi", "1 Timoteo", "2 Timoteo", "Tito", "Filemone", "Ebrei", "Giacomo", "1 Pietro",
                    "2 Pietro", "1 Giovanni", "2 Giovanni", "3 Giovanni", "Giuda", "Apocalisse"],
        "Spanish": ["Génesis", "Éxodo", "Levítico", "Números", "Deuteronomio", "Josué", "Jueces", "Rut", "1 Samuel",
                    "2 Samuel", "1 Reyes", "2 Reyes", "1 Crónicas", "2 Crónicas", "Esdras", "Nehemías", "Ester", "Job",
                    "Salmos", "Proverbios", "Eclesiastés", "Cantares", "Isaías", "Jeremías", "Lamentaciones",
                    "Ezequiel", "Daniel", "Oseas", "Joel", "Amós", "Abdías", "Jonás", "Miqueas", "Nahúm", "Habacuc",
                    "Sofonías", "Hageo", "Zacarías", "Malaquías", "Mateo", "Marcos", "Lucas", "Juan", "Hechos",
                    "Romanos", "1 Corintios", "2 Corintios", "Gálatas", "Efesios", "Filipenses", "Colosenses",
                    "1 Tesalonicenses", "2 Tesalonicenses", "1 Timoteo", "2 Timoteo", "Tito", "Filemón", "Hebreos",
                    "Santiago", "1 Pedro", "2 Pedro", "1 Juan", "2 Juan", "3 Juan", "Judas", "Apocalipsis"],
        "Portuguese": ["Gênesis", "Êxodo", "Levítico", "Números", "Deuteronômio", "Josué", "Juízes", "Rute", "1 Samuel",
                       "2 Samuel", "1 Reis", "2 Reis", "1 Crônicas", "2 Crônicas", "Esdras", "Neemias", "Ester", "Jó",
                       "Salmos", "Provérbios", "Eclesiastes", "Cânticos", "Isaías", "Jeremias", "Lamentações",
                       "Ezequiel", "Daniel", "Oseias", "Joel", "Amós", "Obadias", "Jonas", "Miqueias", "Naum",
                       "Habacuque", "Sofonias", "Ageu", "Zacarias", "Malaquias", "Mateus", "Marcos", "Lucas", "João",
                       "Atos", "Romanos", "1 Coríntios", "2 Coríntios", "Gálatas", "Efésios", "Filipenses",
                       "Colossenses", "1 Tessalonicenses", "2 Tessalonicenses", "1 Timóteo", "2 Timóteo", "Tito",
                       "Filemom", "Hebreus", "Tiago", "1 Pedro", "2 Pedro", "1 João", "2 João", "3 João", "Judas",
                       "Apocalipse"]
    }
    #ENGLISH_BOOKS = ["Genesis", "Exodus", "Leviticus", "Numbers", "Deuteronomy", "Joshua", "Judges", "Ruth", "1 Samuel", "2 Samuel", "1 Kings", "2 Kings", "1 Chronicles", "2 Chronicles", "Ezra", "Nehemiah", "Esther", "Job", "Psalms", "Proverbs", "Ecclesiastes", "Song of Solomon", "Isaiah", "Jeremiah", "Lamentations", "Ezekiel", "Daniel", "Hosea", "Joel", "Amos", "Obadiah", "Jonah", "Micah", "Nahum", "Habakkuk", "Zephaniah", "Haggai", "Zechariah", "Malachi", "Matthew", "Mark", "Luke", "John", "Acts", "Romans", "1 Corinthians", "2 Corinthians", "Galatians", "Ephesians", "Philippians", "Colossians", "1 Thessalonians", "2 Thessalonians", "1 Timothy", "2 Timothy", "Titus", "Philemon", "Hebrews", "James", "1 Peter", "2 Peter", "1 John", "2 John", "3 John", "Jude", "Revelation"]
    #BIBLE_BOOKS_TRANSLATIONS = {"English": ENGLISH_BOOKS, "German": ["Genesis", "Exodus", "Levitikus", "Numeri", "Deuteronomium", "Josua", "Richter", "Ruth", "1. Samuel", "2. Samuel", "1. Könige", "2. Könige", "1. Chronik", "2. Chronik", "Esra", "Nehemia", "Esther", "Hiob", "Psalmen", "Sprüche", "Prediger", "Hohelied", "Jesaja", "Jeremia", "Klagelieder", "Hesekiel", "Daniel", "Hosea", "Joel", "Amos", "Obadja", "Jona", "Micha", "Nahum", "Habakuk", "Zefanja", "Haggai", "Sacharja", "Maleachi", "Matthäus", "Markus", "Lukas", "Johannes", "Apostelgeschichte", "Römer", "1. Korinther", "2. Korinther", "Galater", "Epheser", "Philipper", "Kolosser", "1. Thessalonicher", "2. Thessalonicher", "1. Timotheo", "2. Timotheo", "Titus", "Philemon", "Hebräer", "Jakobus", "1. Petrus", "2. Petrus", "1. Johannes", "2. Johannes", "3. Johannes", "Judas", "Offenbarung"]}
    available_models = get_available_models(st.session_state.get('gemini_key'))

    with st.form("bible_book_form"):
        st.header("1. Select Your Study Parameters")
        col1, col2 = st.columns(2)
        selected_language = col1.selectbox("Choose your language:", list(BIBLE_BOOKS_TRANSLATIONS.keys()))
        selected_book_translated = col2.selectbox("Choose a book to study:", BIBLE_BOOKS_TRANSLATIONS[selected_language])
        
        col3, col4 = st.columns(2)
        bible_translation = col3.text_input("Enter Bible Translation", placeholder="e.g., NIV, KJV, Hoffnung für alle")
        selected_model = col4.selectbox("Select Gemini Model", available_models) if available_models else None
        
        if st.form_submit_button(f"Create Study Guide for {selected_book_translated}"):
            if not all([selected_model, bible_translation]):
                st.error("❌ Please select a model and enter a Bible translation.")
            else:
                book_index = BIBLE_BOOKS_TRANSLATIONS[selected_language].index(selected_book_translated)
                english_book_name = ENGLISH_BOOKS[book_index]
                with st.spinner(f"Your AI Bible Study Team is preparing your guide..."):
                    crew = BibleStudyCrew(selected_model, selected_language, bible_translation)
                    st.session_state.study_guide_content = crew.run_book_study(english_book_name)

    if "study_guide_content" in st.session_state and st.session_state.study_guide_content:
        st.markdown("---"); st.header("3. Your Custom Study Guide"); st.markdown(st.session_state.study_guide_content)
        render_download_buttons(st.session_state.study_guide_content, f"{selected_book_translated.replace(' ', '_')}_study_guide")
        if st.button("🎧 Listen to this  study guide content"):
           st.session_state['text_for_audio'] =st.session_state.study_guide_content
           st.info("Go to the 'Audio Suite' page to generate the audio.")

def render_bible_topic_study_page():
    st.title("🙏 AI Bible Topic Study")
    st.markdown("Your personal theology research partner. Enter a topic to discover relevant scriptures and receive a custom devotional.")
    available_models = get_available_models(st.session_state.get('gemini_key'))

    with st.form("bible_topic_form"):
        st.header("Start Your Study")
        topic = st.text_input("Enter a Topic, Theme, or Question", placeholder="e.g., Faith, Forgiveness, Who is the Holy Spirit?")
        col1, col2 = st.columns(2)
        language = col1.text_input("Language for Results", "English")
        testament = col2.selectbox("Select Testament(s)", ["All", "Old Testament", "New Testament"])
        
        col3, col4 = st.columns(2)
        bible_translation = col3.text_input("Enter Bible Translation", placeholder="e.g., NIV, KJV, Hoffnung für alle")
        selected_model = col4.selectbox("Choose AI Model", available_models) if available_models else None
        
        if st.form_submit_button("Begin Bible Study", use_container_width=True):
            if not all([topic, language, selected_model, bible_translation]):
                st.error("Please fill all fields.")
            else:
                with st.spinner("The AI ministry team is studying God's Word for you..."):
                    crew = BibleStudyCrew(selected_model, language, bible_translation)
                    st.session_state.study_result = crew.run_topic_study(topic, testament)

    if "study_result" in st.session_state and st.session_state.study_result:
        st.markdown("---"); st.subheader(f"Study Results on '{topic}'"); st.markdown(st.session_state.study_result)
        render_download_buttons(st.session_state.study_result, f"bible_study_{topic.replace(' ', '_')}")
        if st.button("🎧 Listen to this  bible study"):
           st.session_state['text_for_audio'] =st.session_state.study_result
           st.info("Go to the 'Audio Suite' page to generate the audio.")


def render_news_page():
    st.title("📰 AI Newsroom Headquarters")
    st.markdown("Welcome, Editor-in-Chief! Commission a complete, up-to-the-minute digital newspaper from your AI journalist crew.")
    available_models = get_available_models(st.session_state.get('gemini_key'))
    LANGUAGES = gemini_supported_languages

    with st.form("news_form"):
        st.header("Step 1: Define Your Newspaper's Focus")
        scope = st.selectbox("Select Newspaper Scope:", ["Global", "National", "Local"])
        location = ""
        if scope == "Local":
            #st.error(scope)
            location = st.text_input("Enter City:", "NewYork")
        elif scope == "National":
            location = st.text_input("Enter Country:", "USA")

        st.markdown("**Select the sections to include:**")
        topic_options = ["Top Story", "Local News", "Business & Finance", "Sports", "Technology", "Fashion & Trends","Arts & Culture ","Obituaries/Recent deaths","Weather, Comics, and Puzzles"]
        selected_topics = [topic for topic in topic_options if st.checkbox(topic, True, key=f"topic_{topic}")]
        target_language = st.selectbox("Choose Newspaper Language:", LANGUAGES)
        selected_model = st.selectbox("Choose a Gemini Model for Reporting:", available_models) if available_models else None
        
        if st.form_submit_button("Assemble Today's Newspaper"):
            if not selected_model:
                st.error("❌ Please select a model.")
            elif not selected_topics:
                st.error("Please select at least one topic.")
            else:
                with st.spinner("Your AI Newsroom is on the story... This will take a few minutes."):
                    crew = NewsroomHQCrew(selected_model, scope, location, selected_topics, target_language)
                    st.session_state.newspaper_content = crew.run()
                    st.session_state.newspaper_location = location if location else scope

    if "newspaper_content" in st.session_state and st.session_state.newspaper_content:
        st.markdown("---")
        st.subheader(f"The {st.session_state.get('newspaper_location', 'Global')} Times")
        st.markdown(st.session_state.newspaper_content, unsafe_allow_html=True)
        render_download_buttons(st.session_state.newspaper_content, "newspaper")
        if st.button("🎧 Listen to this  digital newspapaer"):
           st.session_state['text_for_audio'] =st.session_state.newspaper_content
           st.info("Go to the 'Audio Suite' page to generate the audio.")

def render_viral_video_page():
    st.title("🎬 AI Viral Video Series Studio")
    st.markdown("Create a powerful, 45-second video series concept for social media, broken into 5 compelling clips.")
    available_models = get_available_models(st.session_state.get('gemini_key'))
    LANGUAGES = gemini_supported_languages

    with st.form("viral_video_form"):
        st.header("Step 1: What's Your Message?")
        topic_or_verse = st.text_input("Enter a Christian Topic or Bible Verse:", placeholder="e.g., John 3:16, Saved by Grace")
        target_language = st.selectbox("Choose Prompt & Hook Language:", LANGUAGES)
        selected_model = st.selectbox("Choose a Gemini Model for Video Concepting:", available_models) if available_models else None
        
        if st.form_submit_button("Generate Viral Video Series Concept"):
            if not all([topic_or_verse, selected_model]):
                st.error("🚨 Please provide a topic/verse and select a model.")
            else:
                with st.spinner("Your AI Social Media team is brainstorming a viral series..."):
                    crew = ViralVideoCrew(selected_model, topic_or_verse, target_language)
                    st.session_state.video_series_content = crew.run()

    if "video_series_content" in st.session_state and st.session_state.video_series_content:
        st.markdown("---")
        st.subheader("✅ Your Viral Video Series Concept")
        st.markdown(st.session_state.video_series_content)
        render_download_buttons(st.session_state.video_series_content, "viral_video_series_concept")

def render_single_video_page():
    st.title("📹 Single Video Studio")
    st.markdown("Unleash your creativity! Use AI to generate a compelling short video from your text prompt.")
    
    video_models = get_available_models(st.session_state.get('gemini_key'), task="video-generation")
    
    st.warning("**Cost Warning:** Generating videos with VEO models can be expensive. A single second of video can cost approximately **€0.75**. Please be mindful of prompt length and usage.", icon="⚠️")

    with st.form("single_video_form"):
        prompt = st.text_area("Enter your video prompt here:", height=150, placeholder="A photorealistic video of a hummingbird flying in slow motion in a tropical garden...")
        selected_model = st.selectbox("Choose VEO Model:", video_models) if video_models else None
        
        submitted = st.form_submit_button("Generate Video", use_container_width=True)

    if submitted:
        if not prompt:
            st.error("Please enter a prompt to generate a video.")
        elif not selected_model:
            st.error("No VEO models available. Check your API Key and access permissions.")
        else:
            video_bytes_content = SingleVideoCrew.generate_video(selected_model, prompt)
            if video_bytes_content:
                st.session_state.single_video = video_bytes_content

    if 'single_video' in st.session_state and st.session_state.single_video:
        st.markdown("---")
        st.subheader("Your Generated Video")
        st.video(st.session_state.single_video)
        st.download_button(
            label="⬇️ Download Video (.mp4)",
            data=st.session_state.single_video,
            file_name="generated_veo_video.mp4",
            mime="video/mp4"
        )

def render_school_tutor_page():
    st.title("🎓 AI Tutor (Grades 1-12)")
    st.markdown("Your personal AI learning assistant. Get step-by-step help with your homework.")
    available_models = get_available_models(st.session_state.get('gemini_key'))
    LANGUAGES = gemini_supported_languages

    with st.form("school_tutor_form"):
        st.subheader("Tell us about your homework")
        col1, col2 = st.columns(2)
        country = col1.text_input("Country", "Germany")
        subject = col1.selectbox("Subject", ["Mathematics", "History", "Science", "Literature", "Physics","Biology","Art", "Spanish","German",
                                             "Homescience","Biology","Chemistry","Geography","Histrory","English","French",
                                             "Italian","Agriculture","Business Administration","CRE"])
        grade = col2.selectbox("Grade / Class", [f"Grade {i}" for i in range(1, 14)])
        language = col2.selectbox("Language for Explanation", LANGUAGES)
        selected_model = st.selectbox("Choose a Gemini Model:", available_models) if available_models else None
        question = st.text_area("📝 Enter your question here", height=150)
        
        if st.form_submit_button("Get Help from AI Tutors!", use_container_width=True):
            if not all([question, selected_model]):
                st.error("Please enter a question and select a model.")
            else:
                with st.spinner("🚀 Your AI Tutors are working on it..."):
                    crew = SchoolTutorCrew(selected_model, country, grade, subject, language)
                    st.session_state.school_result = crew.run(question)

    if "school_result" in st.session_state and st.session_state.school_result:
        st.markdown("---"); st.subheader("✨ Here's your explanation:")
        st.markdown(st.session_state.school_result)
        render_download_buttons(st.session_state.school_result, "school_homework_help")
        if st.button("🎧 Listen to this Homework  Tutorial"):
           st.session_state['text_for_audio'] =st.session_state.school_result
           st.info("Go to the 'Audio Suite' page to generate the audio.")


def render_university_tutor_page():
    st.title("🧑‍🏫 University AI Professor")
    st.markdown("Get expert-level academic help for your university courses.")
    available_models = get_available_models(st.session_state.get('gemini_key'))
    LANGUAGES = gemini_supported_languages

    with st.form("uni_tutor_form"):
        st.subheader("Provide your course and question details")
        course = st.text_input("University Course Name", placeholder="e.g., Experimental Physics, Linear Algebra II")
        language = st.selectbox("Language for Explanation", LANGUAGES)
        selected_model = st.selectbox("Choose a Gemini Model:", available_models) if available_models else None
        question = st.text_area("📝 Enter your question or problem here", height=150)
        
        if st.form_submit_button("Consult the Professor", use_container_width=True):
            if not all([course, question, selected_model]):
                st.error("Please fill all fields and select a model.")
            else:
                with st.spinner("🚀 Consulting with the AI academic team..."):
                    crew = UniversityTutorCrew(selected_model, course, language)
                    st.session_state.uni_result = crew.run(question)

    if "uni_result" in st.session_state and st.session_state.uni_result:
        st.markdown("---"); st.subheader("✨ Professor's Explanation:")
        st.markdown(st.session_state.uni_result)
        render_download_buttons(st.session_state.uni_result, "university_homework_help")
        if st.button(f"🎧 Listen to this  task results"):
           st.session_state['text_for_audio'] =st.session_state.uni_result
           st.info("Go to the 'Audio Suite' page to generate the audio.")

def render_audio_suite_page():
    st.title("🎧 AI Audio Suite")
    st.markdown("Your one-stop shop for audio generation, transcription, and translation.")
    LANGUAGES = gemini_supported_languages
    
    tab1, tab2, tab3 = st.tabs(["**Text-to-Audio**", "**Text Translation**", "**Audio Transcription & Translation**"])

    with tab1:
        st.subheader("Convert Text into High-Quality Audio")
        tts_models = get_available_models(st.session_state.get('gemini_key'), task="text-to-speech")
        voices = voice_names


        with st.form("tts_form"):
            st.session_state['text_for_audio'] = st.text_area("Enter text to convert to audio:", height=150,value=st.session_state.get('text_for_audio',''))
            col1, col2 = st.columns(2)
            selected_voice = col1.selectbox("Choose a Voice:", voices)
            selected_tts_model = col2.selectbox("Choose Audio Model:", tts_models) if tts_models else None
            
            if st.form_submit_button("🎤 Generate Audio", use_container_width=True):
                if not  st.session_state['text_for_audio']: st.error("Please enter some text.")
                elif not selected_tts_model: st.error("Please select an audio model.")
                else:
                    with st.spinner("Generating audio..."):
                        st.session_state['audio_data']= AudioSuiteCrew.generate_audio(selected_tts_model,  st.session_state['text_for_audio'], selected_voice)

        if st.session_state.get('audio_data') and st.session_state.get('text_for_audio') :
            st.success("Audio Generated!")
            wav_bytes = pcm_to_wav(st.session_state.get('audio_data'), channels=1, sample_width=2, sample_rate=24000)
            st.audio(wav_bytes, format='audio/wav')
            st.download_button("⬇️ Download Audio", wav_bytes, f"{selected_tts_model}_{selected_voice}.mp3", "audio/mpeg")

    with tab2:
        st.subheader("Translate Text to a Different Language")
        text_models = get_available_models(st.session_state.get('gemini_key'))
        with st.form("translation_form"):
            text_to_translate = st.text_area("Enter text to translate:", height=150)
            col1, col2 = st.columns(2)
            translation_language = col1.selectbox("Translate to:", LANGUAGES, key="translate_lang_text")
            selected_text_model = col2.selectbox("Choose AI Model:", text_models) if text_models else None
            
            if st.form_submit_button("✍️ Translate Text", use_container_width=True):
                if not text_to_translate: st.error("Please enter text to translate.")
                elif not selected_text_model: st.error("Please select a model.")
                else:
                    with st.spinner(f"Translating to {translation_language}..."):
                        crew = AudioSuiteCrew(selected_text_model)
                        st.session_state.translated_text = crew.translate_text(text_to_translate, translation_language)
        
        if 'translated_text' in st.session_state and st.session_state.translated_text:
            st.markdown("---"); st.subheader(f"Translation ({translation_language}):")
            st.markdown(st.session_state.translated_text)
            render_download_buttons(st.session_state.translated_text, "translated_text")
            if st.button(f"🎧 Listen to this  translated_text"):
                st.session_state['text_for_audio'] = st.session_state.translated_text
               # st.info("Go to the 'Audio Suite' page to generate the audio.")

    with tab3:
        st.subheader("Upload an Audio File to Transcribe and Translate")
        text_models = get_available_models(st.session_state.get('gemini_key'),'text-model')
        trst_models = get_available_models(st.session_state.get('gemini_key'))
        with st.form("transcription_form"):
            uploaded_audio = st.file_uploader("Upload an audio file:", type=["wav", "mp3", "m4a"])
            selected_transcribe_model = st.selectbox("Choose AI Model for Transcription:",
                                                     text_models) if text_models else None
            translation_language_audio = st.selectbox("Translate transcript to:", LANGUAGES, key="translate_lang_audio")
            selected_translate_model = st.selectbox("Choose AI Model for Translation:", trst_models) if text_models else None

            if st.form_submit_button("🔬 Transcribe & Translate Audio", use_container_width=True):
                if not uploaded_audio: st.error("Please upload an audio file.")
                elif not selected_transcribe_model: st.error("Please select a model.")
                else:
                    with st.spinner("Uploading and transcribing audio..."):
                        transcribed_text = AudioSuiteCrew.transcribe_audio(selected_transcribe_model,uploaded_audio)
                        st.session_state.transcribed_text = transcribed_text
                    
                    if st.session_state.transcribed_text:
                        with st.spinner(f"Translating transcript to {translation_language_audio}..."):
                            crew = AudioSuiteCrew(selected_translate_model)
                            st.session_state.translated_transcript = crew.translate_text(st.session_state.transcribed_text, translation_language_audio)

        if 'transcribed_text' in st.session_state and st.session_state.transcribed_text:
            st.markdown("---"); st.subheader("Original Transcript:")
            st.markdown(st.session_state.transcribed_text)
        
        if 'translated_transcript' in st.session_state and st.session_state.translated_transcript:
            st.markdown("---"); st.subheader(f"Translated Transcript ({translation_language_audio}):")
            st.markdown(st.session_state.translated_transcript)
            render_download_buttons(st.session_state.translated_transcript, "translated_transcript")
            if st.button(f"🎧 Listen to this  translatedt"):
                st.session_state['text_for_audio'] = st.session_state.translated_transcript
                st.info("Go to the 'Audio Suite' page to generate the audio.")

def main_render_user_guide_page():
    render_user_guide_page()


def render_image_editing_page():
    ImageEditingCrew.render_image_studio_page()



# ==============================================================================
## 4. Main Application Router
# ==============================================================================

def main():
    st.set_page_config(page_title="AI Ministry & Content Suite", layout="wide")

    st.markdown("""<style>.main .block-container { padding-top: 2rem; padding-left: 5rem; padding-right: 5rem; } .stButton>button { border-radius: 20px; border: 1px solid #4CAF50; background-color: #4CAF50; color: white; padding: 10px 24px; font-size: 16px; margin: 4px 2px; cursor: pointer; transition-duration: 0.4s; } .stButton>button:hover { background-color: white; color: black; } h1, h2, h3 { color: #2E4053; }</style>""", unsafe_allow_html=True)

    st.sidebar.title("🔐 Central Configuration")
    st.session_state['gemini_key'] = st.sidebar.text_input("Google Gemini API Key", type="password", value=st.session_state.get('gemini_key', ''))
    st.session_state['serper_key'] = st.sidebar.text_input("Serper.dev API Key", type="password", value=st.session_state.get('serper_key', ''))
    st.sidebar.markdown("---")
    
    st.sidebar.title("Navigation")
    page_options = {
        "Home": "🏠", "User Guide & Help": "💡", "Sermon Generator": "📖", "Flyer Production Studio": "🚀", "Image  Studio": "🖼️", "Worship Song Studio": "🎶",
        "Book Writing Studio": "📚", "Bible Book Study": "🌍", "Bible Topic Study": "🙏", "Newsroom HQ": "📰", 
        "Viral Video Series Studio": "🎬", "Single Video Studio": "📹", "AI Podcast Studio": "🎙️", "AI Chef Studio": "🍳", 
        "AI Language Academy": "🗣️", "AI Tutor (Grades 1-12)": "🎓", "University AI Professor": "🧑‍🏫", "AI Audio Suite": "🎧",
        "Street Evangelism": "✝️","AI Stock Analysis Studio":"📈", "AI Health & Wellness Suite":"❤️‍🩹",
        "AI Swimming Coach": "🏊",
        "AI Fitness Trainer": "🏋️",
        "AI Driving License Guide": "🚗"
    }
    selection = st.sidebar.radio("Go to", list(page_options.keys()))
    
    keys_needed = {
        "Sermon Generator": ['gemini_key'], "Flyer Production Studio": ['gemini_key', 'serper_key'], "Image  Studio": ['gemini_key'],
        "Worship Song Studio": ['gemini_key', 'serper_key'], "Book Writing Studio": ['gemini_key', 'serper_key'],
        "Bible Book Study": ['gemini_key', 'serper_key'], "Bible Topic Study": ['gemini_key', 'serper_key'],
        "Newsroom HQ": ['gemini_key', 'serper_key'], "Viral Video Series Studio": ['gemini_key'], "Single Video Studio": ['gemini_key'],
        "AI Podcast Studio": ['gemini_key'], "AI Chef Studio": ['gemini_key'], "AI Language Academy": ['gemini_key'],
        "AI Tutor (Grades 1-12)": ['gemini_key'], "University AI Professor": ['gemini_key'], "AI Audio Suite": ['gemini_key'],
        "Street Evangelism": ['gemini_key', 'serper_key'],
        "AI Swimming Coach": ['gemini_key', 'serper_key'],
        "AI Fitness Trainer": ['gemini_key', 'serper_key'],
        "AI Driving License Guide": ['gemini_key', 'serper_key']
    }

    if selection not in ["Home", "User Guide & Help"] and not all(st.session_state.get(key) for key in keys_needed.get(selection, [])):
        st.warning(f"Please enter the required API Key(s) to use the {selection}."); st.stop()
    
    if selection == "Home":
        st.title("✨ Welcome to the AI Ministry & Content Suite!")
        st.markdown("This suite combines powerful tools to help you create, communicate, and learn effectively. Select a tool from the sidebar to begin your creative, learning, or ministry journey.")
        
        st.markdown("---")
        st.header("🔗 Connect With Me")
        st.markdown("""
        - **Facebook:** Dive into our community discussions and live events. [Join the Conversation](https://www.facebook.com/profile.php?id=100027104631463&locale=de_DE)
        - **TikTok:** Catch daily inspiration and creative shorts. [Watch Now](https://www.tiktok.com/)
        - **Instagram:** Explore a visual journey of faith and creativity. [Follow Us](https://www.instagram.com/)
        - **YouTube:** Watch in-depth teachings, sermons, and tutorials. [Subscribe Here](https://studio.youtube.com/channel/UCcHJeYbW6sfIO-a_LoWv3DQ/)
        - **Smart Generative App:** Explore another one of my powerful AI applications. [Try it Out](https://smart-app-generative-ai-crew-ai-de.streamlit.app/)
        """)

        st.markdown("---")
        st.header("🔑 How to Get API Keys")
        with st.expander("▶️ How to get your Google Gemini API Key"):
            st.markdown("""
            1.  **Go to Google AI Studio:** Navigate to the [Google AI Studio website](https://aistudio.google.com/).
            2.  **Sign In:** Log in with your Google account.
            3.  **Create API Key:** Click on the "**Get API Key**" button, then choose "**Create API key in new project**".
            4.  **Copy & Paste:** Copy your new key and paste it into the sidebar configuration.
            """)
        with st.expander("▶️ How to get your Serper.dev API Key"):
            st.markdown("""
            1.  **Go to Serper.dev:** Navigate to the [Serper website](https://serper.dev/).
            2.  **Sign Up:** Create a free account (2,500 free queries).
            3.  **Find Your API Key:** Your key is available on your account dashboard.
            4.  **Copy & Paste:** Copy the key and paste it into the sidebar configuration.
            """)

    elif selection == "User Guide & Help":  main_render_user_guide_page()
    elif selection == "Sermon Generator": render_sermon_page()
    elif selection == "Flyer Production Studio": render_flyer_page()
    elif selection == "Image  Studio": render_image_editing_page()
    elif selection == "Worship Song Studio": render_music_page()
    elif selection == "Book Writing Studio": render_book_page()
    elif selection == "Bible Book Study": render_bible_book_study_page()
    elif selection == "Bible Topic Study": render_bible_topic_study_page()
    elif selection == "Newsroom HQ": render_news_page()
    elif selection == "Viral Video Series Studio": render_viral_video_page()
    elif selection == "Single Video Studio": render_single_video_page()
    elif selection == "AI Podcast Studio": render_podcast_studio_page()
    elif selection == "AI Chef Studio": render_chef_page()
    elif selection == "AI Language Academy": l.render_language_academy_page()
    elif selection == "AI Tutor (Grades 1-12)": render_school_tutor_page()
    elif selection == "University AI Professor": render_university_tutor_page()
    elif selection == "AI Audio Suite": render_audio_suite_page()
    elif selection == "Street Evangelism": crew_utis.render_street_evangelism_page()
    elif selection == "AI Stock Analysis Studio":
        render_stock_analyzer_page()
    elif selection == "AI Health & Wellness Suite": render_health_support_page()
    elif selection == "AI Swimming Coach":
        render_swimming_page()
    elif selection == "AI Fitness Trainer":
        render_fitness_page()
    elif selection == "AI Driving License Guide":
        render_driving_license_page()

if __name__ == "__main__":
    main()

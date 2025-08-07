# --- Virtual Environment Setup Guide ---
# To ensure a clean and isolated environment for this project, it's highly recommended to use a Python virtual environment.
# Follow these steps in your terminal:
#
# 1. Navigate to your project directory:
#    cd path/to/your/project
#
# 2. Create a virtual environment (e.g., named 'venv'):
#    python -m venv venv
#
# 3. Activate the virtual environment:
#    - On Windows:
#      .\venv\Scripts\activate
#    - On macOS & Linux:
#      source venv/bin/activate
#
# 4. Once activated, install the required libraries listed below.

# --- Installation ---
# To run this Streamlit application, you need to install the following libraries.
# You can install them using pip:
# pip install streamlit crewai crewai-tools langchain_google_genai google-generativeai python-docx


import streamlit as st
import os
from crewai import Agent, Task, Crew, Process, LLM 
from crewai_tools import SerperDevTool, ScrapeWebsiteTool, FileReadTool
from datetime import datetime
from docx import Document
from io import BytesIO
import google.generativeai as genai
import time
import asyncio
import wave

# --- App Configuration ---
st.set_page_config(
    page_title="AI Ministry Suite",
    page_icon="✨",
    layout="wide"
)

# --- SHARED COMPONENTS & CONFIGURATION ---

# Using session state to hold credentials and other shared data
if 'gemini_key' not in st.session_state:
    st.session_state['gemini_key'] = ''
if 'serper_key' not in st.session_state:
    st.session_state['serper_key'] = ''

st.sidebar.title("🔐 Central Configuration")
st.sidebar.markdown("Enter your credentials here. They will be used across all pages.")

# Input for Gemini API Key (for all generative tasks)
gemini_key_input = st.sidebar.text_input(
    "Enter Your Google Gemini API Key",
    type="password",
    help="Required for all AI agent text, image, and video generation tasks.",
    value=st.session_state.get('gemini_key', '')
)
st.session_state['gemini_key'] = gemini_key_input

# Input for Serper API Key (for web search tools)
serper_key_input = st.sidebar.text_input(
    "Enter Your Serper.dev API Key",
    type="password",
    help="Required for agents that need web search capabilities.",
    value=st.session_state.get('serper_key', '')
)
st.session_state['serper_key'] = serper_key_input

st.sidebar.markdown("---")
st.sidebar.info("For some services like VEO or Lyria, you may need to authenticate your environment via terminal: `gcloud auth application-default login`")

# --- DYNAMIC MODEL FETCHER ---
@st.cache_data
def get_available_models(api_key, task="generateContent"):
    """Fetches and caches the list of available Gemini models for a specific task."""
    if not api_key:
        return []
    try:
        genai.configure(api_key=api_key)
        models = [
            m.name for m in genai.list_models()
            if task in m.supported_generation_methods
        ]
        return sorted([name.replace("models", "gemini") for name in models])
    except Exception as e:
        st.sidebar.error(f"Error fetching models: Invalid API Key or network issue.")
        return []

# --- PAGE 1: SERMON GENERATOR ---

def render_sermon_page():
    st.title("📖 AI Sermon Generator Crew")
    st.markdown("This application uses a team of AI agents to help you create a detailed, biblically-sound sermon.")
    available_models = get_available_models(st.session_state.get('gemini_key'), task="generateContent")
    sermon_topic = st.text_input("Enter the Sermon Topic:", "The Power of Forgiveness")
    
    if available_models:
        default_model = "gemini-1.5-pro-latest"
        selected_model = st.selectbox("Choose a Gemini Model:", available_models, index=available_models.index(default_model) if default_model in available_models else 0)
    else:
        st.warning("Please enter a valid Gemini API Key in the sidebar to load available models.")
        selected_model = None

    target_language = st.selectbox("Choose Sermon Language:", ("English", "German", "French", "Swahili"))

    if st.button("✨ Generate Sermon"):
        if not st.session_state.get('gemini_key'):
            st.error("❌ Please enter your Gemini API Key in the sidebar to proceed.")
        elif not selected_model:
            st.error("❌ Cannot generate. Please provide a valid API key to load models.")
        elif not sermon_topic:
            st.error("❌ Please provide a Sermon Topic to proceed.")
        else:
            generate_sermon(st.session_state['gemini_key'], sermon_topic, target_language, selected_model)

def generate_sermon(api_key, topic, language, model_name):
    """Initializes and runs the sermon generation crew."""
    try:
        os.environ["GOOGLE_API_KEY"] = api_key
        llm = LLM(model=model_name, temperature=0.7,api_key=os.environ["GOOGLE_API_KEY"] )
    except Exception as e:
        st.error(f"Error initializing language model: {e}")
        return

    with st.spinner(f"The AI Sermon Crew is at work using {model_name}..."):
        theologian = Agent(role='Pentecostal Theologian', goal=f'Create a biblically sound outline for a 10-page sermon on "{topic}".', backstory="Experienced theologian skilled in structuring compelling sermons.", llm=llm, verbose=True)
        researcher = Agent(role='Bible Scripture Researcher', goal=f'Find relevant Bible verses for a sermon on "{topic}".', backstory="Meticulous Bible scholar dedicated to scriptural accuracy.", llm=llm, verbose=True)
        historian = Agent(role='Biblical Historian', goal=f'Provide historical and cultural context for scriptures related to "{topic}".', backstory="Expert in ancient history, bringing scriptures to life.", llm=llm, verbose=True)
        curator = Agent(role='Christian Quote Curator', goal=f'Find powerful quotes from famous Christian figures about "{topic}".', backstory="Well-read historian of Christian thought.", llm=llm, verbose=True)
        writer = Agent(role='Gifted Pentecostal Preacher', goal=f'Write a complete, engaging 10-page sermon on "{topic}" in English.', backstory="Seasoned pastor known for powerful storytelling.", llm=llm, verbose=True)
        translator = Agent(role='Expert Theological Translator', goal=f'Translate the final sermon accurately into {language}.', backstory=f"Professional translator specializing in theological texts, native in {language}.", llm=llm, verbose=True)

        task_outline = Task(description=f'Create a comprehensive outline for a 10-page sermon on "{topic}".', agent=theologian, expected_output="A detailed sermon outline.")
        task_research = Task(description='Find relevant Bible verses for each point in the sermon outline.', agent=researcher, context=[task_outline], expected_output="A list of scriptures organized by outline point.")
        task_context = Task(description='Provide historical context for the key scriptures and themes.', agent=historian, context=[task_outline, task_research], expected_output="A summary of historical context.")
        task_quotes = Task(description=f'Find 3-5 powerful quotes about "{topic}".', agent=curator, context=[task_outline], expected_output="A list of relevant quotes.")
        task_write = Task(description='Write a complete 10-page sermon using the outline, scriptures, context, and quotes.', agent=writer, context=[task_outline, task_research, task_context, task_quotes], expected_output="A complete 10-page sermon in English.")
        task_translate = Task(description=f'Translate the final sermon into {language}.',
                              agent=translator, context=[task_write], expected_output=f"The full sermon text translated into {language}.")

        sermon_crew = Crew(agents=[theologian, researcher, historian, curator, writer, translator],
                           tasks=[task_outline, task_research, task_context, task_quotes, task_write, task_translate],
                           process=Process.sequential, verbose=True)
        try:
            final_sermon = sermon_crew.kickoff()
            st.success("Sermon generation complete!")
            st.markdown(final_sermon)
        except Exception as e:
            st.error(f"An error occurred during sermon generation: {e}")

# --- PAGE 2: FLYER GENERATOR ---

def render_flyer_page():
    st.title("🚀 AI Flyer Production Studio")
    st.markdown("From idea to share-ready asset in minutes. Your expert AI crew will design a concept, generate flyers, and write the social media copy.")
    
    available_text_models = get_available_models(st.session_state.get('gemini_key'), task="generateContent")
    available_image_models = 'imagen-4.0-generate-preview-06-06'
    #(get_available_models(st.session_state.get('gemini_key'), task="generateImages"))

    st.header("Step 1: Describe Your Flyer")
    col1, col2 = st.columns(2)
    with col1:
        topic = st.text_input("**Topic or Theme:**", placeholder="e.g., Street Evangelism, Youth Camp")
    with col2:
        text_element = st.text_input("**Key Text (Verse, Slogan, etc.):**", placeholder='e.g., "John 3:16", "Summer Blast 2025"')
    
    flyer_type = st.selectbox("**Flyer Type:**", ("Social Media Post (Square)", "Poster (Portrait)", "Banner (Landscape)"))
    
    if available_text_models:
        selected_text_model = st.selectbox("Choose a Gemini Model for Concept:", available_text_models, index=available_text_models.index("gemini-1.5-pro-latest") if "gemini-1.5-pro-latest" in available_text_models else 0)
    else:
        st.warning("Please enter a valid Gemini API Key to load text models.")
        selected_text_model = None

    if available_image_models:
        default_image_model = "imagen-3.0-generate-002"
        selected_image_model = st.selectbox("Choose an Imagen Model for Generation:", available_image_models, index=available_image_models.index(default_image_model) if default_image_model in available_image_models else 0)
    else:
        st.warning("Please enter a valid Gemini API Key to load image models.")
        selected_image_model = None

    st.header("Step 2: Produce & Distribute")
    if st.button("🚀 Generate Flyer and Social Copy"):
        if not all([st.session_state.get(k) for k in ['gemini_key', 'serper_key']]):
            st.error("❌ Please enter your Gemini and Serper API Keys in the sidebar.")
        elif not selected_text_model: # or not selected_image_model:
            st.error("❌ Cannot generate. Please provide a valid API key to load all required models.")
        elif not topic or not text_element:
            st.error("❌ Please provide both a Topic and Key Text to proceed.")
        else:
            create_flyer_and_copy(api_key=st.session_state['gemini_key'], topic=topic, text_element=text_element, flyer_type=flyer_type, text_model_name=selected_text_model, image_model_name=selected_image_model)

@st.cache_data
def generate_images_with_genai(api_key: str, prompt: str, model_name: str):
    """Generates images using the google-genai client."""
    try:
        #from google import genai
        genai.configure(api_key=api_key)
        client = genai.client()
        response = client.models.generate_images(
            model=f"{model_name}",
            prompt=prompt,
            config=genai.types.GenerateImagesConfig(number_of_images=4)
        )
        
        image_bytes_list = []
        for generated_image in response.generated_images:
            img_byte_arr = BytesIO()
            generated_image.image.save(img_byte_arr, format='PNG')
            image_bytes_list.append(img_byte_arr.getvalue())
        return image_bytes_list
    except Exception as e:
        st.error(f"An error occurred during image generation: {e}")
        return None

def create_flyer_and_copy(api_key, topic, text_element, flyer_type, text_model_name, image_model_name,serper_key=st.session_state['serper_key']):
    """Initializes and runs the flyer design crew."""
    try:
        os.environ["GOOGLE_API_KEY"] = api_key
        os.environ["SERPER_API_KEY"] = serper_key
        llm = LLM(model=text_model_name, temperature=0.8,api_key=os.environ["GOOGLE_API_KEY"] )
        search_tool = SerperDevTool(api_key=os.environ["SERPER_API_KEY"])
    except Exception as e:
        st.error(f"Error initializing language model: {e}")
        return

    with st.spinner("Your AI Design Studio is developing the concept..."):
        brief_specialist = Agent(role='Creative Brief Specialist', goal='Develop a clear creative brief for a flyer.', backstory="Marketing expert skilled in distilling ideas into actionable briefs.", llm=llm, tools=[search_tool], verbose=True)
        concept_developer = Agent(role='Visual Concept Developer', goal='Brainstorm strong visual concepts based on a creative brief.', backstory="Seasoned Art Director with a modern aesthetic sense.", llm=llm, tools=[search_tool], verbose=True)
        prompt_crafter = Agent(role='Google Imagen Prompt Engineer', goal='Craft a detailed, effective image generation prompt for Google\'s Imagen model.', backstory="Technical artist who knows how to translate concepts into effective AI prompts.", llm=llm, verbose=True)
        copywriter = Agent(role='Social Media Copywriter', goal='Write a short, engaging social media post to accompany the flyer image.', backstory="Viral marketing specialist who crafts words that stop the scroll.", llm=llm, tools=[search_tool], verbose=True)

        briefing = Task(description=f"Analyze the request (Topic: '{topic}', Text: '{text_element}', Type: '{flyer_type}') to create a Creative Brief defining Target Audience, Desired Emotion, and Core Message.", agent=brief_specialist, expected_output="A concise creative brief.")
        visualizing = Task(description="Based on the brief, develop a full visual concept, including a core metaphor, color palette, composition, and artistic style.", agent=concept_developer, context=[briefing], expected_output="A detailed visual concept document.")
        crafting_prompt = Task(description="Synthesize the brief and concept into a single, masterful image generation prompt for Google Imagen. The prompt must be a detailed paragraph. **Do NOT include the text/slogan in the prompt itself.**", agent=prompt_crafter, context=[briefing, visualizing], expected_output="A single paragraph: the final image prompt.")
        crafting_copy = Task(description=f"Based on the brief and concept, write a compelling social media post. "
                                         f"Incorporate the key text '{text_element}', use 3-5 relevant hashtags,"
                                         f" and end with a call-to-action.", agent=copywriter, context=[briefing, visualizing],
                             expected_output="A complete social media post with text and hashtags.")

        flyer_crew = Crew(agents=[brief_specialist, concept_developer, prompt_crafter, copywriter],
                          tasks=[briefing, visualizing, crafting_prompt, crafting_copy], process=Process.sequential, verbose=True)

        crew_result = None
        try:
            crew_result = flyer_crew.kickoff()
            st.write(crew_result.tasks_output[2].raw)
            image_prompt = crew_result.tasks_output[2].raw
            social_copy = crew_result.tasks_output[3].raw
            st.success("Concept approved! Prompt and copy are ready.")
            st.subheader("🎨 Generated Image Prompt")
            st.code(image_prompt, language="text")
        except Exception as e:
            st.error(f"An error occurred during AI crew execution: {e}")
            return

    if crew_result:
        with st.spinner("Sending prompt to Google Imagen for final rendering..."):
            image_bytes_list = generate_images_with_genai(
                api_key=api_key,
                prompt=image_prompt,
                model_name=image_model_name
            )
        
        if image_bytes_list:
            st.success("Rendering complete! Here are your flyer options:")
            st.subheader("✅ Your Final Flyers")
            st.session_state['generated_flyers'] = image_bytes_list # Save flyers for video page
            
            cols = st.columns(len(image_bytes_list))
            for i, image_bytes in enumerate(image_bytes_list):
                with cols[i]:
                    st.image(image_bytes, caption=f"Option {i+1}")
                    st.download_button(
                        label=f"Download Option {i+1}",
                        data=image_bytes,
                        file_name=f"flyer_option_{i+1}.png",
                        mime="image/png"
                    )
            
            st.subheader("✍️ Your Social Media Caption (Ready to Copy)")
            st.text_area("", social_copy, height=150)

# --- PAGE 3: WORSHIP SONG STUDIO ---

def render_music_page():
    st.title("🎶 AI Worship Song Studio")
    st.markdown("Partner with our AI Worship Team to compose a song concept, generate a production-ready prompt for **Google's Lyria AI**, and then generate the audio.")
    
    available_models = get_available_models(st.session_state.get('gemini_key'))

    st.header("Step 1: Share Your Inspiration")
    col1, col2 = st.columns(2)
    with col1:
        genre = st.selectbox("**Select the Musical Genre:**", ("Worship (Hillsong/Bethel style)", "Praise (Elevation/Upbeat style)", "African Gospel Praise"))
        topic = st.text_input("**Core Topic or Theme:**", placeholder="e.g., God's Grace, Faithfulness, Salvation")
    with col2:
        verses = st.text_area("**Enter Bible Verses or Inspirational Text:**", placeholder="e.g., John 3:16, Psalm 23", height=150)
    
    if available_models:
        selected_model = st.selectbox("Choose a Gemini Model for Songwriting:", available_models, index=available_models.index("gemini-1.5-pro-latest") if "gemini-1.5-pro-latest" in available_models else 0)
    else:
        st.warning("Please enter a valid Gemini API Key in the sidebar to load available models.")
        selected_model = None

    st.header("Step 2: Compose the Song")
    if st.button("Compose & Generate Lyria Prompt"):
        if not all([st.session_state.get(k) for k in ['gemini_key', 'serper_key']]):
            st.error("❌ Please enter your Gemini and Serper API keys in the sidebar.")
        elif not selected_model:
            st.error("❌ Cannot generate. Please provide a valid API key to load models.")
        elif not topic and not verses:
            st.error("🚨 Please provide a Topic or some Bible Verses to inspire the song.")
        else:
            create_and_run_music_crew(api_key=st.session_state['gemini_key'], genre=genre, verses=verses, topic=topic, model_name=selected_model)

    if "lyria_prompt" in st.session_state:
        st.subheader("✅ Your Final Lyria Prompt")
        st.info("This is the prompt your AI team created. Now you can generate the audio.", icon="📋")
        st.code(st.session_state["lyria_prompt"], language="text")

        st.header("Step 3: Generate Your Audio")
        if st.button("🎵 Generate 30-Second Audio Track"):
            generate_and_display_music(st.session_state["lyria_prompt"])


async def generate_music_async(prompt, audio_chunks):
    """Async function to handle the real-time music generation."""
    try:
        client = genai.Client(http_options={'api_version': 'v1alpha'})
        
        async def receive_audio(session):
            while True:
                async for message in session.receive():
                    if message.server_content and message.server_content.audio_chunks:
                        audio_chunks.append(message.server_content.audio_chunks[0].data)

        async with client.aio.live.music.connect(model='models/lyria-realtime-exp') as session:
            async with asyncio.TaskGroup() as tg:
                tg.create_task(receive_audio(session))
                await session.set_weighted_prompts(prompts=[genai.types.WeightedPrompt(text=prompt, weight=1.0)])
                await session.set_music_generation_config(config=genai.types.LiveMusicGenerationConfig(bpm=120, temperature=1.0))
                await session.play()
                await asyncio.sleep(30) # Generate for 30 seconds
                await session.stop()
    except Exception as e:
        st.error(f"Error during async music generation: {e}")

def pcm_to_wav(pcm_data, channels, sample_width, sample_rate):
    """Converts raw PCM data to a WAV file in memory."""
    buffer = BytesIO()
    with wave.open(buffer, 'wb') as wf:
        wf.setnchannels(channels)
        wf.setsampwidth(sample_width)
        wf.setframerate(sample_rate)
        wf.writeframes(pcm_data)
    buffer.seek(0)
    return buffer.getvalue()

def generate_and_display_music(prompt):
    """Main function to orchestrate async generation and display in Streamlit."""
    with st.spinner("Connecting to Lyria and generating your audio track... This will take about 45 seconds."):
        audio_chunks = []
        try:
            asyncio.run(generate_music_async(prompt, audio_chunks))
            
            if audio_chunks:
                st.success("Audio track generated successfully!")
                raw_audio_data = b''.join(audio_chunks)
                wav_bytes = pcm_to_wav(raw_audio_data, channels=2, sample_width=2, sample_rate=24000)
                
                st.audio(wav_bytes, format='audio/wav')
                st.download_button(
                    label="⬇️ Download WAV file",
                    data=wav_bytes,
                    file_name="generated_worship_song.wav",
                    mime="audio/wav"
                )
            else:
                st.error("No audio data was generated. Please try again.")
        except Exception as e:
            st.error(f"An error occurred: {e}")
            st.info("This feature is experimental. Please ensure your Google Cloud account has access to the Lyria API.")


def create_and_run_music_crew(api_key, genre, verses, topic, model_name,serper_key=st.session_state['serper_key']):
    """Initializes and runs the music creation crew."""
    try:
        os.environ["GOOGLE_API_KEY"] = api_key
        os.environ["SERPER_KEY"] = serper_key
        llm = LLM(model=model_name, temperature=0.7,api_key=os.environ["GOOGLE_API_KEY"])
        search_tool = SerperDevTool(api_key=os.environ["SERPER_KEY"])
    except Exception as e:
        st.error(f"Error initializing language model: {e}")
        return

    with st.spinner("Your AI Worship Team is gathering... This may take a few minutes."):
        lyricist = Agent(role='Theological Lyricist', goal='Extract core theological truths, emotions, and imagery from user input for a worship song.', backstory="Bridges deep biblical study and heartfelt lyrical expression.", llm=llm, tools=[search_tool], verbose=True)
        songwriter = Agent(role='Worship Songwriter', goal='Craft compelling, structured song lyrics (verse, chorus, bridge) based on the theological concepts.', backstory="Seasoned songwriter skilled in creating poetic and accessible lyrics for worship.", llm=llm, verbose=True)
        arranger = Agent(role='Music Arranger', goal='Define the musical arrangement and atmosphere for a song based on the chosen genre.', backstory="Producer who creates sonic landscapes for any genre, from Bethel to African Gospel.", llm=llm, verbose=True)
        prompt_technician = Agent(role='Lyria Prompt Technician', goal='Synthesize lyrics and arrangement into a comprehensive prompt for Google\'s Lyria music AI.', backstory="Specialist in generative AI for music, translating creative vision into precise AI instructions.", llm=llm, verbose=True)

        lyrical_concept_task = Task(description=f"Analyze provided verses ('{verses}') and topic ('{topic}') to create a Lyrical Concept Brief. Identify Core Message, Key Emotions, and Visual Imagery.", agent=lyricist, expected_output="A concise Lyrical Concept Brief.")
        arrangement_task = Task(description=f"Create a detailed Musical Arrangement Guide for a '{genre}' song about '{topic}'. Specify Tempo, Mood, Dynamics, and genre-specific Instrumentation.", agent=arranger, expected_output="A detailed Musical Arrangement Guide.")
        song_writing_task = Task(description="Using the Lyrical Concept Brief, write a complete worship song with a clear structure (Verse, Chorus, Bridge).", agent=songwriter, context=[lyrical_concept_task], expected_output="A complete song with clearly labeled sections.")
        prompt_generation_task = Task(description="Combine the final lyrics and the Musical Arrangement Guide into one single, detailed prompt for Google's Lyria model. The prompt must be a clear, descriptive paragraph.", agent=prompt_technician, context=[song_writing_task, arrangement_task], expected_output="A single, comprehensive text prompt for Lyria.")

        music_crew = Crew(agents=[lyricist, songwriter, arranger, prompt_technician], tasks=[lyrical_concept_task, arrangement_task, song_writing_task, prompt_generation_task], process=Process.sequential, verbose=True)
        
        try:
            final_prompt = music_crew.kickoff()
            st.session_state["lyria_prompt"] = final_prompt # Save for audio generation
        except Exception as e:
            st.error(f"An error occurred during composition: {e}")

# --- PAGE 4: BOOK WRITING STUDIO ---

def render_book_page():
    st.title("📚 AI Book Writing Studio")
    st.markdown("This is your personal AI writer's room. Provide a topic, a detailed prompt, and select your desired language.")
    
    available_models = get_available_models(st.session_state.get('gemini_key'))

    st.header("Step 1: Define Your Book's Vision")
    language = st.selectbox("**Select the language for your book:**", ("English", "German", "French", "Swahili"))
    topic = st.text_input("**Enter the core topic or theme of your book:**", placeholder="e.g., The history of quantum computing")
    user_prompt = st.text_area("**Provide a detailed description of your book idea:**", height=200, placeholder="Describe your book's focus, target audience, and key themes...")
    
    if available_models:
        selected_model = st.selectbox("Choose a Gemini Model for Book Writing:", available_models, index=available_models.index("gemini-1.5-pro-latest") if "gemini-1.5-pro-latest" in available_models else 0)
    else:
        st.warning("Please enter a valid Gemini API Key in the sidebar to load available models.")
        selected_model = None

    st.header("Step 2: Assemble Your Crew and Start Writing")
    if st.button(f"Start Writing My Book in {language}"):
        if not all([st.session_state.get(k) for k in ['gemini_key', 'serper_key']]):
            st.error("❌ Please enter your Gemini and Serper API keys in the sidebar.")
        elif not selected_model:
            st.error("❌ Cannot generate. Please provide a valid API key to load models.")
        elif not topic or not user_prompt:
            st.error("🚨 Please provide a Topic and a detailed description to proceed.")
        else:
            create_and_run_book_crew(
                api_key=st.session_state['gemini_key'],
                topic=topic,
                user_prompt=user_prompt,
                language=language,
                model_name=selected_model,
                serper_api_key=st.session_state['serper_key']
            )

def create_and_run_book_crew(api_key, topic, user_prompt, language, model_name,serper_api_key):
    """Initializes and runs the book writing crew."""
    #st.info(api_key)
    try:
        os.environ["GOOGLE_API_KEY"] = api_key
        os.environ["OPENAI_KEY"] = api_key
        os.environ["SERPER_API_KEY"] = serper_api_key
        #st.info(serper_key)
        llm = LLM(model=model_name, temperature=0.7,api_key=os.environ["GOOGLE_API_KEY"])
       # st.info(llm)
        search_tool = SerperDevTool(api_key=os.environ["SERPER_API_KEY"])
        #scrape_tool = ScrapeWebsiteTool()
        file_tool = FileReadTool()
    except Exception as e:
        st.error(f"Error initializing language model or tools: {e}")
        return

    with st.spinner(f"Your AI Book Writing Crew is assembling to write in {language}... This may take several minutes."):
        architect = Agent(role='Chief Outline Architect',
                          goal=f'Create a comprehensive, chapter-by-chapter outline for a ~20-page book on "{topic}" in {language}.',
                          backstory='A seasoned developmental editor and bestselling author, you excel '
                                    'at structuring complex ideas into engaging book formats.',
                          llm=llm, tools=[search_tool], verbose=True)
        researcher = Agent(role='Research Specialist', goal='Gather, verify, and compile detailed information '
                                                            'for each point in the book outline.',
                           backstory='A meticulous multilingual researcher with a Ph.D., you can find a needle in a digital haystack.',
                           llm=llm, tools=[search_tool], verbose=True)
        writer = Agent(role='Narrative Crafter', goal=f'Write engaging, '
                                                      f'well-structured chapters in {language},'
                                                      f' based on the provided outline and research.',
                       backstory='A master storyteller and ghostwriter, you bring ideas to life with native-level fluency'
                                 ' in several languages.', llm=llm, tools=[search_tool], verbose=True)
        editor = Agent(role='Senior Editor',
                       goal=f'Review, edit, and polish the drafted chapters to ensure stylistic consistency, '
                            f'grammatical correctness, and overall narrative coherence in {language}.',
                       backstory='With a red pen sharpened by years at top publishing houses, you are the final gatekeeper of quality.',
                       llm=llm, verbose=True)

        outline_task = Task(description=f"Analyze the user's book idea (Topic: '{topic}', Prompt: '{user_prompt}') "
                                        f"and develop a comprehensive chapter-by-chapter outline. The entire output MUST be in {language}.", agent=architect, expected_output=f"A detailed book outline, written entirely in {language}.")
        research_task = Task(description=f"For each chapter in the outline, conduct thorough research. Compile all findings into a structured document, tailored for a writer working in {language}.", agent=researcher, context=[outline_task], expected_output="A well-organized research document.")
        writing_task = Task(description=f"Using the outline and research, write the full content for the first three chapters of the book. The entire text must be in {language}.", agent=writer, context=[research_task], expected_output=f"The complete text for the first three chapters, written in {language}.")
        
        output_filename = f'book_final_output_{language.lower()}.md'
        editing_task = Task(description=f"Perform a comprehensive edit of the drafted chapters. Your review and all final edits must be in {language}, ensuring the text sounds like it was written by a native speaker.", agent=editor, context=[writing_task], expected_output=f"The final, polished text for the written chapters, ready for publication in {language}.", output_file=output_filename)

        book_crew = Crew( llm=llm,
            agents=[architect, researcher, writer, editor], tasks=[outline_task, research_task, writing_task, editing_task],
                          process=Process.sequential, #memory=True,
                          verbose=True)

        try:
            result = book_crew.kickoff()
            st.success("Your AI crew has completed its task!")
            st.balloons()
            st.subheader("Final Book Output")
            try:
                with open(output_filename, 'r', encoding='utf-8') as file:
                    final_output = file.read()
                st.markdown(final_output)
            except FileNotFoundError:
                st.error(f"The final output file ('{output_filename}') was not found. Displaying raw result instead.")
                st.write(result)
        except Exception as e:
            st.error(f"An error occurred while running the AI crew: {e}")

# --- PAGE 5: BIBLE STUDY GENERATOR ---

def render_bible_study_page():
    st.title("🌍 Multilingual AI Bible Study Generator")
    st.markdown("Powered by a team of AI Theologians, Pastors, and Historians.")

    def markdown_to_docx(md_content):
        doc = Document()
        for line in md_content.split('\n'):
            if line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=3)
            elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('# '): doc.add_heading(line.replace('# ', ''), level=1)
            elif line.strip() == '---': doc.add_page_break()
            else: doc.add_paragraph(line)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    ENGLISH_BOOKS = ["Genesis", "Exodus", "Leviticus", "Numbers", "Deuteronomy", "Joshua", "Judges", "Ruth", "1 Samuel", "2 Samuel", "1 Kings", "2 Kings", "1 Chronicles", "2 Chronicles", "Ezra", "Nehemiah", "Esther", "Job", "Psalms", "Proverbs", "Ecclesiastes", "Song of Solomon", "Isaiah", "Jeremiah", "Lamentations", "Ezekiel", "Daniel", "Hosea", "Joel", "Amos", "Obadiah", "Jonah", "Micah", "Nahum", "Habakkuk", "Zephaniah", "Haggai", "Zechariah", "Malachi", "Matthew", "Mark", "Luke", "John", "Acts", "Romans", "1 Corinthians", "2 Corinthians", "Galatians", "Ephesians", "Philippians", "Colossians", "1 Thessalonians", "2 Thessalonians", "1 Timothy", "2 Timothy", "Titus", "Philemon", "Hebrews", "James", "1 Peter", "2 Peter", "1 John", "2 John", "3 John", "Jude", "Revelation"]
    BIBLE_BOOKS_TRANSLATIONS = {
        "English": ENGLISH_BOOKS,
        "German": ["Genesis", "Exodus", "Levitikus", "Numeri", "Deuteronomium", "Josua", "Richter", "Ruth", "1. Samuel", "2. Samuel", "1. Könige", "2. Könige", "1. Chronik", "2. Chronik", "Esra", "Nehemia", "Esther", "Hiob", "Psalmen", "Sprüche", "Prediger", "Hohelied", "Jesaja", "Jeremia", "Klagelieder", "Hesekiel", "Daniel", "Hosea", "Joel", "Amos", "Obadja", "Jona", "Micha", "Nahum", "Habakuk", "Zefanja", "Haggai", "Sacharja", "Maleachi", "Matthäus", "Markus", "Lukas", "Johannes", "Apostelgeschichte", "Römer", "1. Korinther", "2. Korinther", "Galater", "Epheser", "Philipper", "Kolosser", "1. Thessalonicher", "2. Thessalonicher", "1. Timotheus", "2. Timotheus", "Titus", "Philemon", "Hebräer", "Jakobus", "1. Petrus", "2. Petrus", "1. Johannes", "2. Johannes", "3. Johannes", "Judas", "Offenbarung"],
        "French": ["Genèse", "Exode", "Lévitique", "Nombres", "Deutéronome", "Josué", "Juges", "Ruth", "1 Samuel", "2 Samuel", "1 Rois", "2 Rois", "1 Chroniques", "2 Chroniques", "Esdras", "Néhémie", "Esther", "Job", "Psaumes", "Proverbes", "Ecclésiaste", "Cantique des Cantiques", "Ésaïe", "Jérémie", "Lamentations", "Ézéchiel", "Daniel", "Osée", "Joël", "Amos", "Abdias", "Jonas", "Michée", "Nahum", "Habacuc", "Sophonie", "Aggée", "Zacharie", "Malachie", "Matthieu", "Marc", "Luc", "Jean", "Actes", "Romains", "1 Corinthiens", "2 Corinthiens", "Galates", "Éphésiens", "Philippiens", "Colossiens", "1 Thessaloniciens", "2 Thessaloniciens", "1 Timothée", "2 Timothée", "Tite", "Philémon", "Hébreux", "Jacques", "1 Pierre", "2 Pierre", "1 Jean", "2 Jean", "3 Jean", "Jude", "Apocalypse"],
        "Swahili": ["Mwanzo", "Kutoka", "Walawi", "Hesabu", "Kumbukumbu la Torati", "Yoshua", "Waamuzi", "Ruthu", "1 Samweli", "2 Samweli", "1 Wafalme", "2 Wafalme", "1 Mambo ya Nyakati", "2 Mambo ya Nyakati", "Ezra", "Nehemia", "Esta", "Ayubu", "Zaburi", "Methali", "Mhubiri", "Wimbo Ulio Bora", "Isaya", "Yeremia", "Maombolezo", "Ezekieli", "Danieli", "Hosea", "Yoeli", "Amosi", "Obadia", "Yona", "Mika", "Nahumu", "Habakuki", "Sefania", "Hagai", "Zekaria", "Malaki", "Mathayo", "Marko", "Luka", "Yohana", "Matendo", "Warumi", "1 Wakorintho", "2 Wakorintho", "Wagalatia", "Waefeso", "Wafilipi", "Wakolosai", "1 Wathesalonike", "2 Wathesalonike", "1 Timotheo", "2 Timotheo", "Tito", "Filemoni", "Waebrania", "Yakobo", "1 Petro", "2 Petro", "1 Yohana", "2 Yohana", "3 Yohana", "Yuda", "Ufunuo"]
    }
    
    available_models = get_available_models(st.session_state.get('gemini_key'))

    st.header("1. Select Your Language and Book")
    selected_language = st.selectbox("Choose your language:", list(BIBLE_BOOKS_TRANSLATIONS.keys()), key="bible_lang")
    selected_book_translated = st.selectbox("Choose a book to study:", BIBLE_BOOKS_TRANSLATIONS[selected_language], key="bible_book")
    
    if available_models:
        selected_model = st.selectbox("Select Gemini Model", available_models, index=available_models.index("gemini-1.5-pro-latest") if "gemini-1.5-pro-latest" in available_models else 0, key="bible_model")
    else:
        st.warning("Please enter a valid Gemini API Key in the sidebar to load available models.")
        selected_model = None

    st.header("2. Generate Your Study Guide")
    if st.button(f"Create Study Guide for {selected_book_translated}"):
        if not all([st.session_state.get(k) for k in ['gemini_key', 'serper_key']]):
            st.error("🚨 Please enter both your Gemini and Serper API keys in the sidebar to continue.")
        elif not selected_model:
            st.error("❌ Cannot generate. Please provide a valid API key to load models.")
        else:
            if "study_guide_content" in st.session_state:
                del st.session_state["study_guide_content"]
            
            book_index = BIBLE_BOOKS_TRANSLATIONS[selected_language].index(selected_book_translated)
            english_book_name = ENGLISH_BOOKS[book_index]

            with st.spinner(f"Your AI Bible Study Team is preparing your guide for '{selected_book_translated}' in {selected_language}..."):
                create_and_run_bible_study_crew(
                    english_book_name=english_book_name,
                    language=selected_language,
                    model_name=selected_model,
                    api_key=st.session_state['gemini_key'],
                    serper_api_key=st.session_state['serper_key']
                )

    if "study_guide_content" in st.session_state:
        st.header("3. Your Custom Study Guide")
        guide_content = st.session_state["study_guide_content"]
        st.markdown(guide_content)
        
        st.header("4. Export Your Guide")
        col1, col2, col3, col4 = st.columns(4)
        filename_base = f"{selected_book_translated.replace(' ', '_')}_study_guide"
        
        with col1:
            st.download_button("⬇️ DOCX", markdown_to_docx(guide_content), f"{filename_base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with col2:
            st.download_button("⬇️ Markdown", guide_content.encode('utf-8'), f"{filename_base}.md", "text/markdown")
        with col3:
            st.download_button("⬇️ HTML", guide_content.encode('utf-8'), f"{filename_base}.html", "text/html")
        with col4:
            st.download_button("⬇️ Text", guide_content.encode('utf-8'), f"{filename_base}.txt", "text/plain")

def create_and_run_bible_study_crew(english_book_name, language, model_name, api_key, serper_api_key):
    """Initializes and runs the bible study crew."""
    try:
        os.environ["GOOGLE_API_KEY"] = api_key
        os.environ["SERPER_API_KEY"] = serper_api_key
        llm = LLM(model=model_name, temperature=0.5, api_key=os.environ["GOOGLE_API_KEY"])
        search_tool = SerperDevTool(api_key=os.environ["SERPER_API_KEY"])
    except Exception as e:
        st.error(f"Error initializing services: {e}")
        return

    historian = Agent(role='Biblical Historian & Archaeologist', goal=f'Provide a comprehensive historical, cultural, and literary background for {english_book_name}, in {language}.', backstory="With a PhD from Jerusalem University, you provide the crucial context that makes the biblical text come alive.", llm=llm, tools=[search_tool], verbose=True)
    theologian = Agent(role='Exegetical Theologian', goal=f'Analyze the text of {english_book_name} to uncover its main theological themes, key verses, and structure, presenting findings in {language}.', backstory="As a systematic theologian, you are an expert at exegesis—drawing out the intended meaning of the text.", llm=llm, tools=[search_tool], verbose=True)
    pastor = Agent(role='Pastoral Guide & Counselor', goal=f'Create practical, thought-provoking application questions and prayer points based on the themes of {english_book_name}, written in {language}.', backstory="A seasoned pastor skilled in crafting questions that bridge the gap between ancient text and modern life.", llm=llm, verbose=True)
    editor = Agent(role='Senior Editor for Christian Publishing', goal=f'Compile all sections into a single, cohesive, and beautifully formatted Bible study guide in {language}.', backstory="You work for an international Christian publishing house, ensuring every manuscript is professional and theologically sound.", llm=llm, verbose=True)

    task1 = Task(description=f"Create the 'Historical Background' section for a study guide on **{english_book_name}**. Your output MUST be in {language}.", agent=historian, expected_output=f"A Markdown section on the historical background of {english_book_name}, written entirely in {language}.")
    task2 = Task(description=f"Create the 'Theological Themes & Key Verses' section for **{english_book_name}**. Your output MUST be in {language}. Use a well-known {language} Bible translation for quotes.", agent=theologian, expected_output=f"A detailed Markdown section on theological themes of {english_book_name}, written entirely in {language}.")
    task3 = Task(description=f"Create the 'Practical Application & Reflection' section for **{english_book_name}**. Your output MUST be in {language}.", agent=pastor, expected_output=f"An encouraging Markdown section with discussion questions and prayer points for {english_book_name}, written entirely in {language}.")
    
    output_filename = f'final_study_guide_{language.lower()}.md'
    task4 = Task(description=f"Compile all sections into a single study guide. The main title should be the {language} translation for 'A Study Guide to the Book of {english_book_name}'.", agent=editor, context=[task1, task2, task3], expected_output=f"A complete, well-formatted Markdown document in {language}.", output_file=output_filename)

    crew = Crew(agents=[historian, theologian, pastor, editor], tasks=[task1, task2, task3, task4],
                process=Process.sequential, verbose=True)
    
    try:
        result = crew.kickoff()
        with open(output_filename, 'r', encoding='utf-8') as file:
            st.session_state["study_guide_content"] = file.read()
        st.success("Your study guide is ready!")
        st.balloons()
    except Exception as e:
        st.error(f"An error occurred: {e}")

# --- PAGE 6: NEWSROOM HQ ---

def render_news_page():
    st.title("📰 AI Newsroom Headquarters")
    st.markdown("Welcome, Editor-in-Chief! Commission a complete, up-to-the-minute digital newspaper from your AI journalist crew.")
    
    available_models = get_available_models(st.session_state.get('gemini_key'))

    st.header("Step 1: Define Your Newspaper's Focus")
    scope_options = ["Global", "National", "Local"]
    scope = st.selectbox("**Select Newspaper Scope:**", scope_options)

    location = ""
    if scope == "Local":
        location = st.selectbox("Select City:", ["Berlin", "Hamburg", "Munich", "Cologne", "Frankfurt"])
    elif scope == "National":
        location = st.text_input("Enter Country:", "Germany")

    st.markdown("**Select the sections to include in your newspaper:**")
    topic_options = ["Top Story", "Business & Stock Market", "Sports", "Technology", "Fashion & Trends"]
    selected_topics = [topic for topic in topic_options if st.checkbox(topic, True, key=f"topic_{topic}")]
    
    if available_models:
        selected_model = st.selectbox("Choose a Gemini Model for Reporting:", available_models, index=available_models.index("gemini-1.5-pro-latest") if "gemini-1.5-pro-latest" in available_models else 0)
    else:
        st.warning("Please enter a valid Gemini API Key in the sidebar to load available models.")
        selected_model = None

    st.header("Step 2: Go to Print!")
    if st.button("Assemble Today's Newspaper"):
        if not all([st.session_state.get(k) for k in ['gemini_key', 'serper_key']]):
            st.error("🚨 Please enter both your Gemini and Serper API keys in the sidebar to continue.")
        elif not selected_model:
            st.error("❌ Cannot generate. Please provide a valid API key to load models.")
        elif not selected_topics:
            st.error("Please select at least one topic to include in the newspaper.")
        else:
            create_and_run_newspaper_crew(
                api_key=st.session_state['gemini_key'],
                serper_api_key=st.session_state['serper_key'],
                scope=scope,
                location=location,
                topics=selected_topics,
                model_name=selected_model
            )

def create_and_run_newspaper_crew(api_key, serper_api_key, scope, location, topics, model_name):
    """Initializes and runs the newspaper creation crew."""
    try:
        os.environ["GOOGLE_API_KEY"] = api_key
        os.environ["SERPER_API_KEY"] = serper_api_key
        llm = LLM(model=model_name, temperature=0.7,api_key=os.environ["GOOGLE_API_KEY"])
        search_tool = SerperDevTool(api_key=os.environ["SERPER_API_KEY"])
    except Exception as e:
        st.error(f"Error initializing services: {e}")
        return

    with st.spinner("Your AI Newsroom is on the story... This will take a few minutes."):
        editor = Agent(role='Managing Editor', goal='Oversee the creation of a high-quality, relevant, and factually accurate newspaper.', backstory="With decades of experience at major news outlets, you are the final word on journalistic integrity.", llm=llm, allow_delegation=True, verbose=True)
        wire_service = Agent(role='News Wire Service', goal='Continuously scan the web for the latest, most significant news stories.', backstory="You are the digital equivalent of the Associated Press, the first to know about any breaking event.", llm=llm, tools=[search_tool], verbose=True)
        reporters = [Agent(role=f'{topic.title()} Reporter', goal=f'Develop in-depth, accurate, and engaging news articles on {topic}.', backstory=f"You are a seasoned journalist with a deep specialization in {topic}.", llm=llm, tools=[search_tool], verbose=True) for topic in topics]

        query_location = location if scope in ["Local", "National"] else "world"
        fetch_task = Task(description=f"Fetch the most recent and significant news stories for a {scope} newspaper focused on {query_location}. Compile a list of headlines, sources, and summaries.", agent=wire_service, expected_output="A structured list of current news stories.")
        reporting_tasks = [Task(description=f"Using the news wire data, write a concise and compelling news article on your beat: '{topic}'. Include a headline, byline, and a 2-3 paragraph body.", agent=reporter, context=[fetch_task], expected_output="A well-formatted news article.") for reporter, topic in zip(reporters, topics)]
        
        output_filename = 'final_newspaper.md'
        editing_task = Task(description="Review all drafted articles and assemble them into a single, cohesive newspaper format in Markdown.", agent=editor, context=reporting_tasks, expected_output="A single, well-formatted Markdown document containing the complete newspaper.", output_file=output_filename)

        crew = Crew(agents=[editor, wire_service] + reporters, tasks=[fetch_task] + reporting_tasks + [editing_task], process=Process.sequential, verbose=True)
        
        try:
            result = crew.kickoff()
            st.success("Today's edition is ready!")
            st.balloons()
            st.subheader(f"The {location if location else scope} Times")
            try:
                with open(output_filename, 'r', encoding='utf-8') as file:
                    final_output = file.read()
                st.markdown(final_output)
            except FileNotFoundError:
                st.error("The final newspaper file was not found. Displaying raw result instead.")
                st.write(result)
        except Exception as e:
            st.error(f"An error occurred while running the AI crew: {e}")


# --- PAGE 7: VIRAL VIDEO STUDIO ---

def render_viral_video_page():
    st.title("🎬 AI Viral Video Studio")
    st.markdown("Create a powerful, 8-second vertical video for social media. Generate a concept, provide an image, and let VEO bring it to life.")
    
    available_models = get_available_models(st.session_state.get('gemini_key'))

    st.header("Step 1: What's Your Message?")
    topic_or_verse = st.text_input("**Enter a Christian Topic or Bible Verse:**", placeholder="e.g., John 3:16, Saved by Grace")
    
    if available_models:
        selected_model = st.selectbox("Choose a Gemini Model for Video Concepting:", available_models, index=available_models.index("gemini-1.5-pro-latest") if "gemini-1.5-pro-latest" in available_models else 0)
    else:
        st.warning("Please enter a valid Gemini API Key in the sidebar to load available models.")
        selected_model = None

    st.header("Step 2: Create the Vision")
    if st.button("Generate Viral Video Prompt"):
        if not all([st.session_state.get(k) for k in ['gemini_key', 'serper_key']]):
            st.error("🚨 Please enter both your Gemini and Serper API keys in the sidebar.")
        elif not selected_model:
            st.error("❌ Cannot generate. Please provide a valid API key to load models.")
        elif not topic_or_verse:
            st.error("🚨 Please provide a topic or verse to get started.")
        else:
            create_and_run_viral_video_crew(
                api_key=st.session_state['gemini_key'],
                serper_api_key=st.session_state['serper_key'],
                topic_or_verse=topic_or_verse,
                model_name=selected_model
            )

    if "veo_prompt" in st.session_state:
        st.subheader("✅ Your Final VEO Prompt")
        st.code(st.session_state["veo_prompt"], language='text')

        st.header("Step 3: Provide an Image & Generate Video")
        
        # Image input options
        input_image_bytes = None
        
        # Option 1: Use a flyer from the Flyer Production Studio
        if 'generated_flyers' in st.session_state and st.session_state['generated_flyers']:
            st.info("You can use one of the flyers you just created!")
            flyer_options = {f"Flyer Option {i+1}": flyer for i, flyer in enumerate(st.session_state['generated_flyers'])}
            selected_flyer_key = st.selectbox("Choose a generated flyer:", list(flyer_options.keys()))
            if selected_flyer_key:
                input_image_bytes = flyer_options[selected_flyer_key]
                st.image(input_image_bytes, width=200)

        # Option 2: Upload a new image
        uploaded_file = st.file_uploader("Or upload your own image:", type=["png", "jpg", "jpeg"])
        if uploaded_file is not None:
            input_image_bytes = uploaded_file.getvalue()
            st.image(input_image_bytes, caption="Uploaded Image", width=200)

        if st.button("🎬 Generate 8-Second Video from Image"):
            if input_image_bytes:
                generate_and_display_video_from_image(st.session_state["veo_prompt"], input_image_bytes)
            else:
                st.error("Please select a generated flyer or upload an image first.")

def generate_and_display_video_from_image(prompt, image_bytes):
    """Generates a video using VEO with an initial image and displays it."""
    try:
        client = genai.Client()
        
        with st.spinner("Uploading your image..."):
            # Step 1: Upload the image file to the VEO service
            image_file = client.files.upload(
                file=BytesIO(image_bytes),
                display_name="Initial image for VEO video"
            )
            st.success("Image uploaded successfully.")

        with st.spinner("Generating video with VEO... This can take a few minutes."):
            # Step 2: Generate video using the uploaded image
            operation = client.models.generate_videos(
                model="models/veo-1.0-generate-0619",
                prompt=prompt,
                image=image_file,
            )

            status_placeholder = st.empty()
            while not operation.done:
                status_placeholder.info("Polling operation status... please wait.")
                time.sleep(10)
                operation = client.operations.get(operation.name)
            
            status_placeholder.success("Video generation complete!")

            generated_video = operation.response.generated_videos[0]
            video_file_resource = client.files.get(generated_video.video.name)
            video_bytes_content = client.files.download(name=video_file_resource.name).content
            
            st.video(video_bytes_content)
            st.download_button(
                label="⬇️ Download Video",
                data=video_bytes_content,
                file_name="viral_video_from_image.mp4",
                mime="video/mp4"
            )

    except Exception as e:
        st.error(f"An error occurred during video generation: {e}")
        st.error("Please ensure your Gemini API Key is correctly configured for VEO access and that you have authenticated with 'gcloud auth application-default login'.")


def create_and_run_viral_video_crew(api_key, serper_api_key, topic_or_verse, model_name):
    """Initializes and runs the viral video creation crew."""
    try:
        os.environ["GOOGLE_API_KEY"] = api_key
        os.environ["SERPER_API_KEY"] = serper_api_key
        llm = LLM(model=model_name, temperature=0.8,api_key=os.environ["GOOGLE_API_KEY"])
        search_tool = SerperDevTool(api_key=os.environ["SERPER_API_KEY"])
    except Exception as e:
        st.error(f"Error initializing services: {e}")
        return

    with st.spinner("Your AI Social Media team is brainstorming... This might take a moment."):
        trend_analyst = Agent(
            role='Social Media Trend Analyst',
            goal='Identify current visual trends on TikTok and Instagram to make a Christian message go viral.',
            backstory='You have a finger on the pulse of what makes short-form video popular. You understand aesthetics, pacing, and audio that grabs attention.',
            llm=llm, tools=[search_tool], verbose=True
        )
        theologian = Agent(
            role='Theological Content Strategist',
            goal='Distill the core, powerful message from a Christian topic or Bible verse into a simple, impactful concept for a modern audience.',
            backstory='You can take deep theological truths and make them understandable and relatable in seconds.',
            llm=llm, verbose=True
        )
        storyteller = Agent(
            role='Viral Video Visual Storyteller',
            goal='Develop a compelling, 8-second visual narrative concept for a vertical video that is emotionally resonant and visually stunning.',
            backstory='You are a master of short-form cinema, able to tell a powerful story without words, using only visuals and motion.',
            llm=llm, verbose=True
        )
        veo_engineer = Agent(
            role='Google VEO Prompt Engineer',
            goal='Craft a detailed, technically sound prompt for Google VEO that translates a visual concept into an executable video generation command.',
            backstory='You are an expert in generative video AI, knowing the precise keywords to achieve cinematic quality, specific camera angles, and emotional tone.',
            llm=llm, verbose=True
        )

        trend_task = Task(description='Analyze current viral video trends (styles, aesthetics, camera movements, effects) that could be adapted for a profound Christian message.', agent=trend_analyst, expected_output="A summary of 2-3 current, relevant visual trends.")
        theology_task = Task(description=f'Distill the core message of "{topic_or_verse}" into a single, powerful sentence or emotional concept.', agent=theologian, expected_output="A single sentence that captures the essence of the topic.")
        visual_task = Task(description='Based on the core message and viral trends, create a shot-by-shot storyboard for an 8-second video. Describe the scene, subject, camera movement, lighting, and overall emotion. The concept must be visually captivating and suitable for a vertical format.', agent=storyteller, context=[trend_task, theology_task], expected_output="A detailed storyboard description for an 8-second video.")
        prompting_task = Task(description='Combine the storyboard, core message, and trend analysis into a single, detailed paragraph prompt for Google VEO. The prompt must be optimized for creating a hyper-realistic, cinematic 8-second vertical video (9:16 aspect ratio).', agent=veo_engineer, context=[visual_task, theology_task], expected_output="A single, comprehensive paragraph: the final VEO prompt.")

        crew = Crew(agents=[trend_analyst, theologian, storyteller, veo_engineer], tasks=[trend_task, theology_task, visual_task, prompting_task], process=Process.sequential, verbose=True)

        try:
            final_prompt = crew.kickoff()
            st.session_state["veo_prompt"] = final_prompt # Save the prompt to session state
            
        except Exception as e:
            st.error(f"An error occurred while creating the video concept: {e}")


# --- MAIN APP ROUTER ---

def main():
    st.sidebar.title("Navigation")
    page_options = {
        "Home": "🏠",
        "Sermon Generator": "📖",
        "Flyer Production Studio": "�",
        "Worship Song Studio": "🎶",
        "Book Writing Studio": "📚",
        "Bible Study Generator": "🌍",
        "Newsroom HQ": "📰",
        "Viral Video Studio": "🎬"
    }
    selection = st.sidebar.radio("Go to", list(page_options.keys()))

    if selection == "Home":
        st.title("✨ Welcome to the AI Ministry & Content Suite!")
        st.markdown("---")
        st.header("Your AI-Powered Partner in Communication")
        st.markdown("""
        This suite combines powerful tools to help you create and communicate your message effectively:

        - **📖 Sermon Generator:** A collaborative team of AI agents to help you craft deep, biblically-sound, and engaging sermons.
        - **🚀 Flyer Production Studio:** An AI design agency that produces a stunning flyer image and compelling social media copy.
        - **🎶 Worship Song Studio:** An AI music collective that writes lyrics and creates a production-ready prompt for generative music AI.
        - **📚 Book Writing Studio:** Your personal AI writer's room to outline and draft a book in multiple languages.
        - **🌍 Bible Study Generator:** An AI team that creates in-depth, multilingual study guides for any book of the Bible.
        - **📰 Newsroom HQ:** Commission a complete digital newspaper with your own team of AI journalists.
        - **🎬 Viral Video Studio:** An AI creative team that concepts a powerful, 8-second vertical video and generates a prompt for Google VEO.

        ### How to Get Started:
        1.  **Configure Credentials:** Enter your **Gemini API Key** and **Serper API Key** in the sidebar. 
        2.  **Navigate:** Use the sidebar navigation to select the tool you want to use.
        3.  **Create:** Follow the instructions on each page to generate your content.
        """)
        st.markdown("---")

    elif selection == "Sermon Generator":
        render_sermon_page()
    elif selection == "Flyer Production Studio":
        render_flyer_page()
    elif selection == "Worship Song Studio":
        render_music_page()
    elif selection == "Book Writing Studio":
        render_book_page()
    elif selection == "Bible Study Generator":
        render_bible_study_page()
    elif selection == "Newsroom HQ":
        render_news_page()
    elif selection == "Viral Video Studio":
        render_viral_video_page()


if __name__ == "__main__":
    main()
